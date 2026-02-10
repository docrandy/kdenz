"""
Generate the corrected VCM Foundational Claims Document.
Fixes applied:
1. Count corrected from "61" to "56" (3 pre-volitional + 53 volitional)
2. Existence rows use Database Definitions (not Mechanisms)
3. Tables expanded from 4 rows to 6 rows (added Mechanism + Intervention)
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

# ============================================================
# ALL 56 ENTRIES: CORRECTED DATA
# Each entry: (id, name, definition, mechanism, gate_classification, decision_computation, intervention_axes, effect_size)
# definition = Database Definition → Existence claim (CORRECTED)
# mechanism = Database Mechanism → NEW separate Mechanism claim
# intervention_axes = Database Intervention Axes → NEW Intervention claim
# ============================================================

ENTRIES = {
    "C0": {
        "gate_name": "Environmental Permeability",
        "constraint_type": "Environment",
        "question": "Can the person physically act? Assesses whether structural, biological, or environmental conditions permit goal-directed action regardless of motivation.",
        "neural": "Sensory Cortex, Lateral PFC, Medial PFC, Dorsal Anterior Insula",
        "framing_note": "C0 is a pre-volitional screening gate. It rules out non-volitional blockers before diagnosing failure within the volitional chain (C1\u2013C6B). If C0 fails, the issue is environmental, biological, or structural\u2014not volitional.",
        "entries": [
            {
                "id": "C0-P",
                "name": "Physical/Biological Limitation",
                "definition": "Medical, neurological, or metabolic conditions block action regardless of motivation. Consolidates subtypes: Psychomotor Retardation (C0-3), Motor Readiness Failure (C0-4), Metabolic Insufficiency (C0-5), HPA Axis Exhaustion (C0-6).",
                "mechanism": "Structural biological constraint prevents motor output regardless of psychological state.",
                "gate_class": "Classified under C0 (Environmental Permeability) \u2014 Pre-volitional screening",
                "decision_comp": "IF Physical_Limitation = TRUE THEN Action_Blocked (regardless of motivation state)",
                "intervention": "Routes to medical evaluation. Not addressable through behavioral intervention. Subtypes: Psychomotor Retardation (dopamine/motor system slowing), Motor Readiness Failure (structural motor pathway damage), Metabolic Insufficiency (ATP/energy depletion), HPA Axis Exhaustion (autonomic burnout).",
                "effect_size": "N/A \u2014 routes to medical evaluation"
            },
            {
                "id": "C0-1",
                "name": "Environmental Friction",
                "definition": "Environmental friction (distance, setup time, clutter) raises initiation threshold above motor readiness.",
                "mechanism": "Environmental friction increases perceived effort cost beyond willingness-to-act threshold.",
                "gate_class": "Classified under C0 (Environmental Permeability) \u2014 Pre-volitional screening",
                "decision_comp": "IF Environmental_Friction > Activation_Threshold THEN Action_Blocked",
                "intervention": "Reduce activation cost, Environmental redesign, Reduce friction.",
                "effect_size": "d = 0.3\u20130.8"
            },
            {
                "id": "C0-2",
                "name": "Situational Blockage",
                "definition": "Structural, resource, or social constraints (time, money, access, legal, family opposition) make goal objectively non-viable.",
                "mechanism": "Objective structural constraints prevent action regardless of internal motivation state.",
                "gate_class": "Classified under C0 (Environmental Permeability) \u2014 Pre-volitional screening",
                "decision_comp": "IF Resource_Available = FALSE OR Access_Blocked = TRUE THEN Goal_Non_Viable",
                "intervention": "Change structure, Modify context, Reduce goal scope.",
                "effect_size": "d = 0.3\u20130.6 (depends on constraint type)"
            },
        ]
    },
    "C1": {
        "gate_name": "Believability",
        "constraint_type": "Perception",
        "question": "Does the person believe the goal is achievable? Assesses self-efficacy, perceived control, and belief formation processes.",
        "neural": "Ventral Striatum, vmPFC, Lateral PFC",
        "framing_note": None,
        "entries": [
            {
                "id": "C1-1",
                "name": "Cognitive Biases (Overconfidence/Planning Fallacy)",
                "definition": "Overconfident predictions override realistic assessment; planning fallacy creates prediction-reality gap.",
                "mechanism": "Inside view dominance; base-rate neglect inflates success expectations.",
                "gate_class": "Classified under C1 (Believability)",
                "decision_comp": "IF Confidence > Evidence_Base THEN Prediction_Error_HIGH",
                "intervention": "Rebuild efficacy, Increase specificity, Mastery experiences.",
                "effect_size": "d = 0.5\u20130.7 (medium to large)"
            },
            {
                "id": "C1-2",
                "name": "Fixed Mindset",
                "definition": "Belief that personal limitations are unchangeable biases efficacy updating downward.",
                "mechanism": "Entity theory of ability; perceived ceiling on growth biases efficacy updating downward, especially in at-risk populations.",
                "gate_class": "Classified under C1 (Believability)",
                "decision_comp": "IF Belief(Ability_Fixed) = TRUE THEN Efficacy_Updating_Biased_Downward",
                "intervention": "Rebuild efficacy, Mastery experiences, Reframe reward.",
                "effect_size": "d = 0.05\u20130.20"
            },
            {
                "id": "C1-3",
                "name": "No Relatable Models",
                "definition": "Absence of proximal role models who share similar background undermines vicarious efficacy.",
                "mechanism": "Vicarious efficacy pathway blocked; no similar-other success evidence available.",
                "gate_class": "Classified under C1 (Believability)",
                "decision_comp": "IF Proximal_Models = 0 THEN Vicarious_Efficacy = LOW",
                "intervention": "Rebuild efficacy, Increase value alignment.",
                "effect_size": "d = 0.15\u20130.35"
            },
            {
                "id": "C1-4",
                "name": "Meta-Cognitive Doubt",
                "definition": "Meta-level self-doubt ('I can't trust my own judgment') halts belief formation.",
                "mechanism": "Second-order doubt; questioning one's own cognitive processes prevents stable belief.",
                "gate_class": "Classified under C1 (Believability)",
                "decision_comp": "IF Trust(Own_Judgment) = LOW THEN Belief_Formation_Halted",
                "intervention": "Rebuild efficacy, Increase perceived control, Mastery experiences.",
                "effect_size": "d = 0.3\u20130.5 (small to medium)"
            },
            {
                "id": "C1-5",
                "name": "Insufficient Evidence",
                "definition": "Insufficient information/knowledge about goal feasibility prevents stable belief.",
                "mechanism": "Belief formation requires minimum evidence threshold; insufficient data prevents commitment.",
                "gate_class": "Classified under C1 (Believability)",
                "decision_comp": "IF Evidence(Feasibility) < Belief_Threshold THEN Belief_Unstable",
                "intervention": "Rebuild efficacy, Mastery experiences, Simulate context.",
                "effect_size": "d = 0.2\u20130.5 (high uncertainty)"
            },
            {
                "id": "C1-6",
                "name": "Fragile Confidence",
                "definition": "Confidence easily undermined by setbacks or external criticism; high sensitivity to negative feedback.",
                "mechanism": "Efficacy representation unstable; single negative signal can collapse entire belief structure.",
                "gate_class": "Classified under C1 (Believability)",
                "decision_comp": "IF Setback_Sensitivity = HIGH AND Setback_Occurs THEN Efficacy_Collapse",
                "intervention": "Rebuild efficacy, Reward persistence, Mastery experiences.",
                "effect_size": "d = 0.4\u20130.6"
            },
            {
                "id": "C1-7",
                "name": "Social Disconfirmation",
                "definition": "Social pressure or negative feedback undermines belief; external skepticism erodes self-efficacy through Bandura's verbal persuasion pathway.",
                "mechanism": "Verbal persuasion pathway (Bandura, 1977): negative social feedback degrades self-efficacy beliefs.",
                "gate_class": "Classified under C1 (Believability)",
                "decision_comp": "IF Social_Feedback = NEGATIVE THEN Verbal_Persuasion_Pathway_Activates; Efficacy_Erosion",
                "intervention": "Rebuild efficacy, Self-affirmation exercises, Increase perceived control, Change structure.",
                "effect_size": "d = 0.2\u20130.4"
            },
            {
                "id": "C1-8",
                "name": "Impostor Attribution Pattern",
                "definition": "Success attributed to external factors (luck, help) rather than internal capability; prevents efficacy consolidation. Related to but distinct from fundamental attribution error (Ross, 1977).",
                "mechanism": "External attribution bias; self-efficacy cannot update from success experiences. Foundational: Weiner (1985) attribution theory.",
                "gate_class": "Classified under C1 (Believability)",
                "decision_comp": "IF Attribution(Success) = EXTERNAL THEN Efficacy_Update_Blocked",
                "intervention": "Rebuild efficacy, Increase perceived control, Attribution retraining.",
                "effect_size": "d = 0.4\u20130.6 (estimated)"
            },
        ]
    },
    "C2": {
        "gate_name": "Desire",
        "constraint_type": "Desire",
        "question": "Does the person genuinely want the goal? Assesses intrinsic motivation, value alignment, and approach-avoidance dynamics.",
        "neural": "Midbrain Dopamine, Ventral Striatum, vmPFC",
        "framing_note": None,
        "entries": [
            {
                "id": "C2-1",
                "name": "Value Misalignment",
                "definition": "Goal pursued for extrinsic reasons, not authentic values; NAcc response blunted when goal conflicts with core values.",
                "mechanism": "Goal-value incongruence; vmPFC cannot assign authentic value to misaligned goal.",
                "gate_class": "Classified under C2 (Desire)",
                "decision_comp": "IF Goal_Value_Alignment < Threshold THEN Intrinsic_Motivation = LOW",
                "intervention": "Increase value alignment, Reframe reward.",
                "effect_size": "d = 0.4\u20130.7"
            },
            {
                "id": "C2-2",
                "name": "Fear-Suppressed Desire",
                "definition": "Desire exists but suppressed by anticipated loss or disruption; fear of consequences blocks approach motivation.",
                "mechanism": "Amygdala threat response overrides approach motivation; desire present but gated by fear.",
                "gate_class": "Classified under C2 (Desire)",
                "decision_comp": "IF Fear(Consequence) > Desire(Goal) THEN Approach_Blocked",
                "intervention": "Reduce cost, Reframe reward, Simulate context.",
                "effect_size": "d = 0.5\u20130.7 (medium to large)"
            },
            {
                "id": "C2-3",
                "name": "Emotional Avoidance",
                "definition": "Emotional discomfort (anxiety, shame, embarrassment) associated with goal blocks desire; avoidance system dominates.",
                "mechanism": "Experiential avoidance; emotional cost of engagement exceeds perceived benefit.",
                "gate_class": "Classified under C2 (Desire)",
                "decision_comp": "IF Emotional_Cost(Engagement) > Perceived_Benefit THEN Avoidance_Dominates",
                "intervention": "Reduce cost, Reframe reward, Simulate context.",
                "effect_size": "d = 0.5\u20130.8 (medium to large)"
            },
            {
                "id": "C2-4",
                "name": "Autonomy Threat (Reactance)",
                "definition": "External pressure or perceived control loss triggers oppositional motivation; reactance activates competing desire for freedom restoration (boomerang effect).",
                "mechanism": "Psychological reactance; perceived freedom threat generates competing desire for freedom restoration. Brehm (1966): desire ENHANCEMENT, not suppression.",
                "gate_class": "Classified under C2 (Desire)",
                "decision_comp": "IF Perceived_Control_Loss = TRUE THEN Competing_Desire_Activated(Freedom_Restoration)",
                "intervention": "Increase perceived control, Autonomy support, Reduce external pressure, Change structure.",
                "effect_size": "d = 0.3\u20130.5"
            },
            {
                "id": "C2-5",
                "name": "Approach-Avoidance Conflict",
                "definition": "Simultaneous desire and anti-desire cancel motivation; net motivation = 0.",
                "mechanism": "Dual-process conflict; approach and avoidance systems simultaneously activated.",
                "gate_class": "Classified under C2 (Desire)",
                "decision_comp": "IF Approach_Force \u2248 Avoidance_Force THEN Oscillation + Anxiety + BIS_Activation",
                "intervention": "Reduce cost, Reframe reward, Increase value alignment.",
                "effect_size": "d = 0.4\u20130.6 (medium)"
            },
            {
                "id": "C2-6",
                "name": "External Motivation",
                "definition": "Goal pursued exclusively for approval/reward from others, not intrinsic value.",
                "mechanism": "SDT framework; external rewards undermine internalization and intrinsic motivation.",
                "gate_class": "Classified under C2 (Desire)",
                "decision_comp": "IF (Intrinsic_Motivation = 0) AND (Only_External_Rewards = TRUE) THEN Motivation_Fragile",
                "intervention": "Increase value alignment, Autonomy support, Internalization facilitation.",
                "effect_size": "d = 0.3\u20130.5"
            },
            {
                "id": "C2-7",
                "name": "Purpose Deficit",
                "definition": "No sense of purpose or meaning attached to goal; subjective value = 0 so vmPFC cannot assign value.",
                "mechanism": "Meaning-making failure; goal lacks connection to superordinate purpose structure.",
                "gate_class": "Classified under C2 (Desire)",
                "decision_comp": "IF Subjective_Value(Goal) = 0 THEN vmPFC_Value_Assignment = FAILED",
                "intervention": "Increase value alignment, Reframe reward.",
                "effect_size": "d = 0.3\u20130.7"
            },
            {
                "id": "C2-8",
                "name": "Lack of Future Relevance",
                "definition": "Goal does not connect to long-term identity or future self.",
                "mechanism": "Temporal discounting inflates distance penalty; future benefit psychologically distant.",
                "gate_class": "Classified under C2 (Desire)",
                "decision_comp": "IF Temporal_Distance(Goal) > Discounting_Threshold THEN Present_Motivation = LOW",
                "intervention": "Increase value alignment, Reframe reward.",
                "effect_size": "d = 0.3\u20130.5 (small to medium)"
            },
        ]
    },
    "C3": {
        "gate_name": "Will",
        "constraint_type": "Will",
        "question": "Is the person willing to bear the cost? Assesses effort computation, discomfort tolerance, and cost-benefit analysis at the commitment threshold.",
        "neural": "dACC, Anterior Insula, dlPFC",
        "framing_note": None,
        "entries": [
            {
                "id": "C3-1",
                "name": "Self-Control Fatigue",
                "definition": "Reduced willingness to engage effortful control following prolonged cognitive demands; reflects motivational shift toward less effortful options rather than capacity loss.",
                "mechanism": "Process model (Inzlicht & Schmeichel): prolonged self-control shifts motivation from 'have-to' to 'want-to' tasks. Motivational reallocation, not resource depletion.",
                "gate_class": "Classified under C3 (Will)",
                "decision_comp": "IF Sustained_Effort_Duration > Fatigue_Threshold THEN Motivation_Shifts_To_Less_Effortful_Options",
                "intervention": "Effort pacing, Energy pacing, Reduce goal scope, Increase task engagement/autonomy.",
                "effect_size": "d = 0.31\u20130.35 with intensive protocols (Dang et al. 2025); near-zero with brief protocols"
            },
            {
                "id": "C3-2",
                "name": "Delayed Gratification Intolerance",
                "definition": "Unable to tolerate delay between effort and reward; steep temporal discounting makes present cost unacceptable.",
                "mechanism": "Hyperbolic discounting; future reward drastically devalued against immediate effort cost.",
                "gate_class": "Classified under C3 (Will)",
                "decision_comp": "IF Discount_Rate(Reward) > Effort_Cost_Tolerance THEN Commitment_Rejected",
                "intervention": "Mini-deadlines, Reframe reward, Increase intrinsic cues.",
                "effect_size": "d = 0.4\u20130.6 (clinical); r = 0.10\u20130.20 (prediction)"
            },
            {
                "id": "C3-3",
                "name": "Effort Overestimation",
                "definition": "Perceived effort/discomfort judged excessive relative to reward; catastrophic effort estimation blocks commitment.",
                "mechanism": "Effort computation error; insula and ACC over-represent anticipated discomfort, leading to inflated effort cost estimates.",
                "gate_class": "Classified under C3 (Will)",
                "decision_comp": "IF Perceived_Effort >> Actual_Effort THEN Commitment_Blocked",
                "intervention": "Reduce cost, Micro-starts, Simulate context.",
                "effect_size": "FLAG: estimated composite, no meta-analytic anchor"
            },
            {
                "id": "C3-4",
                "name": "Insufficient Perceived Reward",
                "definition": "Expected reward is not salient or valued enough to justify effort.",
                "mechanism": "Effort-reward tradeoff: dopamine's role is behavioral activation and effort allocation (Salamone), not hedonic 'liking' (Berridge framing).",
                "gate_class": "Classified under C3 (Will)",
                "decision_comp": "IF NAcc_Activation(Reward) < Effort_Cost THEN Insufficient_Motivation",
                "intervention": "Reframe reward, Increase perceived control.",
                "effect_size": "d = 0.3\u20130.5"
            },
            {
                "id": "C3-5",
                "name": "Failure Cost Aversion",
                "definition": "Anticipated failure cost (emotional, identity) exceeds success reward; expected utility negative.",
                "mechanism": "Loss aversion at identity level; failure cost weighted more heavily than success reward.",
                "gate_class": "Classified under C3 (Will)",
                "decision_comp": "IF Expected_Failure_Cost > Expected_Success_Reward THEN Net_Utility_Negative",
                "intervention": "Reduce cost, Reframe reward, Rebuild efficacy.",
                "effect_size": "d = 0.5\u20130.7 (medium to large)"
            },
            {
                "id": "C3-6",
                "name": "Decision Paralysis (Analysis Paralysis)",
                "definition": "Multiple option combinations create intractable decision problem; endless deliberation loop prevents commitment.",
                "mechanism": "Combinatorial explosion in option space; optimization attempt exceeds cognitive capacity.",
                "gate_class": "Classified under C3 (Will)",
                "decision_comp": "IF Options > Satisficing_Threshold THEN Deliberation_Loop_Infinite",
                "intervention": "Reduce goal scope, Increase specificity, Change timing.",
                "effect_size": "d = 0.0\u20130.5 (highly conditional)"
            },
        ]
    },
    "C4": {
        "gate_name": "Intention",
        "constraint_type": "Intention",
        "question": "Has the person formed a concrete plan? Assesses plan specificity, implementation intentions, and cue-response binding.",
        "neural": "dlPFC, Posterior Parietal Cortex, Mid-lateral PFC",
        "framing_note": None,
        "entries": [
            {
                "id": "C4-1",
                "name": "Conflicting Plans",
                "definition": "Multiple commitments/goals creating priority conflict; attention fragments across incompatible demands.",
                "mechanism": "Executive resource competition; multiple active plans exceed allocation capacity.",
                "gate_class": "Classified under C4 (Intention)",
                "decision_comp": "IF Active_Plans > Executive_Capacity THEN Priority_Fragmentation",
                "intervention": "Increase structure, Reduce goal scope, Change timing.",
                "effect_size": "d = 0.4\u20130.6 (medium)"
            },
            {
                "id": "C4-2",
                "name": "Chronic Replanning",
                "definition": "Ongoing deliberation and planning prevents action; continuous replanning blocks execution.",
                "mechanism": "Planning becomes substitute for action; dopaminergic reward from planning itself.",
                "gate_class": "Classified under C4 (Intention)",
                "decision_comp": "IF Replanning_Frequency > Action_Frequency THEN Execution_Blocked",
                "intervention": "Increase specificity, Use implementation intentions, Mini-deadlines.",
                "effect_size": "d = 0.5\u20130.7 (medium to large)"
            },
            {
                "id": "C4-3",
                "name": "Plan Selection Failure",
                "definition": "Committed to goal but cannot select among viable plan options; implementation stalled despite endorsement.",
                "mechanism": "Option evaluation deadlock; multiple viable paths create selection paralysis.",
                "gate_class": "Classified under C4 (Intention)",
                "decision_comp": "IF Viable_Options > 1 AND Selection_Criteria = UNCLEAR THEN Stalled",
                "intervention": "Increase specificity, Reduce goal scope, Change structure.",
                "effect_size": "d = 0.4\u20130.6 (medium)"
            },
            {
                "id": "C4-4",
                "name": "Cognitive Overload",
                "definition": "Plan complexity exceeds working memory capacity.",
                "mechanism": "Working memory capacity exceeded, plan coherence impossible; information processing saturated.",
                "gate_class": "Classified under C4 (Intention)",
                "decision_comp": "IF Plan_Complexity > Working_Memory_Capacity THEN Coherence_Impossible",
                "intervention": "Increase specificity, Habit stacking.",
                "effect_size": "d = 0.6\u20130.9 (large)"
            },
            {
                "id": "C4-5",
                "name": "Action Ambiguity",
                "definition": "Vague about what action to take right now; next-step specificity below threshold blocks intention formation.",
                "mechanism": "Insufficient action specification; motor system cannot activate without concrete target.",
                "gate_class": "Classified under C4 (Intention)",
                "decision_comp": "IF Next_Step_Specificity < Action_Threshold THEN Intention_Formation_Blocked",
                "intervention": "Increase specificity, Use implementation intentions, Immediate cues.",
                "effect_size": "d = 0.6\u20130.9 (large)"
            },
            {
                "id": "C4-6",
                "name": "Recursive Justification",
                "definition": "Repeatedly generating new justifications to delay or modify plan; recursive justification delays commitment.",
                "mechanism": "Self-justification loop; each justification generates need for further justification.",
                "gate_class": "Classified under C4 (Intention)",
                "decision_comp": "IF Justification_Count > Threshold THEN Action_Delayed_Indefinitely",
                "intervention": "Pre-commitment, Mini-deadlines, Ulysses pact.",
                "effect_size": "d = 0.4\u20130.6 (medium)"
            },
            {
                "id": "C4-7",
                "name": "Weak Cue-Response Binding",
                "definition": "Automatic behavior initiation from situational cues fails; cue-response binding strength below threshold.",
                "mechanism": "Implementation intention not encoded; situation-action link too weak for automatic activation.",
                "gate_class": "Classified under C4 (Intention)",
                "decision_comp": "IF Cue_Response_Strength < Automatic_Threshold THEN Manual_Initiation_Required",
                "intervention": "Habit stacking, Use implementation intentions, Environmental redesign.",
                "effect_size": "d = 0.6\u20130.8 (large)"
            },
        ]
    },
    "C5": {
        "gate_name": "Commitment",
        "constraint_type": "Commitment",
        "question": "Is the person's commitment robust? Assesses goal shielding, priority maintenance, and resistance to competing demands.",
        "neural": "Anterior Cingulum Bundle, dlPFC, vmPFC",
        "framing_note": None,
        "entries": [
            {
                "id": "C5-1",
                "name": "Competing Commitments",
                "definition": "Multiple conflicting commitments create chronic attentional and motivational competition; competing goals win the value computation, diverting effort from primary goal.",
                "mechanism": "Motivational and attentional competition: self-control failures occur because competing goals win the value computation, not because a resource tank ran empty.",
                "gate_class": "Classified under C5 (Commitment)",
                "decision_comp": "IF Competing_Goal_Value > Primary_Goal_Value THEN Attention_Diverted; Primary_Goal_Undermined",
                "intervention": "Increase structure, Reduce goal scope, Add friction to alternatives, Goal prioritization exercises.",
                "effect_size": "d = 0.4\u20130.7"
            },
            {
                "id": "C5-2",
                "name": "Environmental Disruption",
                "definition": "Physical or social environment disrupts goal pursuit; environmental trigger salience overrides goal salience.",
                "mechanism": "Environmental cue competition; salient environmental triggers override internal goal representations.",
                "gate_class": "Classified under C5 (Commitment)",
                "decision_comp": "IF Environmental_Trigger_Salience > Goal_Salience THEN Goal_Overridden",
                "intervention": "Environmental control, Alter environment, Add friction to alternatives.",
                "effect_size": "d = 0.4\u20130.7"
            },
            {
                "id": "C5-3",
                "name": "No External Monitoring",
                "definition": "No external pressure or visibility regarding commitment; lack of monitoring weakens felt obligation.",
                "mechanism": "Social accountability absent; without observation, commitment maintenance degrades.",
                "gate_class": "Classified under C5 (Commitment)",
                "decision_comp": "IF External_Monitoring = 0 THEN Commitment_Decay_Rate = HIGH",
                "intervention": "Pre-commitment, Monitor progress, Raise abandonment cost.",
                "effect_size": "d = 0.32\u20130.48"
            },
            {
                "id": "C5-4",
                "name": "Insufficient Stakes",
                "definition": "No real penalty for non-compliance; consequence magnitude below motivation threshold allows abandonment.",
                "mechanism": "Low abandonment cost; nothing significant lost by quitting.",
                "gate_class": "Classified under C5 (Commitment)",
                "decision_comp": "IF Consequence(Abandonment) < Motivation_Threshold THEN Abandonment_Easy",
                "intervention": "Raise abandonment cost, Pre-commitment, Ulysses pact.",
                "effect_size": "d = 0.2\u20130.7"
            },
            {
                "id": "C5-5",
                "name": "Lack of Urgency",
                "definition": "Distant deadline or unclear deadline creates low-pressure environment; urgency perception insufficient.",
                "mechanism": "Temporal distance reduces motivational salience; low urgency permits indefinite delay.",
                "gate_class": "Classified under C5 (Commitment)",
                "decision_comp": "IF Deadline_Distance > Urgency_Threshold THEN Motivational_Salience = LOW",
                "intervention": "Mini-deadlines, Change timing, Immediate cues.",
                "effect_size": "d = 0.4\u20130.6 (medium)"
            },
            {
                "id": "C5-6",
                "name": "Value Contradiction",
                "definition": "Goal conflicts with another held value; value conflict disrupts goal protection.",
                "mechanism": "Value system inconsistency; competing value activation undermines commitment stability.",
                "gate_class": "Classified under C5 (Commitment)",
                "decision_comp": "IF Goal_Value CONFLICTS_WITH Held_Value THEN Commitment_Unstable",
                "intervention": "Increase value alignment, Reframe reward, Reduce cost.",
                "effect_size": "d = 0.3\u20130.6"
            },
            {
                "id": "C5-7",
                "name": "Identity Misalignment",
                "definition": "Goal inconsistent with identity or self-concept; triggers identity protection against goal-inconsistent behavior.",
                "mechanism": "Self-concept threat; goal pursuit threatens established identity narrative.",
                "gate_class": "Classified under C5 (Commitment)",
                "decision_comp": "IF Goal INCONSISTENT_WITH Self_Concept THEN Identity_Protection_Activated",
                "intervention": "Increase value alignment, Reframe reward, Rebuild efficacy.",
                "effect_size": "d = 0.4\u20130.7"
            },
            {
                "id": "C5-8",
                "name": "Identity-Behavior Dissonance",
                "definition": "Sustained incongruence between stated goal and actual behavior creates chronic self-concept conflict.",
                "mechanism": "Cognitive dissonance accumulation; persistent gap between aspiration and action erodes commitment.",
                "gate_class": "Classified under C5 (Commitment)",
                "decision_comp": "IF (Stated_Goal - Actual_Behavior) > Dissonance_Threshold FOR Duration > T THEN Commitment_Erosion",
                "intervention": "Self-compassion training, Goal revision/downsizing, Micro-starts, Monitor progress.",
                "effect_size": "d = 0.4\u20130.6 (ESTIMATED/INFERRED)"
            },
        ]
    },
    "C6A": {
        "gate_name": "Action Initiation",
        "constraint_type": "Action Initiation",
        "question": "Can the person start? Assesses motor initiation, threshold activation, and barriers at the moment of action.",
        "neural": "SMA, Pre-SMA, ACC motor zone, Basal Ganglia",
        "framing_note": None,
        "entries": [
            {
                "id": "C6A-1",
                "name": "Motor Initiation Threshold",
                "definition": "Motor threshold to begin action is high; basal ganglia gate locked, inertia dominates.",
                "mechanism": "Motor initiation threshold elevated; striatal gate requires higher signal to open.",
                "gate_class": "Classified under C6A (Action Initiation)",
                "decision_comp": "IF Motor_Signal < Activation_Threshold THEN Basal_Ganglia_Gate_Locked",
                "intervention": "Reduce activation cost, Micro-starts, Reduce inertia.",
                "effect_size": "d = 0.4\u20130.8"
            },
            {
                "id": "C6A-2",
                "name": "Last-Minute Deliberation",
                "definition": "Continued deliberation at moment of action prevents initiation; deliberation delays motor execution.",
                "mechanism": "Executive override of motor system; prefrontal deliberation maintains inhibitory control at action moment.",
                "gate_class": "Classified under C6A (Action Initiation)",
                "decision_comp": "IF Deliberation_Active_At_Action_Moment THEN Motor_Execution_Delayed",
                "intervention": "Pre-commitment, Immediate cues, Use implementation intentions.",
                "effect_size": "d = 0.4\u20130.6 (medium)"
            },
            {
                "id": "C6A-3",
                "name": "Attention Diversion",
                "definition": "Attention diverted immediately before action moment; competing stimulus salience captures attention at critical initiation moment.",
                "mechanism": "Attentional capture by competing stimuli at the precise moment of intended action initiation.",
                "gate_class": "Classified under C6A (Action Initiation)",
                "decision_comp": "IF Competing_Stimulus_Salience > Goal_Salience AT Action_Moment THEN Initiation_Blocked",
                "intervention": "Environmental control, Add friction to alternatives, Reduce friction.",
                "effect_size": "d = 0.5\u20130.7"
            },
            {
                "id": "C6A-4",
                "name": "Waiting for Readiness",
                "definition": "Waits for ideal conditions/mood/readiness to start; unmet readiness conditions block action.",
                "mechanism": "Conditional action policy; self-imposed prerequisites prevent initiation.",
                "gate_class": "Classified under C6A (Action Initiation)",
                "decision_comp": "IF Readiness_Conditions_Met = FALSE THEN Action_Postponed",
                "intervention": "Reduce activation cost, Immediate cues, Micro-starts.",
                "effect_size": "d = 0.4\u20130.7"
            },
            {
                "id": "C6A-5",
                "name": "Over-Simulation",
                "definition": "Over-rehearsal/mental practice prevents moving to actual action; outcome simulation reduces physiological energization needed for action.",
                "mechanism": "Positive outcome fantasies reduce physiological energization (Kappes & Oettingen, 2011). Critical: OUTCOME simulation reduces effort; PROCESS simulation increases effort.",
                "gate_class": "Classified under C6A (Action Initiation)",
                "decision_comp": "IF Outcome_Simulation_Count > Threshold THEN Physiological_Energization_Reduced; Action_Drive_Below_Threshold",
                "intervention": "Mental contrasting (WOOP), Process simulation instead of outcome simulation, Immediate cues, Reduce inertia.",
                "effect_size": "d = 0.3\u20130.5"
            },
            {
                "id": "C6A-6",
                "name": "Self-Sabotage",
                "definition": "Unconscious sabotage creates barrier to action; defensive avoidance creates self-imposed barriers.",
                "mechanism": "Protective self-handicapping; creating obstacles preserves self-concept against potential failure.",
                "gate_class": "Classified under C6A (Action Initiation)",
                "decision_comp": "IF Failure_Threat = HIGH THEN Create_Self_Imposed_Barriers",
                "intervention": "Rebuild efficacy, Reduce cost, Increase perceived control.",
                "effect_size": "d = 0.4\u20130.7 (medium to large)"
            },
            {
                "id": "C6A-7",
                "name": "Momentum Dependency",
                "definition": "Cannot cold-start despite desire/planning/capability; continuation easy once momentum achieved. Common in ADHD (low tonic dopamine). Initiation threshold >> continuation threshold.",
                "mechanism": "Dopaminergic initiation deficit; tonic dopamine insufficient for cold-start but adequate for continuation.",
                "gate_class": "Classified under C6A (Action Initiation)",
                "decision_comp": "IF Initiation_Threshold >> Continuation_Threshold THEN Cold_Start_Failure",
                "intervention": "Micro-starts, Reduce inertia, Environmental redesign.",
                "effect_size": "d = 0.4\u20130.8"
            },
            {
                "id": "C6A-8",
                "name": "Trauma-Related Freeze",
                "definition": "Threat-protective immobility; defense cascade via amygdala-PAG circuitry overrides volitional control.",
                "mechanism": "Defense cascade via amygdala-periaqueductal gray (PAG) circuitry: threat detection triggers escalating defensive responses (arousal \u2192 fight/flight \u2192 freeze \u2192 tonic immobility \u2192 collapse).",
                "gate_class": "Classified under C6A (Action Initiation)",
                "decision_comp": "IF Threat_Detection = TRUE THEN Amygdala_PAG_Defense_Cascade_Activated; Volitional_Control_Overridden",
                "intervention": "Requires trauma-informed therapy (EMDR, SE, CPT). Not addressable through behavioral intervention alone.",
                "effect_size": "Clinical \u2014 requires specialized treatment"
            },
        ]
    },
    "C6B": {
        "gate_name": "Action Persistence",
        "constraint_type": "Action Persistence",
        "question": "Can the person sustain action? Assesses effort regulation, distraction resistance, and persistence through obstacles once action has begun.",
        "neural": "Dorsal Striatum, ACC Persistence Circuit, Motor Cortex",
        "framing_note": None,
        "entries": [
            {
                "id": "C6B-1",
                "name": "Task Boredom",
                "definition": "Activity becomes monotonous; hedonic adaptation degrades reward salience, interest wanes mid-task.",
                "mechanism": "Hedonic adaptation; reward signal attenuates with repetition reducing continuation motivation.",
                "gate_class": "Classified under C6B (Action Persistence)",
                "decision_comp": "IF Reward_Salience(Task) < Continuation_Threshold THEN Engagement_Drops",
                "intervention": "Increase intrinsic cues, Reframe reward, Change timing.",
                "effect_size": "d = 0.4\u20130.6 (medium)"
            },
            {
                "id": "C6B-2",
                "name": "Competing Stimuli (Distraction)",
                "definition": "Susceptible to pulling away from task by external/internal temptations; alternative NAcc activation disrupts goal-shielding. Includes both involuntary attentional capture (notifications, novel stimuli) and voluntary comfort-seeking (choosing easier/more pleasant alternatives such as social media, food, rest).",
                "mechanism": "Attentional competition; alternative reward sources capture attention away from current task.",
                "gate_class": "Classified under C6B (Action Persistence)",
                "decision_comp": "IF Alternative_NAcc_Activation > Task_NAcc_Activation THEN Goal_Shielding_Fails",
                "intervention": "Environmental control, Add friction to alternatives, Increase structure.",
                "effect_size": "d = 0.5\u20130.7"
            },
            {
                "id": "C6B-3",
                "name": "Mood-Dependent Execution",
                "definition": "Action output fluctuates with daily mood/energy; execution contingent on mood state.",
                "mechanism": "Mood-contingent action policy; no decoupling between affective state and behavioral output.",
                "gate_class": "Classified under C6B (Action Persistence)",
                "decision_comp": "IF Mood_State = LOW THEN Action_Output = LOW (regardless of plan)",
                "intervention": "Use implementation intentions, Increase structure, Effort pacing.",
                "effect_size": "d = 0.4\u20130.6 (medium)"
            },
            {
                "id": "C6B-4",
                "name": "Single-Error Catastrophizing",
                "definition": "Single small error triggers premature task abandonment; catastrophizing triggers abandonment.",
                "mechanism": "All-or-nothing thinking; single deviation from plan treated as total failure.",
                "gate_class": "Classified under C6B (Action Persistence)",
                "decision_comp": "IF Error_Count >= 1 THEN Catastrophize; Abandon_Task",
                "intervention": "Reframe reward, Cost management, Monitor progress.",
                "effect_size": "d = 0.4\u20130.7 (medium to large)"
            },
            {
                "id": "C6B-5",
                "name": "Perfectionism",
                "definition": "Unwillingness to continue with imperfect output; perfectionistic standards halt progress.",
                "mechanism": "Quality threshold set above achievable level; imperfection intolerance blocks continuation.",
                "gate_class": "Classified under C6B (Action Persistence)",
                "decision_comp": "IF Output_Quality < Perfection_Standard THEN Halt_Progress",
                "intervention": "Reduce cost, Reframe reward, Mini-deadlines.",
                "effect_size": "d = 0.4\u20130.7 (medium to large)"
            },
            {
                "id": "C6B-6",
                "name": "Invisible Progress",
                "definition": "Motivation drops when progress is invisible or slower than expected; motivation contingent on visible progress signal.",
                "mechanism": "Progress feedback loop broken; without visible progress, continuation motivation decays.",
                "gate_class": "Classified under C6B (Action Persistence)",
                "decision_comp": "IF Visible_Progress < Expected_Progress THEN Motivation_Decay",
                "intervention": "Progress feedback, Monitor progress, Mini-deadlines.",
                "effect_size": "d = 0.4\u20130.7"
            },
            {
                "id": "C6B-7",
                "name": "High Restart Cost",
                "definition": "Restarting task after interruption/break is as difficult as initial start; momentum loss creates high restart cost.",
                "mechanism": "Context-switching cost; mental state reconstruction after interruption requires full re-initiation.",
                "gate_class": "Classified under C6B (Action Persistence)",
                "decision_comp": "IF Restart_Cost ~ Initiation_Cost THEN Resumption_Blocked",
                "intervention": "Reduce inertia, Micro-starts, Increase structure.",
                "effect_size": "d = 0.5\u20130.7 (medium to large)"
            },
            {
                "id": "C6B-8",
                "name": "Shame-Driven Withdrawal",
                "definition": "Shame from perceived failure triggers withdrawal and avoidance instead of continued effort. Distinct from guilt, which motivates reparative action.",
                "mechanism": "Shame-driven avoidance: perceived failure activates shame (global self-evaluation: 'I am bad'), triggering withdrawal from the goal domain. NOT guilt ('I did something bad'), which motivates repair.",
                "gate_class": "Classified under C6B (Action Persistence)",
                "decision_comp": "IF Perceived_Failure = TRUE AND Shame_Proneness = HIGH THEN Avoidance_Response_Activated; Withdrawal_From_Goal_Domain",
                "intervention": "Self-compassion training (primary), Error normalization, Cost management, Monitor progress.",
                "effect_size": "d = 0.3\u20130.5"
            },
        ]
    },
}

# ============================================================
# DOCUMENT GENERATION
# ============================================================

def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_paragraph(doc, text, bold=False, italic=False, style=None):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    return p

def add_claim_table(doc, entry):
    """Add a 6-row corrected claim table for a root cause entry."""
    table = doc.add_table(rows=7, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    headers = ['Claim Element', 'Content', 'Status']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = 'Calibri'

    # Data rows
    rows_data = [
        ('Existence', entry['definition'], 'PENDING'),
        ('Gate Classification', entry['gate_class'], 'PENDING'),
        ('Mechanism', entry['mechanism'], 'PENDING'),
        ('Decision Computation', entry['decision_comp'], 'PENDING'),
        ('Intervention', entry['intervention'], 'PENDING'),
        ('Effect Size', entry['effect_size'], 'PENDING'),
    ]

    for row_idx, (element, content, status) in enumerate(rows_data, start=1):
        table.rows[row_idx].cells[0].text = ''
        table.rows[row_idx].cells[1].text = ''
        table.rows[row_idx].cells[2].text = ''

        p0 = table.rows[row_idx].cells[0].paragraphs[0]
        run0 = p0.add_run(element)
        run0.bold = True
        run0.font.size = Pt(10)
        run0.font.name = 'Calibri'

        p1 = table.rows[row_idx].cells[1].paragraphs[0]
        run1 = p1.add_run(content)
        run1.font.size = Pt(10)
        run1.font.name = 'Calibri'

        p2 = table.rows[row_idx].cells[2].paragraphs[0]
        run2 = p2.add_run(status)
        run2.font.size = Pt(10)
        run2.font.name = 'Calibri'
        run2.font.color.rgb = RGBColor(180, 120, 0)  # amber for pending

    # Set column widths
    for row in table.rows:
        row.cells[0].width = Inches(1.5)
        row.cells[1].width = Inches(4.5)
        row.cells[2].width = Inches(1.0)

    doc.add_paragraph()  # spacing


def add_neural_table(doc):
    """Add the neural validation status table."""
    table = doc.add_table(rows=9, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['Gate', 'Neural Systems', 'Validation', 'Key Evidence']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = 'Calibri'

    neural_data = [
        ('C0', 'Sensory Cortex, Lateral PFC, Medial PFC, Dorsal Anterior Insula', 'VALIDATED', 'Environmental perception, interoception circuits well-established'),
        ('C1', 'Ventral Striatum, vmPFC, Lateral PFC', 'VALIDATED', 'Self-efficacy and belief formation literature extensive (Bandura \u2192 imaging studies)'),
        ('C2', 'Midbrain Dopamine, Ventral Striatum, vmPFC', 'VALIDATED', 'Reward valuation circuitry among most-studied in neuroscience'),
        ('C3', 'dACC, Anterior Insula, dlPFC', 'VALIDATED', 'Effort cost evaluation: Neuron, Nature Neuroscience confirm dACC role'),
        ('C4', 'dlPFC, Posterior Parietal Cortex, Mid-lateral PFC', 'VALIDATED', 'Neuron (2009), J Neuroscience (2010, 2020) \u2014 intention formation circuits'),
        ('C5', 'Anterior Cingulum Bundle, dlPFC, vmPFC', 'VALIDATED', 'Nature Human Behaviour (2024, Holton et al.) \u2014 vmPFC drives commitment'),
        ('C6A', 'SMA, Pre-SMA, ACC motor zone, Basal Ganglia', 'VALIDATED', 'Motor initiation literature, basal ganglia gating well-established'),
        ('C6B', 'Dorsal Striatum, ACC Persistence Circuit, Motor Cortex', 'VALIDATED', 'Habit circuits, dorsal striatum persistence \u2014 multiple PMC reviews'),
    ]

    for row_idx, (gate, neural, valid, evidence) in enumerate(neural_data, start=1):
        for col_idx, text in enumerate([gate, neural, valid, evidence]):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = ''
            run = cell.paragraphs[0].add_run(text)
            run.font.size = Pt(10)
            run.font.name = 'Calibri'
            if col_idx == 2:  # Validation column
                run.font.color.rgb = RGBColor(0, 128, 0)  # green for validated

    doc.add_paragraph()


def add_rating_table(doc):
    """Add the evidence rating scale table."""
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['Rating', 'Criteria', 'Action']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = 'Calibri'

    ratings = [
        ('STRONG', 'Multiple peer-reviewed studies directly supporting the claim. Mechanism well-established in neuroscience or behavioral literature.', 'Validated'),
        ('MODERATE', 'At least one direct study or multiple indirect studies. Mechanism plausible and consistent with established theory.', 'Validated with note'),
        ('WEAK', 'Only indirect evidence or clinical observation. No direct empirical test of this specific mechanism.', 'Flag for research'),
        ('UNSUPPORTED', 'No evidence found. Claim may be theoretically plausible but untested.', 'Requires revision'),
        ('CONTRADICTED', 'Available evidence contradicts the claim.', 'Remove or restructure'),
    ]

    for row_idx, (rating, criteria, action) in enumerate(ratings, start=1):
        for col_idx, text in enumerate([rating, criteria, action]):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = ''
            run = cell.paragraphs[0].add_run(text)
            run.font.size = Pt(10)
            run.font.name = 'Calibri'

    doc.add_paragraph()


def build_document():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # ---- TITLE PAGE ----
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_heading('VOLITIONAL CHAIN MODEL', level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)

    add_paragraph(doc, 'Foundational Claims Document', bold=True)
    add_paragraph(doc, 'Validation Foundation for 56 Root Causes Across 8 Constraint Gates')
    add_paragraph(doc, '3 Pre-Volitional Screening Entries (C0) + 53 Volitional Root Causes (C1\u2013C6B)')
    doc.add_paragraph()
    add_paragraph(doc, 'February 2026 \u2014 CONFIDENTIAL')
    add_paragraph(doc, 'Nunez Medical Services Inc. / MindMetrics\u2122')
    doc.add_page_break()

    # ---- SECTION 1: PURPOSE ----
    add_heading(doc, '1. PURPOSE OF THIS DOCUMENT', level=1)
    add_paragraph(doc, 'This document articulates every testable claim made by the Volitional Chain Model (VCM). It serves as the validation foundation \u2014 the complete set of assertions that must be supported by evidence for VCM to hold. Each claim is categorized by type, scope, and the burden of proof required.')
    add_paragraph(doc, 'Three levels of claims are documented: (1) Framework Architecture \u2014 structural claims about how the gating system works, (2) Gate-Level Claims \u2014 what each constraint gate does and which neural systems support it, and (3) Root-Cause-Level Claims \u2014 the 56 specific failure modes including their existence as clinical patterns, mechanisms, decision computations, intervention predictions, and effect sizes.')
    add_paragraph(doc, 'Structural note: The VCM distinguishes between pre-volitional screening (C0: Environmental Permeability) and the volitional chain proper (C1\u2013C6B). C0 rules out non-volitional blockers \u2014 biological, environmental, and structural constraints \u2014 before diagnosing volitional failure. The 53 root causes across gates C1\u2013C6B constitute the core diagnostic taxonomy for conscious volitional behavior.')

    # ---- SECTION 2: FRAMEWORK ARCHITECTURE CLAIMS ----
    add_heading(doc, '2. FRAMEWORK ARCHITECTURE CLAIMS', level=1)
    add_paragraph(doc, 'These are the structural claims about how VCM operates as a system. They are the highest-stakes claims because if any fails, the entire framework requires restructuring.')

    add_heading(doc, 'Claim A1: Sequential Gating with Binary Pass/Fail', level=2)
    add_paragraph(doc, 'Statement: The volitional chain proceeds through eight gates in fixed order (C0 \u2192 C1 \u2192 C2 \u2192 C3 \u2192 C4 \u2192 C5 \u2192 C6A \u2192 C6B). If any gate evaluates to 0 (fail), no downstream activity occurs at that moment. The behavior will not happen unless the failed gate is subsequently activated.')
    add_paragraph(doc, 'Claim Type: Structural / Architectural')
    add_paragraph(doc, 'Strength: Strong claim \u2014 this is the central organizing principle of VCM')
    add_paragraph(doc, 'Validation Requirement: Evidence that volitional processing involves ordered stages where upstream failures prevent downstream engagement. Supporting literature: Heckhausen & Gollwitzer\'s Rubicon Model (phases of action), Shallice\'s Supervisory Attentional System (hierarchical control), Botvinick\'s hierarchical reinforcement learning.')
    add_paragraph(doc, 'Key Nuance: The claim is about momentary evaluation. A person may pass gates 0\u20134, fail at C5, and 30 minutes later re-engage and pass C5. The chain resets per volitional episode, not permanently.')

    add_heading(doc, 'Claim A2: Dominant Root Cause per Episode', level=2)
    add_paragraph(doc, 'Statement: While multiple factors may contribute to a failed gate, there is a dominant root cause in each moment of failure. Different attempts at the same behavior can have different dominant root causes (e.g., first attempt: high friction; second attempt: competing beliefs).')
    add_paragraph(doc, 'Claim Type: Diagnostic / Clinical')
    add_paragraph(doc, 'Strength: Moderate claim \u2014 clinically pragmatic, allows targeted intervention')
    add_paragraph(doc, 'Validation Requirement: Evidence that motivational failures can be decomposed into primary vs. secondary contributors. Supporting: clinical diagnostic practice (DSM principal diagnosis model), factorial designs showing dominant variance explanations in motivation research.')
    add_paragraph(doc, 'Key Nuance: VCM does not claim single-cause purity. It claims that in any given failure moment, one root cause explains the majority of the blocking variance, making it the optimal intervention target.')

    add_heading(doc, 'Claim A3: Gate Exhaustiveness', level=2)
    add_paragraph(doc, 'Statement: The eight gates are exhaustive for mapping ALL volitional failure. Any failure of goal-directed behavior can be attributed to one of these eight gates. The root causes within each gate are NOT yet claimed to be exhaustive \u2014 new root causes may be discovered.')
    add_paragraph(doc, 'Claim Type: Completeness / Theoretical')
    add_paragraph(doc, 'Strength: Strong claim \u2014 requires showing no volitional failure exists outside the eight-gate space')
    add_paragraph(doc, 'Validation Requirement: Must demonstrate that proposed counter-examples (failures that appear to fall outside VCM) can be mapped to existing gates. Strongest validation: adversarial testing \u2014 attempt to find unmappable failures. Supporting frameworks that approach completeness: Rubicon Model (4 phases), COM-B (3 components), Theoretical Domains Framework (14 domains).')

    add_heading(doc, 'Claim A4: Novel Synthesis + Novel Contribution', level=2)
    add_paragraph(doc, 'Statement: VCM is positioned as BOTH a novel synthesis (organizing established mechanisms from motivation science, neuroscience, and behavioral psychology into a unified sequential model) AND a novel theoretical contribution (the specific sequential gating architecture, the 56-cause taxonomy, and the decision computation formalism are new).')
    add_paragraph(doc, 'Claim Type: Positioning / Epistemological')
    add_paragraph(doc, 'Validation Implication: The synthesis claims require demonstrating that each mechanism has independent empirical support. The novelty claims require demonstrating that the specific combination, sequence, and formalization add explanatory or predictive power beyond existing frameworks.')

    # ---- SECTION 3: GATE-LEVEL CLAIMS ----
    add_heading(doc, '3. GATE-LEVEL CLAIMS', level=1)
    add_paragraph(doc, 'Each gate makes three types of claims: (1) the gate represents a distinct constraint type in volitional processing, (2) the neural systems listed are dominant players (not exclusive), and (3) the gate\'s root causes collectively capture the major failure modes for that constraint type.')

    add_heading(doc, '3.1 Neural Validation Status (All Gates)', level=2)
    add_paragraph(doc, 'Claim Type for Neural Assignments: The neural systems listed represent dominant players in each gate\'s processing. This is NOT a claim of exclusivity \u2014 other regions contribute. It IS a claim that these systems are necessary for the gate\'s function, meaning damage or disruption to them would impair the gate\'s operation.')
    add_neural_table(doc)

    # Gate descriptions
    gate_order = ['C0', 'C1', 'C2', 'C3', 'C4', 'C5', 'C6A', 'C6B']
    section_num = 2
    for gate_id in gate_order:
        gate = ENTRIES[gate_id]
        section_num += 1
        add_heading(doc, f'3.{section_num - 1} {gate_id}: {gate["gate_name"]}', level=2)
        add_paragraph(doc, f'Constraint Type: {gate["constraint_type"]}')
        add_paragraph(doc, f'Diagnostic Question: {gate["question"]}')
        add_paragraph(doc, f'Neural Systems: {gate["neural"]}')
        add_paragraph(doc, f'Root Cause Count: {len(gate["entries"])} entries')
        if gate.get('framing_note'):
            add_paragraph(doc, f'Framing Note: {gate["framing_note"]}', italic=True)

    # ---- SECTION 4: ROOT-CAUSE-LEVEL CLAIMS ----
    add_heading(doc, '4. ROOT-CAUSE-LEVEL CLAIMS', level=1)
    add_paragraph(doc, 'Each of the 56 root causes makes the following claims, listed from strongest to weakest:')
    add_paragraph(doc, '(a) Existence: The failure mode exists as a distinct clinical/behavioral pattern')
    add_paragraph(doc, '(b) Gate Classification: The root cause correctly maps to its parent gate')
    add_paragraph(doc, '(c) Mechanism: The mechanism description accurately captures what goes wrong')
    add_paragraph(doc, '(d) Decision Computation: The IF/THEN formula describes the pattern (examined under dual interpretation)')
    add_paragraph(doc, '(e) Intervention: The intervention axes address the mechanism')
    add_paragraph(doc, '(f) Effect Size: The stated effect sizes are supported by evidence')

    add_heading(doc, '4.1 Decision Computation: Dual Interpretation', level=2)
    add_heading(doc, 'Interpretation A: Computational Neuroscience (Strong)', level=3)
    add_paragraph(doc, 'Under this interpretation, the formulas describe what the brain is actually computing. For example, \'IF Environmental_Friction > Activation_Threshold THEN Action_Blocked\' claims that the dACC literally runs a cost comparison and outputs a block signal to downstream motor systems. This requires evidence from computational neuroscience showing these specific operations in these specific circuits.')
    add_paragraph(doc, 'Burden of Proof: HIGH \u2014 requires computational modeling studies, single-cell recording data, or detailed fMRI/EEG studies showing the specific computation.')

    add_heading(doc, 'Interpretation B: Diagnostic Heuristic (Pragmatic)', level=3)
    add_paragraph(doc, 'Under this interpretation, the formulas are clinician-facing shorthand that describe the observable pattern. The formula captures what you would see in the person\'s behavior and self-report, not a literal neural computation. This requires behavioral evidence that the pattern holds reliably and that the formula correctly predicts when the failure mode is active.')
    add_paragraph(doc, 'Burden of Proof: MODERATE \u2014 requires behavioral studies, clinical observation patterns, or self-report data showing the formula\'s variables predict the outcome.')

    add_paragraph(doc, 'VCM Position: Both interpretations are examined. Where computational neuroscience evidence exists, it strengthens the claim. Where only behavioral evidence exists, the formula remains valid as a diagnostic heuristic. The validation process will flag which interpretation each formula supports.')

    # ---- SECTION 5: COMPLETE ROOT CAUSE REGISTRY ----
    add_heading(doc, '5. COMPLETE ROOT CAUSE REGISTRY', level=1)
    add_paragraph(doc, 'All 56 root causes with their core claims. Each entry is validated individually across six dimensions: Existence, Gate Classification, Mechanism, Decision Computation, Intervention, and Effect Size.')
    add_paragraph(doc, 'Corrected structure: Each claim table now contains six rows corresponding to the six validation dimensions defined in Section 6.2, resolving the prior conflation between existence claims (does this pattern exist?) and mechanism claims (does it work this way?).')

    for gate_id in gate_order:
        gate = ENTRIES[gate_id]
        add_heading(doc, f'{gate_id}: {gate["gate_name"]} ({len(gate["entries"])} root causes)', level=2)

        if gate.get('framing_note'):
            add_paragraph(doc, gate['framing_note'], italic=True)

        for entry in gate['entries']:
            add_heading(doc, f'{entry["id"]}: {entry["name"]}', level=3)
            add_claim_table(doc, entry)

    # ---- SECTION 6: VALIDATION CRITERIA ----
    add_heading(doc, '6. VALIDATION CRITERIA AND SCORING', level=1)

    add_heading(doc, '6.1 Evidence Rating Scale', level=2)
    add_rating_table(doc)

    add_heading(doc, '6.2 Validation Dimensions per Root Cause', level=2)
    add_paragraph(doc, 'Each root cause is evaluated on six dimensions:')
    add_paragraph(doc, '1. EXISTENCE \u2014 Does this failure pattern exist as described in behavioral/clinical literature? (Tests whether the pattern is a recognized, distinct phenomenon.)')
    add_paragraph(doc, '2. GATE CLASSIFICATION \u2014 Is this root cause correctly assigned to its parent gate, or does it belong elsewhere in the chain?')
    add_paragraph(doc, '3. MECHANISM \u2014 Does the described mechanism match empirical findings about how this failure works? (Tests the causal explanation.)')
    add_paragraph(doc, '4. COMPUTATION \u2014 Does the IF/THEN formula hold under Interpretation A (computational neuroscience) and/or Interpretation B (diagnostic heuristic)?')
    add_paragraph(doc, '5. INTERVENTION \u2014 Do the recommended intervention axes address this specific mechanism? (Tests therapeutic targeting.)')
    add_paragraph(doc, '6. EFFECT SIZE \u2014 Are the stated effect sizes supported by directly sourced evidence for interventions on this specific mechanism?')

    # ---- SECTION 7: OVERLAP ANALYSIS AND RESOLUTIONS ----
    add_heading(doc, '7. OVERLAP ANALYSIS AND RESOLUTIONS', level=1)
    add_paragraph(doc, 'Phase 3 identified 10 potentially overlapping pairs. Each pair was analyzed for distinctness and resolved with a decision (KEEP SEPARATE or CONSOLIDATE), differentiation criteria, and a diagnostic question for clinical use.')

    add_heading(doc, '7.1 Naming Resolution: C0-1 vs C6A-1', level=2)
    add_paragraph(doc, 'RESOLVED: Both entries were previously named "High Activation Cost," causing diagnostic confusion. Renamed for clarity:')
    add_paragraph(doc, 'C0-1: Renamed to "Environmental Friction" \u2014 external/environmental barriers (distance, setup time, clutter).')
    add_paragraph(doc, 'C6A-1: Renamed to "Motor Initiation Threshold" \u2014 internal neuromotor gating (basal ganglia threshold).')
    add_paragraph(doc, 'Diagnostic: Is there something in your environment making it hard to start, or do you feel stuck even when everything is set up? Environment \u2192 C0-1. Stuck despite readiness \u2192 C6A-1.')

    add_heading(doc, '7.2 Consolidated Pair: C6B-2 + C6B-9', level=2)
    add_paragraph(doc, 'RESOLVED: CONSOLIDATED. C6B-9 (Competing Comforts) merged into C6B-2 (Competing Stimuli). Both shared the core mechanism of alternative reward outcompeting task reward. C6B-2 now includes two subtypes: stimulus-driven capture (involuntary, bottom-up) and hedonic comfort-seeking (voluntary, top-down). Total entry count reduced from 57 to 56; C6B gate count reduced from 9 to 8.')

    add_heading(doc, '7.3 Kept Separate with Differentiation Criteria (8 pairs)', level=2)

    add_paragraph(doc, 'Pair 1: C3-6 (Decision Paralysis) vs C4-3 (Plan Selection Failure) \u2014 KEEP SEPARATE', bold=True)
    add_paragraph(doc, 'C3-6 operates in the deliberative mindset (Gollwitzer, 1990) \u2014 the person has not yet committed to a goal path. C4-3 operates in the implemental mindset \u2014 the person is committed but cannot select among implementation strategies.')
    add_paragraph(doc, 'Diagnostic: Have you decided you definitely want to achieve this goal? YES but don\'t know which strategy \u2192 C4-3. NO, still deciding whether this is the right goal/path \u2192 C3-6.')

    add_paragraph(doc, 'Pair 2: C2-1 (Value Misalignment) vs C2-6 (External Motivation) \u2014 KEEP SEPARATE', bold=True)
    add_paragraph(doc, 'C2-6 is a specific, identifiable subtype of the broader C2-1 pattern. Distinguished by the person\'s ability to identify the external source of motivation. C2-1 captures cases where the misalignment source is diffuse or unconscious.')
    add_paragraph(doc, 'Diagnostic: Can you identify WHY this goal doesn\'t feel like yours? YES, doing it for someone else\'s approval/reward \u2192 C2-6. NO, it just doesn\'t resonate \u2192 C2-1.')

    add_paragraph(doc, 'Pair 3: C2-2 (Fear-Suppressed Desire) vs C2-3 (Emotional Avoidance) \u2014 KEEP SEPARATE', bold=True)
    add_paragraph(doc, 'C2-2 involves a specific feared consequence blocking an acknowledged desire (amygdala \u2192 PAG threat circuit). C2-3 involves diffuse emotional aversiveness that may suppress awareness of the desire itself (anterior insula, medial PFC social cognition networks).')
    add_paragraph(doc, 'Diagnostic: Do you actively want this goal but feel held back by a specific fear, or does the whole topic make you uncomfortable in a way that\'s hard to pin down? Specific fear blocking known desire \u2192 C2-2. Diffuse discomfort, may not acknowledge desire \u2192 C2-3.')

    add_paragraph(doc, 'Pair 4: C5-7 (Identity Misalignment) vs C5-8 (Identity-Behavior Dissonance) \u2014 KEEP SEPARATE', bold=True)
    add_paragraph(doc, 'C5-7 is PROSPECTIVE \u2014 the person rejects a goal before sustained engagement because it doesn\'t fit current identity. C5-8 is RETROSPECTIVE/ACCUMULATED \u2014 repeated failure creates a widening gap between aspiration and action (Brandstatter et al., 2013: cortisol elevation).')
    add_paragraph(doc, 'Diagnostic: Is this goal something that doesn\'t fit who you are, or something you\'ve been trying to do but keep falling short? Doesn\'t fit who I am \u2192 C5-7. Keep trying and failing \u2192 C5-8.')

    add_paragraph(doc, 'Pair 5: C3-2 (Delayed Gratification Intolerance) vs C2-8 (Lack of Future Relevance) \u2014 KEEP SEPARATE', bold=True)
    add_paragraph(doc, 'C3-2 is an impulse control problem (\'hot\' cognition) \u2014 limbic vs prefrontal competition. The person DOES want the goal but can\'t tolerate the delay. C2-8 is an identity/connectedness problem (\'cold\' cognition) \u2014 the future self feels like a stranger (Ersner-Hershfield et al., 2009).')
    add_paragraph(doc, 'Diagnostic: Does the goal feel irrelevant to your future, or do you want it but can\'t stand waiting for the payoff? Irrelevant to future \u2192 C2-8. Want it but need results NOW \u2192 C3-2.')

    add_paragraph(doc, 'Pair 6: C4-1 (Conflicting Plans) vs C5-1 (Competing Commitments) \u2014 KEEP SEPARATE', bold=True)
    add_paragraph(doc, 'C4-1 is conscious resource conflict at the planning stage \u2014 the person knows they have too many active goals and cannot form a coherent implementation plan. C5-1 includes unconscious counter-commitments (Kegan & Lahey, 2009) that surface only during execution.')
    add_paragraph(doc, 'Diagnostic: Are you struggling to organize your goals into a workable plan, or did you have a plan but keep getting pulled away by other obligations? Stuck at planning stage \u2192 C4-1. Had a plan but obligations override it / discover hidden conflicts during execution \u2192 C5-1.')

    add_paragraph(doc, 'Pair 7: C2-5 (Approach-Avoidance Conflict) vs C5-6 (Value Contradiction) \u2014 KEEP SEPARATE', bold=True)
    add_paragraph(doc, 'C2-5 is motivational ambivalence (subcortical BIS activation) \u2014 competing approach and avoidance drives at the desire stage. Pre-reflective and emotional. C5-6 is cognitive values conflict (cortical value representation) \u2014 post-commitment erosion from a competing explicitly-held value.')
    add_paragraph(doc, 'Diagnostic: Do you feel pulled in two directions about wanting this at all, or have you committed but it conflicts with something else you believe in? Torn about wanting it \u2192 C2-5. Committed but clashes with another value \u2192 C5-6.')

    add_paragraph(doc, 'Pair 8: C6A-3 (Attention Diversion) vs C6B-2 (Competing Stimuli) \u2014 KEEP SEPARATE', bold=True)
    add_paragraph(doc, 'Same core mechanism (attentional capture) but at different temporal phases. C6A-3 occurs at the initiation moment (zero perceptual load = maximally vulnerable). C6B-2 occurs during sustained effort. Both may co-occur; intervention emphasis differs (reduce start-friction vs increase persistence structure).')
    add_paragraph(doc, 'Diagnostic: Do you get distracted when you\'re trying to start, or after you\'ve already begun working? Right at start \u2192 C6A-3. After starting \u2192 C6B-2. Both \u2192 flag both, assess which is primary.')

    add_heading(doc, '7.4 Cross-Gate Overlaps (Section 4C)', level=2)
    add_paragraph(doc, 'Additional overlap pairs identified during validation that require cross-referencing but are kept separate with boundary notes.')

    add_paragraph(doc, 'Pair A: C6A-1 (Motor Initiation Threshold) vs C6A-7 (Momentum Dependency) \u2014 KEEP SEPARATE', bold=True)
    add_paragraph(doc, 'Both share dopaminergic literature (Volkow citations). C6A-7 is a specific clinical subtype of C6A-1 with pronounced initiation-continuation asymmetry. C6A-1 is the general motor threshold problem; C6A-7 is the specific pattern where continuation is easy but cold-start is impossible (common in ADHD).')
    add_paragraph(doc, 'Diagnostic: Once you finally start, can you keep going easily? YES \u2192 C6A-7 (momentum dependency). NO, staying going is also hard \u2192 C6A-1 (general motor threshold).')

    add_paragraph(doc, 'Pair B: C6A-4 (Waiting for Readiness) vs C6B-3 (Mood-Dependent Execution) \u2014 KEEP SEPARATE', bold=True)
    add_paragraph(doc, 'Same unitary mechanism (mood-contingent action policy) split across initiation vs persistence phases. Gendolla\'s Mood-Behavior Model does not differentiate pre-task vs mid-task. Model-imposed distinction retained for diagnostic utility.')
    add_paragraph(doc, 'Diagnostic: Do you wait for the right mood to START, or do you stop mid-task when your mood drops? Wait to start \u2192 C6A-4. Stop mid-task \u2192 C6B-3. Both \u2192 flag both.')

    add_paragraph(doc, 'Pair C: C6B-4 (Single-Error Catastrophizing) vs C6B-5 (Perfectionism) \u2014 KEEP SEPARATE', bold=True)
    add_paragraph(doc, 'Shared root in maladaptive perfectionism (Frost et al., 1990). C6B-4 maps to "Concern over Mistakes" dimension; C6B-5 maps to "Personal Standards" + "Doubts about Actions" dimensions. Distinct intervention targets despite shared origin.')
    add_paragraph(doc, 'Diagnostic: Is the problem that one mistake makes you quit, or that nothing ever feels good enough? One mistake = quit \u2192 C6B-4. Never good enough \u2192 C6B-5.')

    add_paragraph(doc, 'Pair D: C0-P Subtypes \u2014 UNDER REVIEW', bold=True)
    add_paragraph(doc, 'Graded biological constraints interact with C1\u2013C3. Consider future split: C0-P1 (Structural Motor Block \u2014 binary) vs C0-P2 (Biological Effort-Cost Elevation \u2014 graded). Current single-entry structure retained pending further clinical validation.')

    add_heading(doc, '7.5 Effect Size Sourcing', level=2)
    add_paragraph(doc, 'Most effect sizes are stated as ranges (e.g., d = 0.4\u20130.7). Validation must determine for each: (a) whether the range is drawn from cited studies directly measuring the intervention on the stated mechanism, (b) whether the range is estimated from related literature, or (c) whether the range is a clinical estimate without direct empirical basis. This will be tracked per root cause.')

    # ---- SECTION 8: VALIDATION ROADMAP ----
    add_heading(doc, '8. VALIDATION ROADMAP', level=1)
    add_paragraph(doc, 'Phase 1 (COMPLETE): Gate-level neural validation \u2014 all 8 gates confirmed against neuroscience literature.')
    add_paragraph(doc, 'Phase 2 (THIS DOCUMENT): Foundational claims articulation \u2014 every testable assertion documented with validation criteria. Corrected to include six validation dimensions per root cause (Existence, Gate Classification, Mechanism, Computation, Intervention, Effect Size).')
    add_paragraph(doc, 'Phase 3 (NEXT): Systematic root-cause validation \u2014 all 56 entries evaluated on 6 dimensions against empirical literature. Will proceed gate by gate, flagging misclassifications, unsupported claims, and evidence gaps.')
    add_paragraph(doc, 'Phase 4: Claims revision \u2014 update database based on validation findings. Remove unsupported entries, reclassify misplaced ones, add missing root causes identified during literature review. Resolve flagged overlaps.')
    add_paragraph(doc, 'Phase 5: Effect size audit \u2014 verify or replace all intervention effect size claims with directly sourced evidence.')

    # ---- SECTION 9: CROSS-CULTURAL VALIDITY ----
    add_heading(doc, '9. CROSS-CULTURAL VALIDITY NOTE', level=1)
    add_paragraph(doc, 'The majority of studies supporting VCM claims were conducted with WEIRD (Western, Educated, Industrialized, Rich, Democratic) samples. This section documents the implications for claim generalizability.')
    add_paragraph(doc, 'Claims most likely to be cross-culturally valid:', bold=True)
    add_paragraph(doc, '(a) Gate architecture claims (A1\u2013A3) \u2014 sequential constraint processing maps to conserved neural systems (basal ganglia gating, prefrontal executive control, dopaminergic reward). These are biologically grounded and less culture-dependent.')
    add_paragraph(doc, '(b) Root causes with primarily neurobiological mechanisms \u2014 e.g., C6A-1 (Motor Initiation Threshold), C3-2 (Delayed Gratification Intolerance), C6B-1 (Task Boredom). The underlying neural computations are universal.')
    add_paragraph(doc, 'Claims requiring cultural validation:', bold=True)
    add_paragraph(doc, '(a) Root causes involving social cognition \u2014 e.g., C5-3 (No External Monitoring), C1-7 (Social Disconfirmation), C2-6 (External Motivation). Social pressure dynamics vary significantly across individualist vs. collectivist cultures.')
    add_paragraph(doc, '(b) Root causes involving identity and values \u2014 e.g., C5-7 (Identity Misalignment), C5-6 (Value Contradiction), C2-1 (Value Misalignment). Self-concept structure differs across cultures (independent vs. interdependent self-construal).')
    add_paragraph(doc, '(c) Effect sizes \u2014 intervention effect magnitudes may differ across cultural contexts, particularly for self-monitoring, social accountability, and value alignment interventions.')
    add_paragraph(doc, 'Recommendation: Cross-cultural validation should be a dedicated phase of the validation roadmap, conducted after the core evidence audit (Phases 3\u20135) establishes baseline validity with the existing literature.')

    # ---- SAVE ----
    output_path = r'c:\Users\randy\.claude\projects\kdenz\Claude code created\VCM_Claims_Document_CORRECTED.docx'
    doc.save(output_path)
    print(f'Document saved to: {output_path}')

    # Count verification
    total = sum(len(gate['entries']) for gate in ENTRIES.values())
    print(f'Total entries generated: {total}')
    c0_count = len(ENTRIES['C0']['entries'])
    volitional_count = total - c0_count
    print(f'Pre-volitional (C0): {c0_count}')
    print(f'Volitional (C1-C6B): {volitional_count}')


if __name__ == '__main__':
    build_document()

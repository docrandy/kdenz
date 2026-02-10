"""
Generate the corrected VCM Root Cause Database document.
Additions from Phase 2:
- Key Studies (2-4 citations per entry)
- Evidence Level (STRONG/MODERATE/WEAK)
- Moderating Factors
- Evidence Notes (caveats, replication concerns)
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

# ============================================================
# GATE-LEVEL DATA
# ============================================================

GATES = {
    "C0": {
        "gate_name": "Environmental Permeability",
        "constraint_type": "Environment",
        "question": "Can the person physically act? Assesses whether structural, biological, or environmental conditions permit goal-directed action regardless of motivation.",
        "neural": "Sensory Cortex, Lateral PFC, Medial PFC, Dorsal Anterior Insula",
        "vocal_signature": "Flat prosody, monotone delivery, low volume, slow speech rate, frequent pauses",
        "framing_note": "C0 is a pre-volitional screening gate. It rules out non-volitional blockers before diagnosing failure within the volitional chain (C1-C6B). If C0 fails, the issue is environmental, biological, or structural -- not volitional.",
    },
    "C1": {
        "gate_name": "Believability",
        "constraint_type": "Perception",
        "question": "Does the person believe the goal is achievable? Assesses self-efficacy, perceived control, and belief formation processes.",
        "neural": "Ventral Striatum, vmPFC, Lateral PFC",
        "vocal_signature": "Low confidence, pitch instability, hedging language, rising intonation, hesitation clusters",
        "framing_note": None,
    },
    "C2": {
        "gate_name": "Desire",
        "constraint_type": "Desire",
        "question": "Does the person genuinely want the goal? Assesses intrinsic motivation, value alignment, and approach-avoidance dynamics.",
        "neural": "Midbrain Dopamine, Ventral Striatum, vmPFC",
        "vocal_signature": "Flat affect, low enthusiasm, monotone when discussing goal, voice brightens on tangents",
        "framing_note": None,
    },
    "C3": {
        "gate_name": "Will",
        "constraint_type": "Will",
        "question": "Is the person willing to bear the cost? Assesses effort computation, discomfort tolerance, and cost-benefit analysis at the commitment threshold.",
        "neural": "dACC, Anterior Insula, dlPFC",
        "vocal_signature": "Sighing, groaning, heavy breath before speaking, voice trails off mid-sentence, energy drop",
        "framing_note": None,
    },
    "C4": {
        "gate_name": "Intention",
        "constraint_type": "Intention",
        "question": "Has the person formed a concrete plan? Assesses plan specificity, implementation intentions, and cue-response binding.",
        "neural": "dlPFC, Posterior Parietal Cortex, Mid-lateral PFC",
        "vocal_signature": "Vague language, excessive hedging, circular speech, restarts without resolution",
        "framing_note": None,
    },
    "C5": {
        "gate_name": "Commitment",
        "constraint_type": "Commitment",
        "question": "Is the person's commitment robust? Assesses goal shielding, priority maintenance, and resistance to competing demands.",
        "neural": "Anterior Cingulum Bundle, dlPFC, vmPFC",
        "vocal_signature": "Deflection, topic-switching, minimizing language, rationalization patterns, verbal distancing",
        "framing_note": None,
    },
    "C6A": {
        "gate_name": "Action Initiation",
        "constraint_type": "Action Initiation",
        "question": "Can the person start? Assesses motor initiation, threshold activation, and barriers at the moment of action.",
        "neural": "SMA, Pre-SMA, ACC motor zone, Basal Ganglia",
        "vocal_signature": "Long pauses before speaking, false starts, voice catches, trailing off at action words",
        "framing_note": None,
    },
    "C6B": {
        "gate_name": "Action Persistence",
        "constraint_type": "Action Persistence",
        "question": "Can the person sustain action? Assesses effort regulation, distraction resistance, and persistence through obstacles once action has begun.",
        "neural": "Dorsal Striatum, ACC Persistence Circuit, Motor Cortex",
        "vocal_signature": "Frustration bursts, speed fluctuation, giving-up language, comparative complaints",
        "framing_note": None,
    },
}

# ============================================================
# ALL 56 ENTRIES: COMPLETE DATA (Original + Phase 2 Research + Phase 3 Corrections)
# ============================================================

ENTRIES = {
    # ---- C0: ENVIRONMENTAL PERMEABILITY ----
    "C0-P": {
        "gate": "C0", "name": "Physical/Biological Limitation",
        "definition": "Medical, neurological, or metabolic conditions block action regardless of motivation. Consolidates subtypes: Psychomotor Retardation (C0-3), Motor Readiness Failure (C0-4), Metabolic Insufficiency (C0-5), HPA Axis Exhaustion (C0-6).",
        "mechanism": "Structural biological constraint prevents motor output regardless of psychological state.",
        "decision_comp": "IF Physical_Limitation = TRUE THEN Action_Blocked (regardless of motivation state)",
        "diagnostic_signals": "Chronic fatigue reports | Medical history of relevant conditions | Consistent inability across all goal domains",
        "intervention": "Routes to medical evaluation. Not addressable through behavioral intervention. Subtypes: Psychomotor Retardation (dopamine/motor system slowing), Motor Readiness Failure (structural motor pathway damage), Metabolic Insufficiency (ATP/energy depletion), HPA Axis Exhaustion (autonomic burnout).",
        "effect_size": "N/A -- routes to medical evaluation",
        "key_studies": [
            "Heinz et al. (2023). Psychomotor Retardation in Major Depression. Molecular Psychiatry. n=1,519; largest neuroimaging study; confirms ~70% prevalence of psychomotor retardation in MDD.",
            "Bennabi et al. (2013). Psychomotor Retardation in Depression: A Systematic Review. BioMed Research International. Systematic review of psychomotor retardation involving PFC deficits and dopamine dysregulation.",
            "McEwen (1998). Protective and Damaging Effects of Stress Mediators. NEJM. Chronic HPA activation produces allostatic load damaging brain structures.",
            "Buyukdura et al. (2011). Psychomotor Retardation in Depression. Progress in Neuro-Psychopharmacology and Biological Psychiatry. Measurable slowing with roots in basal ganglia-thalamocortical circuits.",
            "Koller et al. (2022). Objective Motor Activity in Depression. Psychological Medicine. 34 studies, n=1,804; objective actigraphy confirmation SMD=-0.78.",
            "Salamone & Correa (2024). Motivational Functions of Dopamine. Annual Review of Psychology, 75. Effort-cost framework; challenges binary biology/motivation divide.",
            "Walitt et al. (2024). Deep Phenotyping of ME/CFS. Nature Communications. NIH study; effort preference alteration in ME/CFS.",
            "Guidi et al. (2023). Allostatic Framework Update. Psychotherapy and Psychosomatics, 92(5). Updates McEwen; intervention responsiveness.",
        ],
        "evidence_level": "STRONG",
        "moderating_factors": "Depression severity/subtype | Chronicity of stress exposure | Medication effects | Comorbid medical conditions",
        "evidence_notes": "CORRECTIONS: (1) Berger et al. (2017) REMOVED -- that paper evaluates CPTs for ADHD diagnosis (excess motor activity), opposite of motor block; belongs in C3-C4 territory, not C0-P. (2) 70% prevalence figure REATTRIBUTED from Bennabi et al. (2013) to Heinz et al. (2023), Molecular Psychiatry (n=1,519) -- the 70% figure does not appear in Bennabi et al. Well-established overall. HPA exhaustion trajectory debated.",
    },
    "C0-1": {
        "gate": "C0", "name": "Environmental Friction",
        "definition": "Environmental friction (distance, setup time, clutter) raises initiation threshold above motor readiness.",
        "mechanism": "Environmental friction increases perceived effort cost beyond willingness-to-act threshold.",
        "decision_comp": "IF Environmental_Friction > Activation_Threshold THEN Action_Blocked",
        "diagnostic_signals": "Setup time complaints | Distance/access barriers cited | Clutter/disorganization prevents starting",
        "intervention": "Reduce activation cost, Environmental redesign, Reduce friction.",
        "effect_size": "d = 0.3-0.8 (Mertens 2022 d=0.43 overall; Hu 2025 d=0.27 unadjusted; wide range reflects intervention heterogeneity)",
        "key_studies": [
            "Thaler & Sunstein (2008). Nudge. Yale University Press. Small environmental changes dramatically alter behavior.",
            "Duckworth et al. (2016). Situational Strategies for Self-Control. Perspectives on Psychological Science. Modifying situations more effective than willpower.",
            "Mani et al. (2013). Poverty Impedes Cognitive Function. Science. Scarcity consumes cognitive bandwidth. FLAG: O'Donnell et al. 2021 failed to replicate; Szecsi & Szaszi 2024 found 'moderate evidence against.'",
            "Madrian & Shea (2001). Power of Suggestion: Inertia in 401(k). QJE. Opt-in to opt-out: 49% to 86% participation.",
            "Mertens et al. (2022). Choice Architecture Interventions: Meta-Analysis. PNAS. 447 effects, n=2.1M; d=0.43 overall; decision structure nudges outperform others.",
            "Hu (2025). Nudging Behavior: Meta-Analysis. Journal of Behavioral Decision Making. 1,638 studies, ~30M participants; most comprehensive synthesis; publication bias concerns.",
            "OECD (2024). Fixing Frictions. 16 governments; real-world sludge audit validation of friction-reduction interventions.",
        ],
        "evidence_level": "STRONG",
        "moderating_factors": "Magnitude of friction | Individual executive function capacity | Habit strength | Motivational intensity",
        "evidence_notes": "CORRECTIONS: (1) Effect size corrected d=0.5-0.8 → d=0.3-0.8 (lower bound adjusted down per Hu 2025 d=0.27). (2) Mani et al. (2013) FLAGGED: O'Donnell et al. 2021 failed replication; Szecsi & Szaszi 2024 found 'moderate evidence against'; weakest fit for C0-1. (3) Madrian author name already corrected (was 'Madian'). Robust across domains overall. DIFFERENTIATION: Distinct from C6A-1 (Motor Initiation Threshold). C0-1 is external/environmental -- physical barriers, distance, setup time. C6A-1 is internal/neuromotor -- the brain's basal ganglia gating system. DIAGNOSTIC: Is there something in your environment making it hard to start, or do you feel stuck even when everything is set up? Environment → C0-1. Stuck despite readiness → C6A-1.",
    },
    "C0-2": {
        "gate": "C0", "name": "Situational Blockage",
        "definition": "Structural, resource, or social constraints (time, money, access, legal, family opposition) make goal objectively non-viable.",
        "mechanism": "Objective structural constraints prevent action regardless of internal motivation state.",
        "decision_comp": "IF Resource_Available = FALSE OR Access_Blocked = TRUE THEN Goal_Non_Viable",
        "diagnostic_signals": "Chronic environment excuses (legitimate) | Resource constraints verified | Social/legal barriers present",
        "intervention": "Change structure, Modify context, Reduce goal scope.",
        "effect_size": "d = 0.3-0.6 (depends on constraint type)",
        "key_studies": [
            "Lewin (1943). Defining the Field at a Given Time. Psychological Review. B=f(LS); barriers create frustration proportional to goal importance.",
            "Mullainathan & Shafir (2013). Scarcity. Times Books. Resource scarcity creates tunneling and cascading constraints.",
            "Gifford (2011). Dragons of Inaction. American Psychologist. Seven categories of PSYCHOLOGICAL barriers to action (not structural as previously characterized).",
            "Herd & Moynihan (2019). Administrative Burden. Russell Sage Foundation. Learning, compliance, and psychological cost framework for structural barriers.",
            "Peters & O'Connor (1980). Situational Constraints and Work Outcomes. Academy of Management Review. Eight categories blocking performance independent of ability/motivation.",
            "Wrosch & Scheier (2003, 2020). Goal Adjustment Theory. Foundation for adaptive goal disengagement under constraint.",
            "Barlow, Wrosch, & McGrath (2020). Goal Adjustment Capacities. Journal of Personality, 88(2). 31 samples; meta-analytic support for goal adjustment.",
        ],
        "evidence_level": "STRONG",
        "moderating_factors": "Objective constraint severity | Perceived vs. actual constraint | Social capital/support networks | Time horizon",
        "evidence_notes": "CORRECTION: Gifford (2011) identifies seven PSYCHOLOGICAL barriers ('dragons of inaction'), not structural barriers as previously characterized. Herd & Moynihan (2019) added for structural barriers framework. Mullainathan & Shafir (2013) has replication concerns: only 4 of 18 studies replicated at significance; 30% directionally opposite. Lewin foundational but historically qualitative. Peters & O'Connor provides strongest direct empirical support but domain-specific.",
    },

    # ---- C1: BELIEVABILITY ----
    "C1-1": {
        "gate": "C1", "name": "Cognitive Biases (Overconfidence/Planning Fallacy)",
        "definition": "Overconfident predictions override realistic assessment; planning fallacy creates prediction-reality gap.",
        "mechanism": "Inside view dominance; base-rate neglect inflates success expectations.",
        "decision_comp": "IF Confidence > Evidence_Base THEN Prediction_Error_HIGH",
        "diagnostic_signals": "History of underestimating timelines | Surprise at failures | No contingency planning",
        "intervention": "Rebuild efficacy, Increase specificity, Mastery experiences.",
        "effect_size": "d = 0.5-0.7 (medium to large)",
        "key_studies": [
            "Buehler et al. (1994). Exploring the Planning Fallacy. JPSP. N=37 honors students predicted 33.9 days, took 55.5 days (exceeded even worst-case estimates of 48.6 days).",
            "Kahneman & Tversky (1979). Intuitive Prediction: Biases and Corrective Procedures. TIMS. Inside view vs outside view.",
            "Kruger & Dunning (1999). Unskilled and Unaware. JPSP. Bottom-quartile estimated 62nd percentile (actual 12th).",
            "Buehler et al. (2010). The Planning Fallacy: Cognitive, Motivational, and Social Origins. AESP.",
            "Moore & Healy (2008). The Trouble with Overconfidence. Psychological Review. Three types of overconfidence distinction.",
            "Jansen, Rafferty, & Griffiths (2021). A Rational Model of Dunning-Kruger. Nature Human Behaviour, 5. Rational model reframes the effect.",
        ],
        "evidence_level": "STRONG",
        "moderating_factors": "Task familiarity | Accountability/incentives | Need for cognition | Group vs individual prediction",
        "evidence_notes": "Planning fallacy among most replicated findings. Dunning-Kruger has statistical criticisms (regression to mean; Gignac & Zajenkowski 2020).",
    },
    "C1-2": {
        "gate": "C1", "name": "Fixed Mindset",
        "definition": "Belief that personal limitations are unchangeable biases efficacy updating downward.",
        "mechanism": "Entity theory of ability; perceived ceiling on growth biases efficacy updating downward, especially in at-risk populations. Effect is real but substantially smaller than originally proposed.",
        "decision_comp": "IF Belief(Ability_Fixed) = TRUE THEN Efficacy_Updating_Biased_Downward",
        "diagnostic_signals": "'I'm just not good at this' | 'Some people can, I can't' | Avoids challenge",
        "intervention": "Rebuild efficacy, Mastery experiences, Reframe reward.",
        "effect_size": "d = 0.05-0.20 (CRITICAL: 3-12x smaller than previously claimed d=0.4-0.6; Macnamara & Burgoyne 2023 d=0.05 overall; Sisk 2018 r=0.10)",
        "key_studies": [
            "Dweck & Leggett (1988). Social-Cognitive Approach to Motivation. Psychological Review. Entity vs incremental theories.",
            "Macnamara & Burgoyne (2023). Do Growth Mindset Interventions Impact Students' Academic Achievement? Psychological Bulletin, 149. 63 studies, N=97,672: d=0.05 overall; highest-quality studies d=0.02.",
            "Sisk et al. (2018). Growth Mind-Sets Important to Achievement? Two Meta-Analyses. Psychological Science. r=0.10, d=0.08 -- very small effects.",
            "Yeager et al. (2019). A National Experiment Reveals Where a Growth Mindset Improves Achievement. Nature, 573. ~12,500 students; 0.1 GPA improvement, lower-achievers only.",
            "Burgoyne, Hambrick, & Macnamara (2020). Is the Rotherham Cohort Study a Fair Test of Growth Mindset? Psychological Science, 31(3). Only 2 of 6 mindset premises significant.",
            "Gazmuri (2025). Growth Mindset Interventions: A Systematic Review. Review of Education. 24 RCTs; strongest studies d=-0.01 to +0.065; 7/24 had conflicts of interest.",
        ],
        "evidence_level": "WEAK-to-MODERATE",
        "moderating_factors": "Domain specificity | Developmental stage | Challenge context | Cultural context | At-risk student status (effects concentrated here)",
        "evidence_notes": "CRITICAL CORRECTION: Effect size was overstated 3-12x. Original d=0.4-0.6 claim corrected to d=0.05-0.20 based on three independent meta-analyses. Rating downgraded from MODERATE to WEAK-to-MODERATE. Mechanism softened from 'blocks belief' to 'biases efficacy updating downward.' The construct is real but substantially smaller and more context-dependent than originally proposed. Effects concentrated in at-risk/lower-achieving students.",
    },
    "C1-3": {
        "gate": "C1", "name": "No Relatable Models",
        "definition": "Absence of proximal role models who share similar background undermines vicarious efficacy.",
        "mechanism": "Vicarious efficacy pathway blocked; no similar-other success evidence available.",
        "decision_comp": "IF Proximal_Models = 0 THEN Vicarious_Efficacy = LOW",
        "diagnostic_signals": "'Nobody like me has done this' | 'I don't know anyone who succeeded'",
        "intervention": "Rebuild efficacy, Increase value alignment.",
        "effect_size": "d = 0.15-0.35 (Lawner et al. 2019 field d=0.20, lab d=0.04)",
        "key_studies": [
            "Bandura (1977). Self-Efficacy: Toward a Unifying Theory of Behavioral Change. Psychological Review. Vicarious experience as one of four efficacy sources.",
            "Bandura (1997). Self-Efficacy: The Exercise of Control. Freeman. Model similarity is key moderator.",
            "Lockwood & Kunda (1997). Superstars and Me. JPSP. Attainable role models inspire; unattainable deflate.",
            "Lawner et al. (2019). Exposure to Role Models and Its Effect on Aspirations. Social Psychology of Education, 22. Most rigorous effect size: field d=0.20, lab d=0.04.",
            "Gladstone & Cimpian (2021). Role Models in STEM Education. International Journal of STEM Education, 8. Most comprehensive review; underrepresented group models help all.",
            "Verniers et al. (2024). Role Model Double-Edged Sword Effects. Review of Research in Education, 48(1). Role model backfire effects in some contexts.",
        ],
        "evidence_level": "STRONG",
        "moderating_factors": "Perceived similarity to model | Attainability of achievement | Number/diversity of models | Domain relevance",
        "evidence_notes": "Bandura's four-source model among most validated. Most evidence tests model presence not absence effects.",
    },
    "C1-4": {
        "gate": "C1", "name": "Meta-Cognitive Doubt",
        "definition": "Meta-level self-doubt ('I can't trust my own judgment') halts belief formation.",
        "mechanism": "Second-order doubt; questioning one's own cognitive processes prevents stable belief.",
        "decision_comp": "IF Trust(Own_Judgment) = LOW THEN Belief_Formation_Halted",
        "diagnostic_signals": "'I can't trust my own decisions' | 'What if I'm wrong about being able to do this?'",
        "intervention": "Rebuild efficacy, Increase perceived control, Mastery experiences.",
        "effect_size": "d = 0.3-0.5 (small to medium)",
        "key_studies": [
            "Petty et al. (2002). Thought Confidence as Determinant of Persuasion. JPSP. Low thought confidence neutralizes impact of primary cognitions.",
            "Brinol, P. & Petty, R. E. (2022). Self-Validation Theory. Psychological Review. Second-order doubt prevents first-order beliefs becoming consequential.",
            "Carroll, Brinol, Ketcham, & Petty (2020). Self-Validation and Metacognitive Doubt. JESP, 89. Direct metacognitive confidence mechanism evidence.",
            "Carroll (2025). Doubt in Goal Pursuit. Self and Identity. CRITICAL: doubt in goal doubts increased commitment.",
            "Wells (2009). Metacognitive Therapy for Anxiety and Depression. Guilford. Metacognitive beliefs maintain psychopathology.",
        ],
        "evidence_level": "MODERATE",
        "moderating_factors": "Trait anxiety | History of decisional regret | Cognitive complexity of domain | Social feedback on judgment quality",
        "evidence_notes": "Self-validation theory primarily in persuasion contexts, not self-efficacy. Wells' clinical model is treatment context. Construct is novel synthesis across multiple adjacent literatures.",
    },
    "C1-5": {
        "gate": "C1", "name": "Insufficient Evidence",
        "definition": "Insufficient information/knowledge about goal feasibility prevents stable belief.",
        "mechanism": "Belief formation requires minimum evidence threshold; insufficient data prevents commitment.",
        "decision_comp": "IF Evidence(Feasibility) < Belief_Threshold THEN Belief_Unstable",
        "diagnostic_signals": "'I don't know enough to know if I can do this' | Excessive research without action",
        "intervention": "Rebuild efficacy, Mastery experiences, Simulate context.",
        "effect_size": "d = 0.2-0.5 (high uncertainty; Tversky & Shafir replication V=0.11-0.30)",
        "key_studies": [
            "Ellsberg (1961). Risk, Ambiguity, and the Savage Axioms. QJE. Ambiguity aversion: preference for known over unknown probabilities.",
            "Fox & Tversky (1995). Ambiguity Aversion and Comparative Ignorance. QJE. Amplified by awareness of what one doesn't know.",
            "Tversky & Shafir (1992). Disjunction Effect in Choice Under Uncertainty. Psychological Science. People defer decisions when lacking clear reason.",
            "Lipshitz & Strauss (1997). Coping with Uncertainty: Naturalistic Decision-Making. OBHDP. Three types of uncertainty; inadequate understanding causes most severe paralysis.",
            "Xu & Tracey (2019/2020). Longitudinal Ambiguity Aversion and Self-Efficacy. Journal of Counseling Psychology. n=371; ambiguity aversion predicts reduced self-efficacy over time.",
            "Kobayashi et al. (2020). Uncertainty and Information-Seeking. Cognitive Research, 5. Unresolved uncertainty drives info-seeking but also blocks action.",
        ],
        "evidence_level": "MODERATE",
        "moderating_factors": "Trait ambiguity tolerance | Perceived stakes | Availability of comparative information | Time pressure",
        "evidence_notes": "Ambiguity aversion robust in behavioral economics. Application to personal goal feasibility beliefs is extrapolation from economic decision contexts.",
    },
    "C1-6": {
        "gate": "C1", "name": "Fragile Confidence",
        "definition": "Confidence easily undermined by setbacks or external criticism; high sensitivity to negative feedback.",
        "mechanism": "Efficacy representation unstable; single negative signal can collapse entire belief structure.",
        "decision_comp": "IF Setback_Sensitivity = HIGH AND Setback_Occurs THEN Efficacy_Collapse",
        "diagnostic_signals": "Confidence drops after minor criticism | One bad day erases weeks of progress",
        "intervention": "Rebuild efficacy, Reward persistence, Mastery experiences.",
        "effect_size": "d = 0.4-0.6 (medium; upper bound d=0.7 clinical populations only)",
        "key_studies": [
            "Kernis (2003). Toward Conceptualization of Optimal Self-Esteem. Psychological Inquiry. Fragile vs secure self-esteem distinction.",
            "Crocker & Park (2004). Costly Pursuit of Self-Esteem. Psychological Bulletin. Contingent self-worth creates domain-specific vulnerability.",
            "Kernis et al. (1989). Stability and Level of Self-Esteem. JPSP. Unstable high self-esteem shows most extreme defensive reactions.",
            "Brown (2010). High Self-Esteem Buffers Negative Feedback. Cognition and Emotion. Stable self-esteem buffers; unstable doesn't.",
            "Koszegi (2022). Fragile Self-Esteem. Review of Economic Studies, 89(4). Formal economic model of fragile self-esteem dynamics.",
            "Reitz et al. (2022). Self-Esteem Development and Life Events. Social and Personality Psychology Compass. Self-esteem development trajectories.",
        ],
        "evidence_level": "STRONG",
        "moderating_factors": "Domain of contingency | Stability vs level of self-esteem | Implicit/explicit discrepancy | Social support",
        "evidence_notes": "Well-validated with strong psychometrics. Most research correlational/cross-sectional. Real-time efficacy collapse dynamics more theoretical than experimentally demonstrated. CROSS-REFERENCE: C1-6 has dual citizenship -- C1 vulnerability (fragile initial belief) AND C6B persistence failure (collapse during action). Trigger 'Setback_Occurs' implies person may have already passed C1; consider C6B if collapse occurs mid-action.",
    },
    "C1-7": {
        "gate": "C1", "name": "Social Disconfirmation",
        "definition": "Social pressure or negative feedback undermines belief; external skepticism erodes self-efficacy through Bandura's verbal persuasion pathway.",
        "mechanism": "Verbal persuasion pathway (Bandura, 1977): negative social feedback degrades self-efficacy beliefs. Others' expressed doubt becomes internalized doubt, reducing perceived capability.",
        "decision_comp": "IF Social_Feedback = NEGATIVE THEN Verbal_Persuasion_Pathway_Activates; Efficacy_Erosion",
        "diagnostic_signals": "'Everyone says I can't' | 'My family doesn't believe in me'",
        "intervention": "Rebuild efficacy, Self-affirmation exercises, Increase perceived control, Change structure.",
        "effect_size": "d = 0.2-0.4 (small to medium; verbal persuasion is weakest of Bandura's four efficacy sources)",
        "key_studies": [
            "Bandura (1977). Self-Efficacy: Toward a Unifying Theory of Behavioral Change. Psychological Review. Verbal persuasion identified as one of four sources of self-efficacy; weakest source but potent when negative.",
            "Bandura (1997). Self-Efficacy: The Exercise of Control. Freeman. Comprehensive framework: verbal persuasion effects moderated by credibility of source and prior efficacy level.",
            "Usher & Pajares (2008). Sources of Self-Efficacy in School: Critical Review. Review of Educational Research. Empirical test of all four Bandura sources; social persuasion significant but weakest predictor.",
            "Cohen & Sherman (2014). The Psychology of Change: Self-Affirmation and Social Psychological Intervention. Annual Review of Psychology. Self-affirmation buffers against social disconfirmation effects.",
        ],
        "evidence_level": "WEAK",
        "moderating_factors": "Source credibility | Relationship with evaluator | Prior efficacy level | Domain identification | Self-affirmation resources",
        "evidence_notes": "CRITICAL CORRECTION: Previous citations (Steele & Aronson 1995, Spencer 1999, Nguyen & Ryan 2008, Flore & Wicherts 2015) were ALL stereotype threat literature -- wrong construct entirely. Social disconfirmation operates through Bandura's verbal persuasion pathway, not stereotype threat. Rating downgraded from MODERATE to WEAK because the correct evidence base (verbal persuasion) shows it is the weakest of Bandura's four efficacy sources. Foundational citations were critically missing.",
    },
    "C1-8": {
        "gate": "C1", "name": "Impostor Attribution Pattern",
        "definition": "Success attributed to external factors (luck, help) rather than internal capability; prevents efficacy consolidation. Related to but distinct from fundamental attribution error (Ross, 1977).",
        "mechanism": "External attribution bias; self-efficacy cannot update from success experiences. Weiner's (1985) attribution theory provides foundational framework: internal-stable attributions build efficacy; external-unstable attributions prevent updating.",
        "decision_comp": "IF Attribution(Success) = EXTERNAL THEN Efficacy_Update_Blocked",
        "diagnostic_signals": "'I got lucky' | 'Someone helped me' | 'It was easy, doesn't count'",
        "intervention": "Rebuild efficacy, Increase perceived control, Attribution retraining.",
        "effect_size": "d = 0.4-0.6 (estimated; no clean meta-analytic d exists for this specific pattern)",
        "key_studies": [
            "Weiner (1985). An Attributional Theory of Achievement Motivation and Emotion. Psychological Review. FOUNDATIONAL: internal/external, stable/unstable attribution dimensions determine efficacy updating.",
            "Clance & Imes (1978). Imposter Phenomenon in High Achieving Women. Psychotherapy. Original identification of impostor phenomenon.",
            "Bravata et al. (2020). Prevalence, Predictors, and Treatment of Impostor Syndrome. JGIM. 62 studies, prevalence 9-82%, comorbid with depression/anxiety/burnout.",
            "Leary et al. (2000). Impostor Phenomenon: Self-Perceptions and Interpersonal Strategies. Journal of Personality. External attribution stable across situations.",
            "Gisselbaek et al. (2025). Impostor Phenomenon in Medical Education: Umbrella Review. Medical Education. Comprehensive synthesis across healthcare domains.",
        ],
        "evidence_level": "MODERATE",
        "moderating_factors": "Achievement domain | Minority/marginalized status | Perfectionism | Organizational culture",
        "evidence_notes": "CORRECTIONS: (1) Rating downgraded STRONG → MODERATE: effect sizes are unsourced (no clean meta-analytic d exists for this pattern); causal direction genuinely debated. (2) Entry renamed from 'Attribution Error' to 'Impostor Attribution Pattern' to avoid conflict with 'fundamental attribution error' (Ross 1977). (3) Foundational Weiner (1985) attribution theory was critically missing and has been added. Well-documented over 4 decades but no RCTs for treatment (as of Bravata 2020). Wide prevalence range (9-82%) reflects measurement inconsistency.",
    },

    # ---- C2: DESIRE ----
    "C2-1": {
        "gate": "C2", "name": "Value Misalignment",
        "definition": "Goal pursued for extrinsic reasons, not authentic values; NAcc response blunted when goal conflicts with core values.",
        "mechanism": "Goal-value incongruence; vmPFC cannot assign authentic value to misaligned goal.",
        "decision_comp": "IF Goal_Value_Alignment < Threshold THEN Intrinsic_Motivation = LOW",
        "diagnostic_signals": "'I should want this but I don't' | 'Everyone says this is important' | Flat tone discussing goal",
        "intervention": "Increase value alignment, Reframe reward.",
        "effect_size": "d = 0.4-0.7 (medium to large; upper bound optimistic)",
        "key_studies": [
            "Sheldon & Elliot (1999). Self-Concordance Model. JPSP. Low self-concordance = less effort, lower well-being.",
            "Hare et al. (2009). Self-Control in Decision-Making. Science. vmPFC encodes subjective value; dlPFC modulates.",
            "Holton et al. (2024). Value-Based Goal Commitment. Nature Human Behaviour. Strongest direct neural evidence: vmPFC encodes goal commitment strength.",
            "Ryan & Deci (2000). SDT and Intrinsic Motivation. American Psychologist. Non-autonomous goals = reduced persistence.",
            "Ryan et al. (2023). Meta-Review of 60 Meta-Analyses Evaluating SDT. Broadest synthesis of self-determination evidence.",
        ],
        "evidence_level": "STRONG",
        "moderating_factors": "Degree of value articulation | Awareness of misalignment | Cultural context | Proximity to identity-central values",
        "evidence_notes": "Self-concordance well-replicated within SDT framework. Neural mechanism inferred from separate research streams rather than single integrated study. DIFFERENTIATION: C2-6 (External Motivation) is the identifiable variant of value misalignment where the person recognizes the goal is externally driven. C2-1 captures cases where the misalignment source is diffuse or unconscious. DIAGNOSTIC: Can you identify WHY this goal doesn't feel like yours? YES, doing it for someone else's approval/reward → C2-6. NO, it just doesn't resonate → C2-1.",
    },
    "C2-2": {
        "gate": "C2", "name": "Fear-Suppressed Desire",
        "definition": "Desire exists but suppressed by anticipated loss or disruption; fear of consequences blocks approach motivation.",
        "mechanism": "Amygdala threat response overrides approach motivation; desire present but gated by fear.",
        "decision_comp": "IF Fear(Consequence) > Desire(Goal) THEN Approach_Blocked",
        "diagnostic_signals": "'I want to but what if...' | 'It could ruin everything' | Visible tension when imagining success",
        "intervention": "Reduce cost, Reframe reward, Simulate context.",
        "effect_size": "d = 0.5-0.7 (medium to large)",
        "key_studies": [
            "Mobbs et al. (2007). When Fear Is Near. Science. Brain shifts from vmPFC to PAG (defensive freezing) as threat increases; vmPFC→PAG shift, not amygdala→PAG. Amygdala contains competing approach-avoidance circuits.",
            "Miller (1944). Experimental Studies of Conflict. Ronald Press. Avoidance gradients steeper than approach gradients.",
            "LeDoux et al. (2017). Birth, Death and Resurrection of Avoidance. Molecular Psychiatry. Desire not eliminated but suppressed by defensive circuits.",
            "Aupperle & Paulus (2010). Neural Systems in Anxiety Disorders. Dialogues in Clinical Neuroscience. Hyperactive amygdala overrides NAcc approach motivation.",
        ],
        "evidence_level": "STRONG",
        "moderating_factors": "Trait anxiety | Threat proximity/imminence | Perceived reversibility | Prior fear conditioning",
        "evidence_notes": "Neural architecture well-established. Most direct neural evidence from physical threat; extrapolation to abstract social fears involves inference. DIFFERENTIATION: Distinct from C2-3 (Emotional Avoidance). C2-2 involves a specific feared consequence blocking an acknowledged desire -- the person KNOWS they want the goal but fears a particular outcome (loss, failure, disruption). Neural pathway: amygdala → PAG threat circuit. DIAGNOSTIC: Do you actively want this goal but feel held back by a specific fear, or does the whole topic make you uncomfortable in a way that's hard to pin down? Specific fear blocking known desire → C2-2. Diffuse discomfort, may not acknowledge desire → C2-3.",
    },
    "C2-3": {
        "gate": "C2", "name": "Emotional Avoidance",
        "definition": "Emotional discomfort (anxiety, shame, embarrassment) associated with goal blocks desire; avoidance system dominates.",
        "mechanism": "Experiential avoidance; emotional cost of engagement exceeds perceived benefit.",
        "decision_comp": "IF Emotional_Cost(Engagement) > Perceived_Benefit THEN Avoidance_Dominates",
        "diagnostic_signals": "Avoidance behavior | Negative self-talk about goal | Physical discomfort discussing goal",
        "intervention": "ACT-based interventions (acceptance, defusion, values clarification), Psychological flexibility training, Mindfulness-based exposure, Reduce cost, Reframe reward.",
        "effect_size": "d = 0.5-0.8 (medium to large)",
        "key_studies": [
            "Hayes et al. (1996). Experiential Avoidance and Behavioral Disorders. JCCP. Experiential avoidance underlies wide range of behavioral disorders.",
            "Chawla & Ostafin (2007). Experiential Avoidance: Empirical Review. JCP. Associated with reduced goal engagement across domains.",
            "A-Tjak et al. (2015). ACT Meta-Analysis. Psychotherapy and Psychosomatics. ACT targeting experiential avoidance: g=0.57.",
            "Elliot & Church (1997). Approach and Avoidance Achievement Motivation. JPSP. Achievement motivation framing; tangentially relevant to experiential avoidance.",
            "Tyndall et al. (2019). AAQ-II Discriminant Validity Concerns. JCP. AAQ-II conflates with neuroticism/distress; measure limitations.",
        ],
        "evidence_level": "STRONG",
        "moderating_factors": "Emotional intensity | Psychological flexibility/mindfulness | History of emotional invalidation | Domain-specificity",
        "evidence_notes": "Backed by entire ACT research program. Most evidence from clinical populations; subclinical applicability less directly studied. DIFFERENTIATION: Distinct from C2-2 (Fear-Suppressed Desire). C2-3 involves diffuse emotional aversiveness (shame, anxiety, embarrassment) that may suppress awareness of the desire itself. The person may not consciously acknowledge wanting the goal. Neural pathway: anterior insula, medial PFC social cognition networks. DIAGNOSTIC: Do you actively want this goal but feel held back by a specific fear, or does the whole topic make you uncomfortable in a way that's hard to pin down? Specific fear blocking known desire → C2-2. Diffuse discomfort, may not acknowledge desire → C2-3.",
    },
    "C2-4": {
        "gate": "C2", "name": "Autonomy Threat (Reactance)",
        "definition": "External pressure or perceived control loss triggers oppositional motivation; reactance activates competing desire for freedom restoration (boomerang effect).",
        "mechanism": "Psychological reactance; perceived freedom threat generates competing desire for freedom restoration. Brehm (1966) showed reactance ENHANCES desire for restricted option (boomerang effect), not suppression. The person wants the restricted behavior MORE, not less.",
        "decision_comp": "IF Perceived_Control_Loss = TRUE THEN Competing_Desire_Activated(Freedom_Restoration); Original_Goal_Desire_Undermined",
        "diagnostic_signals": "'Don't tell me what to do' | 'The more they push, the less I want to' | External pressure motivation",
        "intervention": "Increase perceived control, Autonomy support, Reduce external pressure, Change structure.",
        "effect_size": "d = 0.3-0.5 (general population); d = 0.5-0.7 (high-stakes/high-reactance contexts)",
        "key_studies": [
            "Brehm (1966). A Theory of Psychological Reactance. Academic Press. Foundational theory: reactance ENHANCES desire for restricted option (boomerang effect), not suppression.",
            "Steindl et al. (2015). Understanding Psychological Reactance. Zeitschrift fur Psychologie. Comprehensive review confirming reactance as cognition + anger.",
            "Rains (2013). Nature of Psychological Reactance Revisited: Meta-Analysis. Human Communication Research. r=.20; freedom-threatening messages produce anger and oppositional motivation.",
            "Li & Shi (2025). Psychological Reactance and Persuasion: A Meta-Analysis. Human Communication Research. 28 articles, 146 effects; r=.17-.21; most comprehensive recent meta-analysis.",
        ],
        "evidence_level": "STRONG",
        "moderating_factors": "Trait reactance level | Perceived legitimacy of controlling agent | Importance of threatened freedom | Relationship with pressure source | High-stakes context",
        "evidence_notes": "50+ years of research. CORRECTION: Brehm (1966) describes desire ENHANCEMENT (boomerang effect), not suppression. 60 years of reactance research consistently show the restricted option becomes MORE desirable. Effect size adjusted to d=0.3-0.5 general (Rains r=.20; Li & Shi r=.17-.21); larger in high-stakes contexts.",
    },
    "C2-5": {
        "gate": "C2", "name": "Approach-Avoidance Conflict",
        "definition": "Simultaneous desire and anti-desire cancel motivation; net motivation = 0.",
        "mechanism": "Dual-process conflict; approach and avoidance systems simultaneously activated.",
        "decision_comp": "IF Approach_Force ~ Avoidance_Force THEN Oscillation + Anxiety + BIS_Activation",
        "diagnostic_signals": "'Part of me wants to, part of me doesn't' | Oscillating commitment | Decision paralysis about wanting",
        "intervention": "Reduce cost, Reframe reward, Increase value alignment.",
        "effect_size": "d = 0.4-0.6 (medium)",
        "key_studies": [
            "Miller (1944). Experimental Studies of Conflict. Ronald Press. Foundational: simultaneous gradients produce immobilization at intersection.",
            "McNaughton & Corr (2004). Two-Dimensional Neuropsychology of Defense. Neuroscience and Biobehavioral Reviews. BIS resolves approach-avoidance conflict via hippocampal-septal circuits.",
            "Aupperle et al. (2011). Reverse Translational Approach to Approach-Avoidance Conflict. BBR. Human behavioral paradigm validates Miller's model.",
            "Kirlic et al. (2017). Animal to Human Translational Paradigms. Behaviour Research and Therapy. Neural overlap confirmed in amygdala, ACC, striatal circuits.",
        ],
        "evidence_level": "STRONG",
        "moderating_factors": "Relative strength of approach vs avoidance | Trait BIS sensitivity | Goal proximity | Cognitive reappraisal availability",
        "evidence_notes": "Miller's model validated with modern neuroscience. 'Net motivation = 0' is simplification; conflict produces oscillation, anxiety, and BIS activation. Neural distinction is temporal/phenomenological rather than purely subcortical/cortical. Buttlar, Pauer, & van Harreveld (2024) AC/DC model in European Review of Social Psychology provides updated framework. Moscarello & Bhatt (2022) in Nature Communications identified two opposing hippocampus-to-PFC pathways for approach-avoidance. DIFFERENTIATION: Distinct from C5-6 (Value Contradiction). C2-5 is motivational ambivalence (BIS activation) -- competing approach and avoidance drives produce oscillation and felt tension at the desire stage. Pre-reflective and emotional. Prevents commitment formation. DIAGNOSTIC: Do you feel pulled in two directions about wanting this at all, or have you committed but it conflicts with something else you believe in? Torn about wanting it → C2-5. Committed but clashes with another value → C5-6.",
    },
    "C2-6": {
        "gate": "C2", "name": "External Motivation",
        "definition": "Goal pursued exclusively for approval/reward from others, not intrinsic value.",
        "mechanism": "SDT framework; external rewards undermine internalization and intrinsic motivation.",
        "decision_comp": "IF (Intrinsic_Motivation = 0) AND (Only_External_Rewards = TRUE) THEN Motivation_Fragile",
        "diagnostic_signals": "'I'm doing this to impress others' | 'Without external reward, I lose interest'",
        "intervention": "Increase value alignment, Autonomy support, Internalization facilitation.",
        "effect_size": "d = 0.3-0.5 (medium; Deci 1999 d=-0.28 to -0.40; previously overstated as d=0.6-0.9)",
        "key_studies": [
            "Deci et al. (1999). Extrinsic Rewards and Intrinsic Motivation: Meta-Analysis. Psychological Bulletin. 128 experiments: tangible rewards undermine intrinsic motivation (d=-0.28 to -0.40 depending on reward type).",
            "Ryan & Deci (2000). Intrinsic and Extrinsic Motivations. Contemporary Educational Psychology. SDT continuum.",
            "Sheldon & Elliot (1999). Self-Concordance Model. JPSP. External regulation = least persistence.",
            "Howard et al. (2017). SDT Continuum Structure: Meta-Analysis. Psychological Bulletin. External regulation = weakest association with persistence.",
        ],
        "evidence_level": "STRONG",
        "moderating_factors": "Type of extrinsic reward | Perceived autonomy support | Causality orientations | Controlling vs informational framing",
        "evidence_notes": "CORRECTION: Effect size corrected from d=0.6-0.9 to d=0.3-0.5 (LARGEST discrepancy in C2). Intervention changed from 'Pre-commitment' to 'Autonomy support' -- pre-commitment can exacerbate the external motivation problem; autonomy support facilitates internalization. Boundary condition: undermining effect is specific to tangible, expected, contingent rewards on interesting tasks. One of most well-established findings in motivation science. DIFFERENTIATION: C2-6 is a specific, identifiable subtype of the broader C2-1 (Value Misalignment) pattern. Distinguished by the person's ability to identify the external source of motivation. DIAGNOSTIC: Can you identify WHY this goal doesn't feel like yours? YES, doing it for someone else's approval/reward → C2-6. NO, it just doesn't resonate → C2-1.",
    },
    "C2-7": {
        "gate": "C2", "name": "Purpose Deficit",
        "definition": "No sense of purpose or meaning attached to goal; subjective value = 0 so vmPFC cannot assign value.",
        "mechanism": "Meaning-making failure; goal lacks connection to superordinate purpose structure.",
        "decision_comp": "IF Subjective_Value(Goal) = 0 THEN vmPFC_Value_Assignment = FAILED",
        "diagnostic_signals": "'What's the point?' | 'This doesn't matter' | Low pleasure anticipation",
        "intervention": "Increase value alignment, Reframe reward.",
        "effect_size": "d = 0.3-0.7 (lower bound widened; heterogeneous across populations)",
        "key_studies": [
            "Steger et al. (2006). Meaning in Life Questionnaire. Journal of Counseling Psychology. Low presence of meaning correlates with reduced goal-directed motivation.",
            "Frankl (1963). Man's Search for Meaning. Beacon Press. NON-EMPIRICAL (philosophical/clinical memoir). Will to meaning as primary motive; absence = existential vacuum. Theoretical origin only.",
            "McKnight & Kashdan (2009). Purpose in Life as System. Review of General Psychology. Purpose stimulates goal-directed behavior; absence = disengagement.",
            "Czekierda et al. (2017). Meaning in Life and Physical Health: Meta-Analysis. Health Psychology Review. NOTE: This is a meaning-to-HEALTH meta-analysis, not meaning-to-MOTIVATION directly. Partially misapplied.",
            "Boreham & Schutte (2023). Purpose and Well-Being Meta-Analysis. Journal of Clinical Psychology, k=99, N=66,468. Purpose broadly associated with well-being and goal engagement.",
        ],
        "evidence_level": "MODERATE",
        "moderating_factors": "Presence vs search for meaning | Age/developmental stage | Cultural/religious frameworks | Availability of community/vocation",
        "evidence_notes": "Association well-documented but mostly correlational. Causal direction debated. Frankl influential but NON-EMPIRICAL (philosophical/clinical memoir, zero quantitative evidence). Czekierda (2017) is meaning-to-health, not meaning-to-motivation -- partially misapplied. Rated MODERATE due to cross-sectional evidence.",
    },
    "C2-8": {
        "gate": "C2", "name": "Lack of Future Relevance",
        "definition": "Goal does not connect to long-term identity or future self.",
        "mechanism": "Temporal discounting inflates distance penalty; future benefit psychologically distant.",
        "decision_comp": "IF Temporal_Distance(Goal) > Discounting_Threshold THEN Present_Motivation = LOW",
        "diagnostic_signals": "'I won't care about this in 5 years' | 'This doesn't matter for my future'",
        "intervention": "Increase value alignment, Reframe reward.",
        "effect_size": "d = 0.3-0.5 (small to medium)",
        "key_studies": [
            "Ersner-Hershfield et al. (2009). Saving for the Future Self. SCAN. Less neural overlap for future self = steeper discounting.",
            "Hershfield (2011). Future Self-Continuity. Annals of NYAS. Psychological disconnection from future self increases discounting.",
            "Nurra & Oyserman (2018). From Future Self to Current Action. Self and Identity. Connected future self increases current action.",
            "Bartels & Urminsky (2011). On Intertemporal Selfishness. JCR. Reduced connectedness leads to 'intertemporal selfishness.'",
            "Stendardi et al. (2021). vmPFC and Future Self-Reference. SCAN. vmPFC lesion evidence for future self-reference processing.",
            "Carrillo et al. (2019). Best Possible Self Interventions. PLOS ONE. 29 studies; d=0.325 well-being, d=0.511 positive affect.",
        ],
        "evidence_level": "STRONG",
        "moderating_factors": "Time horizon length | Vividness of future self imagery | Age/developmental stage | Identity stability",
        "evidence_notes": "Well-supported with convergent evidence from neuro, behavioral econ, and social cognition. Most studied in financial saving domain; extrapolation to abstract goals involves some inference. DIFFERENTIATION: Distinct from C3-2 (Delayed Gratification Intolerance). C2-8 is an identity/connectedness problem ('cold' cognition) -- the future self feels like a stranger (Ersner-Hershfield et al., 2009). The person doesn't connect to the goal's future beneficiary. Intervention targets identity continuity. DIAGNOSTIC: Does the goal feel irrelevant to your future, or do you want it but can't stand waiting for the payoff? Irrelevant to future → C2-8. Want it but need results NOW → C3-2.",
    },

    # ---- C3: WILL ----
    "C3-1": {
        "gate": "C3", "name": "Self-Control Fatigue",
        "definition": "Reduced willingness to engage effortful control following prolonged cognitive demands; reflects motivational shift toward less effortful options rather than capacity loss.",
        "mechanism": "Process model (Inzlicht & Schmeichel): prolonged self-control shifts motivation from 'have-to' to 'want-to' tasks. Not resource depletion but motivational reallocation; fatigue signals opportunity cost of continued effort.",
        "decision_comp": "IF Sustained_Effort_Duration > Fatigue_Threshold THEN Motivation_Shifts_To_Less_Effortful_Options",
        "diagnostic_signals": "'I'm too tired' | 'I have nothing left' | Effort avoidance across domains | Recovery with increased incentives or task engagement",
        "intervention": "Effort pacing, Energy pacing, Reduce goal scope, Increase task engagement/autonomy.",
        "effect_size": "d = 0.31-0.35 with intensive protocols (Dang et al. 2025); near-zero with brief protocols",
        "key_studies": [
            "Baumeister et al. (1998). Ego Depletion: Is the Active Self a Limited Resource? JPSP. Original study (resource model now superseded).",
            "Hagger et al. (2016). Multilab Replication of Ego-Depletion Effect. Perspectives on Psychological Science. 23 labs, N=2,141, d=0.04 (null with brief protocols).",
            "Vohs et al. (2021). Multisite Paradigmatic Test. Psychological Science. 36 labs, N=3,531, d=0.06 (null with brief protocols).",
            "Dang et al. (2025). Multi-Lab Replication with Intensive Manipulation. 14 samples, N=2,078, d=0.31-0.35 (significant with 30-40 min protocols, I-squared=0).",
            "Inzlicht & Schmeichel (2012). Process Model of Self-Control Fatigue. Perspectives on Psychological Science, 7(5), 450-463. Motivation shift, not resource depletion.",
            "Berkman, Hutcherson, Livingston, Kahn, & Inzlicht (2017). Self-Control as Value-Based Choice. Current Directions in Psychological Science. Self-control is dynamic value computation, not resource expenditure.",
            "Wiehler, Brochier, & Pessiglione (2022). A Neuro-Metabolic Account of Why Daylong Cognitive Work Alters Choice. Current Biology, 32(16). Glutamate accumulation in lPFC after prolonged cognitive work.",
            "Pessiglione, Vinckier, Bhatt, & Wiehler (2025). MetaMotiF Framework. Trends in Cognitive Sciences. Bridges metabolic and motivational accounts of cognitive fatigue.",
            "Ordali et al. (2024). Prolonged Self-Control Causes Local Sleep in Frontal Cortex. PNAS. Neural fatigue mechanism.",
        ],
        "evidence_level": "MODERATE",
        "moderating_factors": "Task duration/intensity (30+ min shows effect) | Beliefs about willpower (Job, Dweck & Walton 2010) | Incentive level (increased incentives reverse 'depletion') | Sleep quality | Individual motivation",
        "evidence_notes": "REVISED FROM WEAK (Feb 2026). Original resource model (Baumeister 1998) failed two large pre-registered replications (d=0.04, d=0.06). However, Dang et al. (2025) multi-lab study with intensive 30-40 min protocols found consistent effects (d=0.31-0.35, I-squared=0). Current consensus: mechanism is metabolic-then-motivational -- prolonged cognitive work causes glutamate accumulation in lPFC (Wiehler et al. 2022), which triggers motivational shift away from effortful tasks (process model, Inzlicht & Schmeichel 2012). Not resource depletion but neurometabolic cost signaling. Effect is context-dependent: requires intensive/prolonged demands, reverses with increased incentives or engagement. Renamed from 'Energy Depletion' to 'Self-Control Fatigue' to reflect updated mechanism.",
    },
    "C3-2": {
        "gate": "C3", "name": "Delayed Gratification Intolerance",
        "definition": "Unable to tolerate delay between effort and reward; steep temporal discounting makes present cost unacceptable.",
        "mechanism": "Hyperbolic discounting; future reward drastically devalued against immediate effort cost.",
        "decision_comp": "IF Discount_Rate(Reward) > Effort_Cost_Tolerance THEN Commitment_Rejected",
        "diagnostic_signals": "'I need to see results now' | 'Why bother if payoff is years away'",
        "intervention": "Mini-deadlines, Reframe reward, Increase intrinsic cues.",
        "effect_size": "d = 0.4-0.6 (clinical); r = 0.10-0.20 (prediction of real-world outcomes; Mischel ~0.1 SD after SES controls)",
        "key_studies": [
            "Mazur (1987). Adjusting Procedure for Studying Delayed Reinforcement. Erlbaum (book chapter). Hyperbolic discounting function V=A/(1+kD).",
            "McClure et al. (2004). Separate Neural Systems Value Immediate and Delayed Rewards. Science. Limbic vs PFC/parietal systems.",
            "Kable & Glimcher (2007). Neural Correlates of Subjective Value. Nature Neuroscience. vmPFC tracks discounted value.",
            "Mischel et al. (1989). Delay of Gratification in Children. Science. Foundational delay research.",
        ],
        "evidence_level": "STRONG",
        "moderating_factors": "Reward magnitude | Certainty of delivery | Trait impulsivity | Visceral/emotional state",
        "evidence_notes": "Among most robust findings in neuroeconomics. Note: Watts, Duncan & Quan (2018) substantially revised Mischel's longitudinal predictive claims after controlling SES/cognitive ability. Core discounting mechanism still valid. DIFFERENTIATION: Distinct from C2-8 (Lack of Future Relevance). C3-2 is an impulse control problem ('hot' cognition) -- limbic vs prefrontal competition (McClure et al., 2004). The person DOES want the goal but can't tolerate the delay. Intervention targets impulse management (mini-deadlines, reward chunking). DIAGNOSTIC: Does the goal feel irrelevant to your future, or do you want it but can't stand waiting for the payoff? Irrelevant to future → C2-8. Want it but need results NOW → C3-2.",
    },
    "C3-3": {
        "gate": "C3", "name": "Effort Overestimation",
        "definition": "Perceived effort/discomfort judged excessive relative to reward; catastrophic effort estimation blocks commitment.",
        "mechanism": "Effort computation error; insula and ACC over-represent anticipated discomfort, leading to inflated effort cost estimates.",
        "decision_comp": "IF Perceived_Effort >> Actual_Effort THEN Commitment_Blocked",
        "diagnostic_signals": "'This will be horrible' | 'I can't handle that much work' | Task rationalization",
        "intervention": "Reduce cost, Micro-starts, Simulate context.",
        "effect_size": "FLAG: no meta-analytic anchor exists; estimated composite d = 0.3-0.5",
        "key_studies": [
            "Meyniel et al. (2013). When to Have a Break. PNAS. POSTERIOR insula (SII) and ventromedial thalamus track accumulated effort costs (not anterior insula as previously stated); model shows rational cost-based break-taking, not overestimation.",
            "Massar et al. (2015). Subjective Value During Delay and Effort Discounting. NeuroImage. ACC (not specifically dACC), anterior inferior frontal gyrus and DLPFC encoded effortful reward value.",
            "Kool et al. (2010). Decision Making and Avoidance of Cognitive Demand. JEP:General. Humans avoid effort even when objective cost minimal.",
            "Shenhav et al. (2017). Rational and Mechanistic Account of Mental Effort. Annual Review of Neuroscience. Comprehensive review.",
        ],
        "evidence_level": "MODERATE",
        "moderating_factors": "Current fatigue | Depression/anhedonia (Treadway et al. 2012) | Prior task experience | Interoceptive sensitivity",
        "evidence_notes": "CORRECTIONS: (1) Brain region errors fixed: Meyniel (2013) = POSTERIOR insula, not anterior; Massar (2015) = ACC, not specifically dACC. (2) Effect size flagged as estimated composite -- no meta-analytic anchor exists for 'effort overestimation' as a distinct construct. Most research documents effort discounting, not estimation accuracy. (3) Consider retitling to 'Effort Cost Overweighting.' Depression research (Treadway 2012) provides strongest evidence for pathological overestimation.",
    },
    "C3-4": {
        "gate": "C3", "name": "Insufficient Perceived Reward",
        "definition": "Expected reward is not salient or valued enough to justify effort.",
        "mechanism": "Effort-reward tradeoff computation: NAcc activation inadequate relative to effort cost. Dopamine's role is behavioral activation and effort allocation (Salamone & Correa, 2012), not hedonic 'liking' or 'incentive salience' (which is Berridge's framing).",
        "decision_comp": "IF NAcc_Activation(Reward) < Effort_Cost THEN Insufficient_Motivation",
        "diagnostic_signals": "'The payoff isn't worth it' | 'I don't care enough about the result'",
        "intervention": "Reframe reward, Increase perceived control.",
        "effect_size": "d = 0.3-0.5 (Blouzard 2023 g=0.43; Treadway 2012 d~0.39-0.41)",
        "key_studies": [
            "Salamone & Correa (2012). Mysterious Motivational Functions of Mesolimbic Dopamine. Neuron. CORRECTION: Salamone's emphasis is effort and behavioral activation, not hedonic 'liking'; 'incentive salience' is Berridge's framing.",
            "Berridge & Robinson (2016). Liking, Wanting, and Incentive-Salience Theory. American Psychologist. Wanting distinct from liking.",
            "Salamone et al. (2016). Activational Aspects of Motivation. Brain. Low NAcc dopamine linked to amotivation.",
            "Pessiglione et al. (2006). Dopamine-Dependent Prediction Errors. Nature. About reward LEARNING (prediction errors), not reward magnitude perception; tangentially relevant.",
            "Blouzard et al. (2023). Effort-Based Decision Making in Schizophrenia: A Meta-Analysis. JAMA Psychiatry. 20 studies, N=1,503; g=0.43.",
            "Husain & Roiser (2018). Neuroscience of Apathy and Anhedonia. Nature Reviews Neuroscience, 19. Critical distinction: apathy vs anhedonia.",
            "Robinson & Berridge (2025). Incentive-Sensitization Theory 30-Year Update. Annual Review of Psychology, 76. Updated wanting/liking framework.",
            "Salamone & Correa (2024). Activational Aspects of Motivation. Annual Review of Psychology, 75. Latest dopamine-effort synthesis.",
        ],
        "evidence_level": "STRONG",
        "moderating_factors": "Dopaminergic functioning | Anhedonia/depression | Competing reward availability | Neurobiological state",
        "evidence_notes": "CORRECTIONS: (1) Salamone attribution corrected from 'incentive salience' to 'effort and behavioral activation' -- 'incentive salience' is Berridge's framing, not Salamone's. (2) Effect size corrected d=0.5-0.8 → d=0.3-0.5 based on Blouzard 2023 meta-analysis g=0.43. (3) Pessiglione (2006) flagged as tangentially relevant (about reward learning, not reward magnitude). Among best-supported constructs overall. Important distinction: insufficient reward salience is not anhedonia (Husain & Roiser 2018).",
    },
    "C3-5": {
        "gate": "C3", "name": "Failure Cost Aversion",
        "definition": "Anticipated failure cost (emotional, identity) exceeds success reward; expected utility negative.",
        "mechanism": "Loss aversion at identity level; failure cost weighted more heavily than success reward.",
        "decision_comp": "IF Expected_Failure_Cost > Expected_Success_Reward THEN Net_Utility_Negative",
        "diagnostic_signals": "'What if I fail and everyone sees' | 'I'd rather not try than fail'",
        "intervention": "Reduce cost, Reframe reward, Rebuild efficacy.",
        "effect_size": "d = 0.5-0.7 (medium to large)",
        "key_studies": [
            "Kahneman & Tversky (1979). Prospect Theory. Econometrica. Losses weighted more heavily than gains (value function asymmetry; lambda~2.25 estimated in Tversky & Kahneman, 1992).",
            "Crocker & Park (2004). Costly Pursuit of Self-Esteem. Psychological Bulletin. Contingent self-worth creates failure-sensitive domains.",
            "Elliot & Church (1997). Approach and Avoidance Achievement Motivation. JPSP. Fear of failure predicts performance-avoidance goals.",
            "Atkinson (1957). Motivational Determinants of Risk-Taking. Psychological Review. When motive to avoid failure > motive to succeed, avoidance dominates.",
        ],
        "evidence_level": "STRONG",
        "moderating_factors": "Domain of self-worth contingency | Trait fear of failure | Fixed vs growth mindset | Public visibility of potential failure",
        "evidence_notes": "Converges from prospect theory, contingencies of self-worth, achievement motivation. Exact 2:1 loss aversion ratio questioned by Gal & Rucker (2018) as context-dependent.",
    },
    "C3-6": {
        "gate": "C3", "name": "Decision Paralysis (Analysis Paralysis)",
        "definition": "Multiple option combinations create intractable decision problem; endless deliberation loop prevents commitment.",
        "mechanism": "Option COMPARABILITY, not count, is the key variable: hard-to-compare options create evaluation deadlock. Effect is highly conditional on preference uncertainty, option similarity, and choice set complexity.",
        "decision_comp": "IF Option_Comparability = LOW AND Preference_Uncertainty = HIGH THEN Deliberation_Loop",
        "diagnostic_signals": "'I can't decide' | 'What if I choose wrong' | Endless research without action",
        "intervention": "Reduce goal scope, Increase specificity, Satisficing criteria, Change timing.",
        "effect_size": "d = 0.0-0.5 (HIGHLY CONDITIONAL; Scheibehenne et al. 2010 d=0.02 overall; effect real only under specific boundary conditions)",
        "key_studies": [
            "Iyengar & Lepper (2000). When Choice Is Demotivating. JPSP. 24 options = 3% purchase vs 30% for 6.",
            "Scheibehenne, Greifeneder, & Todd (2010). Can There Ever Be Too Many Options? JCR. d=0.02 overall -- directly contradicts large effect claim.",
            "Schwartz (2004). Paradox of Choice. HarperCollins. Excessive options produce anxiety and avoidance. NOTE: Popular science book, not peer-reviewed.",
            "Chernev et al. (2015). Choice Overload: Meta-Analysis. Journal of Consumer Psychology. Effect confirmed but heavily moderated by complexity, difficulty, preference uncertainty.",
            "Shafir et al. (1993). Reason-Based Choice. Cognition. Hard-to-compare options lead to deferral.",
        ],
        "evidence_level": "WEAK-to-MODERATE",
        "moderating_factors": "Number/similarity of options | Maximizer vs satisficer | Decision reversibility | Preference uncertainty",
        "evidence_notes": "CORRECTIONS: (1) Effect size corrected d=0.4-0.7 → d=0.0-0.5 (HIGHLY CONDITIONAL). Scheibehenne et al. 2010 meta-analysis found d=0.02 overall. (2) Rating downgraded to WEAK-to-MODERATE. (3) Chernev et al. journal corrected: Journal of Consumer Psychology (JCP), not JCR. (4) Iyengar & Lepper purchase rate corrected: 30%, not 31%. (5) Schwartz (2004) flagged as popular science book, not peer-reviewed. (6) Gate classification note: C3-6 is pre-decisional (deliberative phase), which raises questions about whether it belongs in C3 (Will/cost-aversion) vs C4 (Intention). DIFFERENTIATION: Distinct from C4-3 (Plan Selection Failure). C3-6 operates in the deliberative mindset (Gollwitzer, 1990) -- the person has not yet committed to a goal path. C4-3 operates in the implemental mindset. DIAGNOSTIC: Have you decided you definitely want to achieve this goal? YES but don't know which strategy → C4-3. NO, still deciding whether this is the right goal/path → C3-6.",
    },

    # ---- C4: INTENTION ----
    "C4-1": {
        "gate": "C4", "name": "Conflicting Plans",
        "definition": "Multiple commitments/goals creating priority conflict; attention fragments across incompatible demands.",
        "mechanism": "Executive resource competition; multiple active plans exceed allocation capacity.",
        "decision_comp": "IF Active_Plans > Executive_Capacity THEN Priority_Fragmentation",
        "diagnostic_signals": "'I have too many things going on' | 'I can't focus on just one'",
        "intervention": "Increase structure, Reduce goal scope, Change timing.",
        "effect_size": "d = 0.4-0.6 (medium)",
        "key_studies": [
            "Emmons & King (1988). Conflict Among Personal Strivings. JPSP. Conflict associated with negative affect, depression, psychosomatic complaints.",
            "Gray, Ozer, & Rosenthal (2017). Goal Conflict and Subjective Well-Being. Journal of Research in Personality. 54 studies, r=.34; definitive goal conflict meta-analysis.",
            "Gray & McNaughton (2000). Neuropsychology of Anxiety. Oxford University Press. BIS activated by goal conflict.",
            "Gorges & Grund (2017). Intraindividual Goal Conflict. Frontiers in Psychology. Goal conflict increases rumination and reduces progress.",
        ],
        "evidence_level": "STRONG",
        "moderating_factors": "Number/importance of competing goals | Degree of resource overlap | Individual executive function capacity | Inherent vs resource-based conflict",
        "evidence_notes": "CORRECTION: Inzlicht & Schmeichel (2012) ego depletion citation REPLACED with Gray et al. (2017) definitive goal conflict meta-analysis (multilab depletion failure d=0.06). Emmons & King robust and widely replicated. DIFFERENTIATION: Distinct from C5-1 (Competing Commitments). C4-1 is conscious resource conflict at the planning stage -- the person knows they have too many active goals and cannot form a coherent implementation plan. Operates before execution begins. DIAGNOSTIC: Are you struggling to organize your goals into a workable plan, or did you have a plan but keep getting pulled away by other obligations? Stuck at planning stage → C4-1. Had a plan but obligations override it / discover hidden conflicts during execution → C5-1.",
    },
    "C4-2": {
        "gate": "C4", "name": "Chronic Replanning",
        "definition": "Ongoing deliberation and planning prevents action; continuous replanning blocks execution.",
        "mechanism": "Planning becomes substitute for action; dopaminergic reward from planning itself.",
        "decision_comp": "IF Replanning_Frequency > Action_Frequency THEN Execution_Blocked",
        "diagnostic_signals": "'I need a better plan first' | Multiple plan iterations without execution | Generic planning",
        "intervention": "Increase specificity, Use implementation intentions, Mini-deadlines.",
        "effect_size": "d = 0.5-0.7 (medium to large)",
        "key_studies": [
            "Sirois & Pychyl (2013). Procrastination and Short-Term Mood Regulation. Social and Personality Psychology Compass. Planning can serve as mood-regulating substitute for action.",
            "Steel (2007). Nature of Procrastination: Meta-Analysis. Psychological Bulletin. Temporal motivation theory: planning about distant outcomes < proximate alternatives.",
            "Ainslie (2001). Breakdown of Will. Cambridge University Press. Repeated preference reversals create chronic replanning equilibrium.",
        ],
        "evidence_level": "MODERATE",
        "moderating_factors": "Individual impulsivity | Task aversiveness | Alternative mood-regulating activities | Abstract vs concrete planning",
        "evidence_notes": "'Dopaminergic reward from planning itself' theoretically grounded but lacks direct neuroimaging isolating planning-as-reward. Construct is synthesis, not independently studied phenomenon. Wang, Wang, & Gai (2021) MCII meta-analysis (g=0.336) supports implementation intentions as antidote to chronic replanning.",
    },
    "C4-3": {
        "gate": "C4", "name": "Plan Selection Failure",
        "definition": "Committed to goal but cannot select among viable plan options; implementation stalled despite endorsement.",
        "mechanism": "Option evaluation deadlock; multiple viable paths create selection paralysis.",
        "decision_comp": "IF Viable_Options > 1 AND Selection_Criteria = UNCLEAR THEN Stalled",
        "diagnostic_signals": "'I don't know which approach to take' | 'They all seem okay but I can't pick'",
        "intervention": "Increase specificity, Reduce goal scope, Change structure.",
        "effect_size": "d = 0.4-0.6 (medium)",
        "key_studies": [
            "Iyengar & Lepper (2000). When Choice Is Demotivating. JPSP. Excessive viable alternatives suppress action.",
            "Schwartz (2004). Paradox of Choice. Harper Perennial. POPULAR BOOK (not peer-reviewed). Maximizers experience greater paralysis.",
            "Chernev et al. (2015). Choice Overload: Meta-Analysis. Journal of Consumer Psychology. Effect conditional on complexity, difficulty, preference uncertainty, goal clarity.",
            "Scheibehenne, Greifeneder, & Todd (2010). Choice Overload Counter-Evidence. Journal of Consumer Research. d~0.02 overall; choice overload effect is near-zero without boundary conditions.",
            "Gollwitzer (1999). Implementation Intentions: Strong Effects. American Psychologist. Implemental mindset narrows focus to HOW, distinct from deliberative mindset.",
        ],
        "evidence_level": "STRONG",
        "moderating_factors": "Preference uncertainty | Decision complexity | Maximizer vs satisficer | Time pressure",
        "evidence_notes": "Well-established with boundary conditions. Application to action strategies vs consumer choice is extrapolation but mechanisms transfer. DIFFERENTIATION: Distinct from C3-6 (Decision Paralysis). C4-3 presupposes goal commitment -- the person is past the 'whether' question and stuck on 'how.' C3-6 is pre-commitment paralysis about which goal/path to pursue. DIAGNOSTIC: Have you decided you definitely want to achieve this goal? YES but don't know which strategy → C4-3. NO, still deciding whether this is the right goal/path → C3-6.",
    },
    "C4-4": {
        "gate": "C4", "name": "Cognitive Overload",
        "definition": "Plan complexity exceeds working memory capacity.",
        "mechanism": "Working memory capacity exceeded, plan coherence impossible; information processing saturated.",
        "decision_comp": "IF Plan_Complexity > Working_Memory_Capacity THEN Coherence_Impossible",
        "diagnostic_signals": "'This plan is too complicated' | 'I can't remember all the steps'",
        "intervention": "Increase specificity, Habit stacking.",
        "effect_size": "d = 0.6-0.9 (large)",
        "key_studies": [
            "Sweller (1988). Cognitive Load During Problem Solving. Cognitive Science. Working memory has strict capacity limits.",
            "Cowan (2001). Magical Number 4. BBS. WM capacity 3-5 chunks (revising Miller's 7+/-2).",
            "Paas & van Merrienboer (2020). Cognitive-Load Theory. Current Directions in Psychological Science. Overload produces confusion, frustration, coherence failure.",
            "Sweller, van Merrienboer, & Paas (2019). Cognitive Architecture and Instructional Design: 20 Years of CLT. Educational Psychology Review. 20-year CLT update.",
        ],
        "evidence_level": "STRONG",
        "moderating_factors": "Individual WM capacity | Plan chunking/schema formation | External scaffolding | Domain expertise",
        "evidence_notes": "35+ years of consistent support. Application to personal plan complexity is reasonable extrapolation from WM/CLT research.",
    },
    "C4-5": {
        "gate": "C4", "name": "Action Ambiguity",
        "definition": "Vague about what action to take right now; next-step specificity below threshold blocks intention formation.",
        "mechanism": "Insufficient action specification; action initiation is impaired without a concrete behavioral target.",
        "decision_comp": "IF Next_Step_Specificity < Action_Threshold THEN Intention_Formation_Blocked",
        "diagnostic_signals": "'I know I should do something but what?' | Untriggered intention cues",
        "intervention": "Increase specificity, Use implementation intentions, Immediate cues.",
        "effect_size": "d = 0.6-0.9 (large)",
        "key_studies": [
            "Locke & Latham (2002). Goal Setting and Task Motivation: 35-Year Odyssey. American Psychologist. Specific goals > vague 'do your best' in 1,000+ studies.",
            "Gollwitzer (1999). Implementation Intentions: Strong Effects. American Psychologist. If-then plans create situation-action links.",
            "Gollwitzer & Sheeran (2006). Implementation Intentions Meta-Analysis. Advances in Experimental Social Psychology. 94 tests, d=0.65.",
            "Bieleke, Keller, & Gollwitzer (2021). Implementation Intentions. European Review of Social Psychology, 32(1). Updated review with boundary conditions; WOOP/MCII integration.",
        ],
        "evidence_level": "STRONG",
        "moderating_factors": "Goal abstractness level | Environmental cue availability | Task novelty vs familiarity | Implementation intention formation",
        "evidence_notes": "Among most robust findings in motivation science. 50+ years of support. d=0.65 remarkably stable. Caveat: works best for simple discrete behaviors.",
    },
    "C4-6": {
        "gate": "C4", "name": "Recursive Justification",
        "definition": "Repeatedly generating new justifications to delay or modify plan; recursive justification delays commitment.",
        "mechanism": "Self-justification loop; each justification generates need for further justification.",
        "decision_comp": "IF Justification_Count > Threshold THEN Action_Delayed_Indefinitely",
        "diagnostic_signals": "'I have a good reason to wait' | 'Let me think about this more' | New excuses each time",
        "intervention": "Pre-commitment, Mini-deadlines, Ulysses pact.",
        "effect_size": "d = 0.4-0.6 (medium)",
        "key_studies": [
            "Festinger (1957). A Theory of Cognitive Dissonance. Stanford University Press. Inconsistency drives rationalization.",
            "Tavris & Aronson (2007). Mistakes Were Made. Houghton Mifflin Harcourt. POPULAR BOOK (not peer-reviewed). Self-justification is self-reinforcing 'pyramid of choice.'",
            "Jarcho et al. (2011). Neural Basis of Rationalization. SCAN. First fMRI of real-time dissonance reduction.",
        ],
        "evidence_level": "MODERATE",
        "moderating_factors": "Dissonance magnitude | Social context | Self-awareness/metacognition | External attribution availability",
        "evidence_notes": "Dissonance theory extremely well-established. But 'recursive justification as delay mechanism' is novel synthesis. Most dissonance research is post-decisional, not pre-actional justification loops.",
    },
    "C4-7": {
        "gate": "C4", "name": "Weak Cue-Response Binding",
        "definition": "Automatic behavior initiation from situational cues fails; cue-response binding strength below threshold.",
        "mechanism": "Implementation intention not encoded; situation-action link too weak for automatic activation.",
        "decision_comp": "IF Cue_Response_Strength < Automatic_Threshold THEN Manual_Initiation_Required",
        "diagnostic_signals": "Forgets to start even when conditions met | Needs external reminders | Weak cue-response binding",
        "intervention": "Habit stacking, Use implementation intentions, Environmental redesign.",
        "effect_size": "d = 0.6-0.8 (large)",
        "key_studies": [
            "Gollwitzer (1999). Implementation Intentions. American Psychologist. If-then plans create strong situation-action links.",
            "Wood & Neal (2007). Habits and the Habit-Goal Interface. Psychological Review. Habits driven by context cues, not goals.",
            "Lally et al. (2010). How Are Habits Formed. EJSP. Median 66 days for automaticity (range 18-254).",
            "Neal et al. (2006). Habits -- A Repeat Performance. Current Directions. Without cue-response links, behavior requires effortful initiation.",
            "Singh, Murphy, Maher, & Smith (2024). Habit Formation Meta-Analysis. Healthcare. 20 studies, N=2,601; most comprehensive habit formation meta-analysis.",
            "Gardner (2024). Current State of Habit Theory. Social and Personality Psychology Compass. Latest synthesis.",
            "Wood (2024). Latest Synthesis from Wood's Lab. Current Directions in Psychological Science.",
        ],
        "evidence_level": "STRONG",
        "moderating_factors": "Context consistency | Number of repetitions | Behavior complexity | Competing habitual responses",
        "evidence_notes": "Well-replicated across implementation intentions and habit research. Lally study small sample (N=96). 66-day median is for simple behaviors; complex actions likely need more.",
    },

    # ---- C5: COMMITMENT ----
    "C5-1": {
        "gate": "C5", "name": "Competing Commitments",
        "definition": "Multiple conflicting commitments create chronic attentional and motivational competition; competing goals win the value computation, diverting effort from primary goal.",
        "mechanism": "Motivational and attentional competition: self-control failures occur because competing goals win the value computation, not because a resource tank ran empty. Multiple active commitments create chronic interference in executive prioritization (Gray et al. 2017: r=.34 goal conflict-wellbeing).",
        "decision_comp": "IF Competing_Goal_Value > Primary_Goal_Value THEN Attention_Diverted; Primary_Goal_Undermined",
        "diagnostic_signals": "'I have too many obligations' | 'Something always comes up' | Goal displacement",
        "intervention": "Increase structure, Reduce goal scope, Add friction to alternatives, Goal prioritization exercises.",
        "effect_size": "d = 0.4-0.7 (medium to large)",
        "key_studies": [
            "Gray, Ozer, & Rosenthal (2017). Goal Conflict and Subjective Well-Being. Journal of Research in Personality. 54 studies, r=.34; definitive goal conflict meta-analysis.",
            "Emmons & King (1988). Conflict Among Personal Strivings. JPSP. Conflict degrades performance on all goals.",
            "Vohs et al. (2021). Multisite Paradigmatic Test of Ego Depletion. Psychological Science. 36 labs, N=3,531, d=0.06; ego depletion model no longer tenable.",
            "Berkman et al. (2017). Self-Control as Value-Based Choice. Current Directions in Psychological Science. Self-control is dynamic value computation, not resource expenditure.",
            "Inzlicht & Roberts (2024). The Fable of State Self-Control. Current Opinion in Psychology. Self-control failures reflect motivational priority, not resource depletion.",
            "Milyavskaya & Inzlicht (2017). What's So Great About Self-Control? SPPS. Successful self-regulation = fewer temptations experienced, not more resistance.",
        ],
        "evidence_level": "MODERATE",
        "moderating_factors": "Number of active commitments | Goal importance hierarchy clarity | Awareness of competing commitments | Executive function capacity",
        "evidence_notes": "CRITICAL CORRECTIONS: (1) Rating downgraded STRONG → MODERATE. (2) Depletion/resource mechanism REMOVED entirely -- ego depletion model no longer tenable (Vohs et al. 2021: 36 labs d=0.06; Hagger et al. 2016: 23 labs d=0.04). Self-control failures happen because competing goals win the value computation, not because a tank ran empty (Berkman et al. 2017; Inzlicht & Roberts 2024). (3) Kegan & Lahey (2009) 'Immunity to Change' is an unvalidated practitioner trade book with NO RCTs and NO controlled studies -- demoted to practitioner reference only, not primary evidence. Gray et al. (2017) provides definitive meta-analytic evidence for goal conflict effects. DIFFERENTIATION: Distinct from C4-1 (Conflicting Plans). C5-1 operates during commitment maintenance -- the person may not realize what's competing until they try to follow through. C4-1 is conscious planning-stage conflict. DIAGNOSTIC: Are you struggling to organize your goals into a workable plan, or did you have a plan but keep getting pulled away by other obligations? Stuck at planning stage → C4-1. Had a plan but obligations override it / discover hidden conflicts during execution → C5-1.",
    },
    "C5-2": {
        "gate": "C5", "name": "Environmental Disruption",
        "definition": "Physical or social environment disrupts goal pursuit; environmental trigger salience overrides goal salience.",
        "mechanism": "Environmental cue competition; salient environmental triggers override internal goal representations.",
        "decision_comp": "IF Environmental_Trigger_Salience > Goal_Salience THEN Goal_Overridden",
        "diagnostic_signals": "'My environment keeps pulling me away' | 'Too many distractions at home'",
        "intervention": "Environmental control, Alter environment, Add friction to alternatives.",
        "effect_size": "d = 0.5-0.8 (medium to large)",
        "key_studies": [
            "Wood & Neal (2007). Habits and the Habit-Goal Interface. Psychological Review. Context cues trigger habitual responses regardless of goals.",
            "Neal et al. (2012). Perceived and Actual Triggers of Habits. JESP, 48(2), 492-498. Strong habits primarily influenced by context, not goals.",
            "Papies (2016). Goal Priming as Situated Intervention. Current Opinion in Psychology. Environment both supports and undermines goal pursuit.",
            "Wood, Mazar, & Neal (2022). Habits and Goals in Human Behavior. Perspectives on Psychological Science. Definitive 2022 update on habit-goal interface.",
            "Wood (2024). Habits in Everyday Life. Current Directions in Psychological Science. Latest synthesis from Wood's lab.",
            "Mertens et al. (2022). Choice Architecture Interventions. PNAS, 200+ studies, N>2M, d=0.43. Largest nudging meta-analysis; environmental restructuring effective.",
        ],
        "evidence_level": "STRONG",
        "moderating_factors": "Strength of environmental habit triggers | Environmental stability | Cue-goal alignment | Individual cue sensitivity/impulsivity",
        "evidence_notes": "Robust through experience sampling and diary studies. Neal et al. year corrected from 2011 to 2012 (received 2011, published March 2012). Most research on simple habitual behaviors, not complex goal pursuit.",
    },
    "C5-3": {
        "gate": "C5", "name": "No External Monitoring",
        "definition": "No external pressure or visibility regarding commitment; lack of monitoring weakens felt obligation.",
        "mechanism": "Social accountability absent; without observation, commitment maintenance degrades.",
        "decision_comp": "IF External_Monitoring = 0 THEN Commitment_Decay_Rate = HIGH",
        "diagnostic_signals": "'Nobody knows if I do it or not' | 'It's just me'",
        "intervention": "Pre-commitment, Monitor progress, Raise abandonment cost.",
        "effect_size": "d = 0.32-0.48 (Harkin et al. 2016 meta-analytic 95% CI)",
        "key_studies": [
            "Matthews (2015). Impact of Commitment, Accountability, and Written Goals. Dominican University. Conference presentation (UNPUBLISHED, not peer-reviewed). Accountability doubled goal attainment (70% vs 35%).",
            "Lerner & Tetlock (1999). Accounting for Accountability. Psychological Bulletin. Accountability increases cognitive effort and commitment.",
            "Harkin et al. (2016). Monitoring Goal Progress: Meta-Analysis. Psychological Bulletin. 138 studies, N=19,951, d=0.40 (95% CI [0.32, 0.48]); larger when progress reported to others.",
            "Lee et al. (2021). Large-Scale stickK Platform Data. CHI '21. Real-world commitment contract platform analysis.",
            "Li, Liu, & Yu (2022). Felt Accountability. Frontiers in Psychology. Felt accountability increases performance but also exhaustion.",
        ],
        "evidence_level": "STRONG",
        "moderating_factors": "Quality of accountability relationship | Monitoring frequency | Whether monitor can impose consequences | Social desirability pressure",
        "evidence_notes": "Matthews study is UNPUBLISHED conference presentation, not peer-reviewed -- treat as directional only, not definitive. Harkin et al. meta-analysis provides strongest quantitative evidence. Long-term sustainability of accountability effects uncertain; external monitoring may not build intrinsic commitment.",
    },
    "C5-4": {
        "gate": "C5", "name": "Insufficient Stakes",
        "definition": "No real penalty for non-compliance; consequence magnitude below motivation threshold allows abandonment.",
        "mechanism": "Low abandonment cost; nothing significant lost by quitting.",
        "decision_comp": "IF Consequence(Abandonment) < Motivation_Threshold THEN Abandonment_Easy",
        "diagnostic_signals": "'Nothing bad happens if I stop' | 'Who cares if I don't'",
        "intervention": "Raise abandonment cost, Pre-commitment, Ulysses pact.",
        "effect_size": "d = 0.2-0.7 (heterogeneous; physical activity SMD=0.13; financial stakes higher)",
        "key_studies": [
            "Kahneman & Tversky (1979). Prospect Theory. Econometrica. Loss aversion: commitment sustained when abandonment carries meaningful losses.",
            "Ariely & Wertenbroch (2002). Procrastination, Deadlines, and Performance. Psychological Science. Self-imposed stakes improve performance. NOTE: Study 2 FAILED pre-registered replication (Hyndman & Bisin 2025).",
            "Gneezy & Rustichini (2000). A Fine Is a Price. JLS. Small fines can backfire by transforming social obligation into market transaction.",
            "Boonmanunt et al. (2023). Deposit Contracts for Health Behavior Change. Annals of Behavioral Medicine, 35 RCTs. RR=1.63 for deposit contracts.",
            "Brown et al. (2024). Loss Aversion Estimates. JEL, 62(2), 607 estimates. Lambda=1.955; ~50% of population loss-tolerant (Chapman et al. 2024).",
        ],
        "evidence_level": "STRONG",
        "moderating_factors": "Loss magnitude relative to goal value | Self-imposed vs externally imposed | Loss vs gain framing | Social vs material stakes | Crowding-out risk (external stakes may undermine intrinsic motivation)",
        "evidence_notes": "Prospect theory Nobel-Prize-winning work. Gneezy & Rustichini provides critical boundary condition. CAUTION: Ariely & Wertenbroch (2002) Study 2 failed pre-registered replication (Hyndman & Bisin 2025). Effect is heterogeneous -- strong for financial stakes, weak for health behavior (SMD=0.13). Risk of crowding out intrinsic motivation with external penalties.",
    },
    "C5-5": {
        "gate": "C5", "name": "Lack of Urgency",
        "definition": "Distant deadline or unclear deadline creates low-pressure environment; urgency perception insufficient.",
        "mechanism": "Temporal distance reduces motivational salience; low urgency permits indefinite delay.",
        "decision_comp": "IF Deadline_Distance > Urgency_Threshold THEN Motivational_Salience = LOW",
        "diagnostic_signals": "'I have plenty of time' | 'No rush' | Lack of urgency",
        "intervention": "Mini-deadlines, Change timing, Immediate cues.",
        "effect_size": "d = 0.4-0.6 (medium)",
        "key_studies": [
            "Steel (2007). Nature of Procrastination: Meta-Analysis. Psychological Bulletin. Motivation inversely proportional to delay (hyperbolic).",
            "Liberman & Trope (1998). Feasibility and Desirability in Temporal Construal. JPSP. Distant events = abstract, near = concrete.",
            "Ariely & Wertenbroch (2002). Procrastination and Deadlines. Psychological Science. Intermediate deadlines restore urgency.",
        ],
        "evidence_level": "STRONG",
        "moderating_factors": "Delay sensitivity/impulsivity | Self-imposed vs external deadlines | Intermediate milestones | Construal level",
        "evidence_notes": "Steel (2007) is most cited procrastination paper (~2,784 citations per Semantic Scholar). Impulsivity strongest predictor (r=0.41, 95% CI [.37, .46], k=22, N=4,005). CAUTION: Ariely & Wertenbroch (2002) cited here also has replication concerns (see C5-4).",
    },
    "C5-6": {
        "gate": "C5", "name": "Value Contradiction",
        "definition": "Goal conflicts with another held value; value conflict disrupts goal protection.",
        "mechanism": "Value system inconsistency; competing value activation undermines commitment stability.",
        "decision_comp": "IF Goal_Value CONFLICTS_WITH Held_Value THEN Commitment_Unstable",
        "diagnostic_signals": "'This goes against what I believe' | 'I feel torn' | Emotional interference",
        "intervention": "Increase value alignment, Reframe reward, Reduce cost.",
        "effect_size": "d = 0.3-0.6 (self-concordance anchor d=0.37)",
        "key_studies": [
            "Sheldon & Elliot (1999). Self-Concordance Model. JPSP. Value-incongruent goals receive less effort.",
            "Kegan & Lahey (2009). Immunity to Change. Harvard Business Press. PRACTITIONER REFERENCE (no RCTs, no controlled studies). Value contradictions create competing commitments.",
            "Brandstatter et al. (2013). Action Crisis. PSPB. Goal-value conflict triggers shift from implemental to deliberative mindset, elevated cortisol.",
            "Brandstatter & Schuler (2013). Action Crisis Experimental Evidence. JESP. Experimental action crisis induction.",
            "Oyserman (2024). Identity-Based Motivation Update. Social and Personality Psychology Compass. Updated IBM framework.",
        ],
        "evidence_level": "MODERATE",
        "moderating_factors": "Degree of value conflict | Awareness of contradiction | Importance of conflicting value | Whether conflict can be reframed",
        "evidence_notes": "Self-concordance well-supported longitudinally. Kegan & Lahey practice-based. Brandstatter provides physiological evidence. Most evidence correlational/longitudinal rather than experimental. DIFFERENTIATION: Distinct from C2-5 (Approach-Avoidance Conflict). C5-6 is cognitive values conflict (cortical value representation) -- the person committed to the goal but a competing explicitly-held value destabilizes commitment. Post-commitment erosion. DIAGNOSTIC: Do you feel pulled in two directions about wanting this at all, or have you committed but it conflicts with something else you believe in? Torn about wanting it → C2-5. Committed but clashes with another value → C5-6.",
    },
    "C5-7": {
        "gate": "C5", "name": "Identity Misalignment",
        "definition": "Goal inconsistent with identity or self-concept; triggers identity protection against goal-inconsistent behavior.",
        "mechanism": "Self-concept threat; goal pursuit threatens established identity narrative.",
        "decision_comp": "IF Goal INCONSISTENT_WITH Self_Concept THEN Identity_Protection_Activated",
        "diagnostic_signals": "'That's not who I am' | 'People like me don't do that'",
        "intervention": "Increase value alignment, Reframe reward, Rebuild efficacy.",
        "effect_size": "d = 0.4-0.7 (upper bound optimistic as general claim)",
        "key_studies": [
            "Oyserman (2009). Identity-Based Motivation. JCP. Identity-inconsistent goals: difficulty = 'not for people like me.'",
            "Oyserman (2024). Identity-Based Motivation Update. Social and Personality Psychology Compass. Latest IBM framework update.",
            "Higgins (1987). Self-Discrepancy Theory. Psychological Review. Actual-ideal discrepancies produce dejection.",
            "Sheldon & Houser-Marko (2001). Self-Concordance and Happiness: Upward Spiral. JPSP. Identity-aligned goals sustain effort.",
        ],
        "evidence_level": "STRONG",
        "moderating_factors": "Centrality of threatened identity | Reframing capability | Self-concept malleability | Social context",
        "evidence_notes": "Oyserman's IBM tested through RCTs in education. Higgins (5,000+ citations). Most research in structured contexts. DIFFERENTIATION: Distinct from C5-8 (Identity-Behavior Dissonance). C5-7 is PROSPECTIVE -- the person rejects a goal before sustained engagement because it doesn't fit current identity ('that's not who I am'). May occur sequentially: overcoming C5-7 but then repeatedly failing can lead to C5-8. DIAGNOSTIC: Is this goal something that doesn't fit who you are, or something you've been trying to do but keep falling short? Doesn't fit who I am → C5-7. Keep trying and failing → C5-8.",
    },
    "C5-8": {
        "gate": "C5", "name": "Identity-Behavior Dissonance",
        "definition": "Sustained incongruence between stated goal and actual behavior creates chronic self-concept conflict.",
        "mechanism": "Cognitive dissonance accumulation; persistent gap between aspiration and action erodes commitment.",
        "decision_comp": "IF (Stated_Goal - Actual_Behavior) > Dissonance_Threshold FOR Duration > T THEN Commitment_Erosion",
        "diagnostic_signals": "'I keep saying I'll change but never do' | 'I'm a fraud' | Chronic self-concept conflict",
        "intervention": "Self-compassion training, Goal revision/downsizing, Micro-starts, Monitor progress.",
        "effect_size": "d = 0.4-0.6 (ESTIMATED/INFERRED -- no direct meta-analytic anchor for this specific construct)",
        "key_studies": [
            "Festinger (1957). Cognitive Dissonance. Stanford University Press. Persistent inconsistency produces escalating discomfort.",
            "Higgins (1987). Self-Discrepancy Theory. Psychological Review. Chronic actual-ideal gaps accumulate dejection.",
            "Brandstatter et al. (2013). Action Crisis. PSPB. Sustained gaps trigger cortisol elevation and commitment erosion.",
            "Carver & Scheier (1998). Self-Regulation of Behavior. Cambridge University Press. Chronic negative discrepancy feedback produces disengagement.",
            "Neff (2023). Self-Compassion: Theory, Practice, and Evidence. Annual Review of Psychology, 74. Self-compassion buffers dissonance-driven disengagement.",
            "Wolf, Herrmann, & Brandstatter (2018). Self-Efficacy Moderates Action Crisis. Journal of Research in Personality. Self-efficacy moderates action crisis trajectory.",
        ],
        "evidence_level": "MODERATE",
        "moderating_factors": "Duration/magnitude of gap | Alternative dissonance-reduction strategies | Self-compassion (Neff 2003/2023) | Attribution style | Self-efficacy (Wolf et al. 2018)",
        "evidence_notes": "Dissonance and self-discrepancy theories foundational. Specific accumulation dynamic less directly studied as named construct. Brandstatter provides strongest direct evidence. Effect size is ESTIMATED/INFERRED from component literatures -- no direct meta-analytic anchor exists. Self-compassion training (Neff 2023) and goal revision are primary interventions for breaking the aspiration-failure cycle. DIFFERENTIATION: Distinct from C5-7 (Identity Misalignment). C5-8 is RETROSPECTIVE/ACCUMULATED -- the person committed and attempted but repeated failure creates a widening gap between aspiration and action ('I keep saying I'll change but never do'). Brandstatter et al. (2013) show this triggers action crisis with cortisol elevation. DIAGNOSTIC: Is this goal something that doesn't fit who you are, or something you've been trying to do but keep falling short? Doesn't fit who I am → C5-7. Keep trying and failing → C5-8.",
    },

    # ---- C6A: ACTION INITIATION ----
    "C6A-1": {
        "gate": "C6A", "name": "Motor Initiation Threshold",
        "definition": "Motor threshold to begin action is high; basal ganglia gate locked, inertia dominates.",
        "mechanism": "Motor initiation threshold elevated; striatal gate requires higher signal to open.",
        "decision_comp": "IF Motor_Signal < Activation_Threshold THEN Basal_Ganglia_Gate_Locked",
        "diagnostic_signals": "'I just can't start' | 'I sit there wanting to begin but can't' | Procrastination despite readiness",
        "intervention": "Reduce activation cost, Micro-starts, Reduce inertia.",
        "effect_size": "d = 0.4-0.8 (upper bound optimistic as general claim; strongest in clinical populations)",
        "key_studies": [
            "Klaus et al. (2019). What, If, and When to Move: Basal Ganglia Circuits. Annual Review of Neuroscience. Striatum modulates threshold for self-paced movement.",
            "Mink (1996). Basal Ganglia: Focused Selection and Inhibition. Progress in Neurobiology. Tonic inhibition of motor generators; sufficient signal needed to release.",
            "Cui et al. (2013). Concurrent Activation of Striatal Pathways. Nature. Both direct/indirect pathways co-activate before movement.",
            "Volkow et al. (2009). Dopamine Reward Pathway in ADHD. JAMA. PET: decreased dopamine receptors/transporters in ventral striatum.",
            "Ging-Jehli et al. (2025). Human Intracranial Recordings Validating Dynamic Threshold. PLOS Biology, 23(1). Direct human evidence for dynamic action threshold.",
            "Markowitz et al. (2023). Striatal Dopamine Gates Spontaneous Action Initiation. Nature, 614. Dopamine gates spontaneous action.",
            "Frank (2025). D1/D2 Opponency Framework Update. Annual Review of Neuroscience, 48. Updated computational model of basal ganglia gating.",
        ],
        "evidence_level": "STRONG",
        "moderating_factors": "Tonic dopamine level | Task reward salience | Fatigue/sleep deprivation | Comorbid ADHD or Parkinson's",
        "evidence_notes": "Go/No-Go gating well-established. 'Locked gate' metaphor oversimplifies. Most direct evidence from animal models; human evidence from neuroimaging/clinical populations. DIFFERENTIATION: Distinct from C0-1 (Environmental Friction). C6A-1 is internal motor threshold -- the person cannot start even when the environment is fully set up. C0-1 is environmental friction blocking action. DIAGNOSTIC: Is there something in your environment making it hard to start, or do you feel stuck even when everything is set up? Environment → C0-1. Stuck despite readiness → C6A-1.",
    },
    "C6A-2": {
        "gate": "C6A", "name": "Last-Minute Deliberation",
        "definition": "Continued deliberation at moment of action prevents initiation; deliberation delays motor execution.",
        "mechanism": "Executive override of motor system; prefrontal deliberation maintains inhibitory control at action moment.",
        "decision_comp": "IF Deliberation_Active_At_Action_Moment THEN Motor_Execution_Delayed",
        "diagnostic_signals": "'Wait, let me think about this one more time' | 'Am I sure?'",
        "intervention": "Pre-commitment, Immediate cues, Use implementation intentions.",
        "effect_size": "d = 0.4-0.6 (medium)",
        "key_studies": [
            "Tversky & Shafir (1992). Choice Under Conflict. Psychological Science. Disjunction effect: uncertainty prevents action.",
            "Gollwitzer (1999). Implementation Intentions. American Psychologist. Deliberative vs implemental mindsets.",
            "Le Bouc, R. & Pessiglione, M. (2022). A Neuro-Computational Account of Procrastination Behavior. Nature Communications, 13, 5753. dmPFC effort cost signal predicts deferral.",
            "Shafir et al. (1993). Reason-Based Choice. Cognition. Hard-to-justify choices deferred.",
        ],
        "evidence_level": "MODERATE",
        "moderating_factors": "Decision complexity | Trait indecisiveness | Time pressure | Stakes magnitude",
        "evidence_notes": "Deliberative-implemental distinction well-documented. Specific 'at the moment of action' deliberation less directly studied. CORRECTION: Le Bouc & Pessiglione (2022), not 'Zhang et al. (2022)' -- no such paper exists. Le Bouc & Pessiglione provide neuro-computational procrastination model; recent, awaits replication.",
    },
    "C6A-3": {
        "gate": "C6A", "name": "Attention Diversion",
        "definition": "Attention diverted immediately before action moment; competing stimulus salience captures attention at critical initiation moment.",
        "mechanism": "Attentional capture by competing stimuli at the precise moment of intended action initiation.",
        "decision_comp": "IF Competing_Stimulus_Salience > Goal_Salience AT Action_Moment THEN Initiation_Blocked",
        "diagnostic_signals": "'I was about to start but then...' | 'I got sidetracked right when I was going to begin'",
        "intervention": "Environmental control, Add friction to alternatives, Reduce friction.",
        "effect_size": "d = 0.5-0.7 (medium to large)",
        "key_studies": [
            "Theeuwes (2010). Top-Down and Bottom-Up Visual Selection. Acta Psychologica. Salient stimuli automatically capture attention.",
            "Folk et al. (1992). Involuntary Covert Orienting. JEP:HPP. Capture contingent on attentional control settings.",
            "Lavie (2005). Distracted and Confused? TICS. Perceptual load determines distraction vulnerability.",
            "Forster & Lavie (2008). Failures to Ignore Irrelevant Distractors. JEP:Applied. Even irrelevant distractors capture under low load.",
        ],
        "evidence_level": "STRONG",
        "moderating_factors": "Perceptual load | Stimulus salience | Working memory capacity | ADHD/attentional control deficits",
        "evidence_notes": "Among most extensively studied phenomena in cognitive psychology. Most research during ongoing tasks, not at initiation moment specifically. Pre-action = low load = more vulnerable. DIFFERENTIATION: Distinct from C6B-2 (Competing Stimuli). Shares the same core mechanism (attentional capture) but at the initiation moment specifically. Pre-action = zero perceptual load = maximally vulnerable to capture. Intervention emphasis: reduce start-friction to shorten the vulnerability window. DIAGNOSTIC: Do you get distracted when you're trying to start, or after you've already begun working? Right at start → C6A-3. After starting → C6B-2. Both → flag both, assess which is primary.",
    },
    "C6A-4": {
        "gate": "C6A", "name": "Waiting for Readiness",
        "definition": "Waits for ideal conditions/mood/readiness to start; unmet readiness conditions block action.",
        "mechanism": "Conditional action policy; self-imposed prerequisites prevent initiation.",
        "decision_comp": "IF Readiness_Conditions_Met = FALSE THEN Action_Postponed",
        "diagnostic_signals": "'I'll start when I feel ready' | 'The timing isn't right' | 'I need to be in the mood'",
        "intervention": "Reduce activation cost, Immediate cues, Micro-starts.",
        "effect_size": "d = 0.5-0.7 (medium to large)",
        "key_studies": [
            "Sirois & Pychyl (2013). Procrastination and Mood Regulation. SPPC. 'I'll start when I feel better' = mood-contingent action policy.",
            "Steel (2007). Nature of Procrastination. Psychological Bulletin. Task aversiveness, impulsiveness strongest predictors.",
            "Gollwitzer (1999). Implementation Intentions. American Psychologist. If-then plans bypass readiness requirements.",
            "Ferrari et al. (1995). Procrastination and Task Avoidance. Plenum Press. Chronic procrastinators believe required mood will arrive later.",
        ],
        "evidence_level": "STRONG",
        "moderating_factors": "Trait perfectionism | Emotional regulation capacity | Task aversiveness | Mood-as-prerequisite belief strength",
        "evidence_notes": "Steel meta-analysis comprehensive. Overlap with C6B-3 (Mood-Dependent Execution) -- same mechanism at different phases.",
    },
    "C6A-5": {
        "gate": "C6A", "name": "Over-Simulation",
        "definition": "Over-rehearsal/mental practice prevents moving to actual action; outcome simulation reduces physiological energization needed for action.",
        "mechanism": "Positive outcome fantasies reduce systolic blood pressure (Kappes & Oettingen, 2011), lowering physiological energization below action threshold. Critical distinction: OUTCOME simulation (imagining success) reduces effort; PROCESS simulation (imagining steps) increases effort (Pham & Taylor, 1999).",
        "decision_comp": "IF Outcome_Simulation_Count > Threshold THEN Physiological_Energization_Reduced; Action_Drive_Below_Threshold",
        "diagnostic_signals": "'I've thought about it a million times' | 'I know exactly what to do but can't start'",
        "intervention": "Mental contrasting (WOOP), Process simulation instead of outcome simulation, Immediate cues, Reduce inertia.",
        "effect_size": "d = 0.3-0.5 (small to medium)",
        "key_studies": [
            "Kappes & Oettingen (2011). Positive Fantasies Sap Energy. JESP. Positive fantasies lower systolic BP (reduced energization), not motor capacity depletion.",
            "Kappes & Oettingen (2016). Mental Simulation as Substitute for Experience. SPPC. Simulating achievement substitutes for actual pursuit.",
            "Pham & Taylor (1999). Thought to Action: Process vs Outcome Simulation. PSPB. CRITICAL: Outcome simulation = complacency; process simulation = INCREASED effort.",
            "Oettingen & Mayer (2002). Motivating Function of Thinking About Future. JPSP. Positive fantasies predicted low effort across 4 domains.",
            "Oettingen (2012). Future Thought and Behaviour Change. ERSP. Comprehensive mental contrasting review; MCII as intervention.",
        ],
        "evidence_level": "MODERATE",
        "moderating_factors": "Process vs outcome simulation type | Mental contrasting usage | Initial desire strength | Trait optimism",
        "evidence_notes": "CORRECTION: 'Motor Command Capacity Depletion' does not appear in scientific literature and has been removed. Oettingen's work shows reduced physiological energization (BP drop), not motor capacity depletion. The mechanism is that positive fantasies provide premature satisfaction, reducing the motivational gap that drives action. Intervention implication: mental contrasting (WOOP/MCII) counters this by pairing fantasy with obstacle recognition.",
    },
    "C6A-6": {
        "gate": "C6A", "name": "Self-Sabotage",
        "definition": "Unconscious sabotage creates barrier to action; defensive avoidance creates self-imposed barriers.",
        "mechanism": "Protective self-handicapping; creating obstacles preserves self-concept against potential failure.",
        "decision_comp": "IF Failure_Threat = HIGH THEN Create_Self_Imposed_Barriers",
        "diagnostic_signals": "Repeated 'accidents' preventing action | 'Something always goes wrong' | Freeze/task switching",
        "intervention": "Rebuild efficacy, Reduce cost, Increase perceived control.",
        "effect_size": "d = 0.4-0.7 (medium to large)",
        "key_studies": [
            "Jones & Berglas (1978). Self-Handicapping Strategies. PSPB. Participants chose performance-inhibiting drug to create external attribution.",
            "Schwinger et al. (2014). Self-Handicapping and Achievement: Meta-Analysis. JEP. 36 studies, N=25,550: consistent negative relationship.",
            "Schwinger, Trautner, Putz, et al. (2022). Antecedents of Self-Handicapping: Meta-Analysis. Journal of Educational Psychology, 114(3), 576-596. 159 studies, N=81,630. Low self-esteem, performance-avoidance goals, entity theories strongest antecedents.",
            "Urdan & Midgley (2001). Academic Self-Handicapping. Educational Psychology Review. Claimed vs behavioral self-handicapping distinction.",
        ],
        "evidence_level": "STRONG",
        "moderating_factors": "Self-esteem instability | Identity-centrality of performance domain | Public vs private context | Gender",
        "evidence_notes": "Well-established since 1978 with large meta-analyses. 'Unconscious' aspect harder to verify -- may be strategic rather than unconscious.",
    },
    "C6A-7": {
        "gate": "C6A", "name": "Momentum Dependency",
        "definition": "Cannot cold-start despite desire/planning/capability; continuation easy once momentum achieved. Common in ADHD (low tonic dopamine). Initiation threshold >> continuation threshold.",
        "mechanism": "Dopaminergic initiation deficit; tonic dopamine insufficient for cold-start but adequate for continuation.",
        "decision_comp": "IF Initiation_Threshold >> Continuation_Threshold THEN Cold_Start_Failure",
        "diagnostic_signals": "'Once I start I'm fine, it's just starting' | 'I need a push to get going'",
        "intervention": "Micro-starts, Reduce inertia, Environmental redesign.",
        "effect_size": "d = 0.4-0.8 (lower bound = behavioral strategies; upper bound = pharmacological in ADHD)",
        "key_studies": [
            "Volkow et al. (2009). Dopamine Reward Pathway in ADHD. JAMA. PET: decreased dopamine receptors/transporters in ventral striatum.",
            "Volkow et al. (2011). Motivation Deficit in ADHD. Molecular Psychiatry. Dopamine markers specifically associated with motivation.",
            "Volkow et al. (2012). Methylphenidate-Elicited Dopamine Increases. J Neuroscience. Dopamine increase magnitude predicted treatment response.",
            "Arnsten (2006). Stimulants: Therapeutic Actions in ADHD. Neuropsychopharmacology. PFC requires optimal catecholamine levels.",
            "Faraone et al. (2021). Comprehensive ADHD Consensus. Neuroscience and Biobehavioral Reviews, 128. International consensus on ADHD neurobiology.",
            "Barkley (2012). Executive Functions. Guilford. Task initiation as distinct EF domain.",
        ],
        "evidence_level": "STRONG",
        "moderating_factors": "Stimulant medication status | Task novelty/intrinsic reward | Time of day/circadian dopamine | Comorbid anxiety",
        "evidence_notes": "Dopamine hypothesis well-supported by genetics, neuroimaging, pharmacology. 'Cold-start vs continuation' distinction clinically observed but not isolated as discrete experimental variable. Tonic/phasic dopamine claim should be qualified -- the distinction is well-established in animal models but human evidence is largely indirect (PET, pharmacological). CROSS-REFERENCE: C6A-7 shares significant dopaminergic literature with C6A-1 (Motor Initiation Threshold). C6A-7 is a specific clinical subtype with pronounced initiation-continuation asymmetry, particularly in ADHD.",
    },
    "C6A-8": {
        "gate": "C6A", "name": "Trauma-Related Freeze",
        "definition": "Threat-protective immobility (evolutionary survival mechanism); defense cascade via amygdala-PAG circuitry overrides volitional control.",
        "mechanism": "Defense cascade via amygdala-periaqueductal gray (PAG) circuitry: threat detection triggers escalating defensive responses (arousal → fight/flight → freeze → tonic immobility → collapse). Amygdala-PAG pathway is well-established independently of polyvagal theory.",
        "decision_comp": "IF Threat_Detection = TRUE THEN Amygdala_PAG_Defense_Cascade_Activated; Volitional_Control_Overridden",
        "diagnostic_signals": "Physical immobility | 'I shut down' | 'I can't move or think'",
        "intervention": "Requires trauma-informed therapy (EMDR, SE, CPT). Not addressable through behavioral intervention alone.",
        "effect_size": "Clinical -- requires specialized treatment",
        "key_studies": [
            "Kozlowska et al. (2015). Fear and the Defense Cascade. Harvard Review of Psychiatry. Defense cascade: arousal -> fight/flight -> freeze -> tonic immobility -> collapse.",
            "Hagenaars et al. (2014). Updating Freeze. Neuroscience and Biobehavioral Reviews. Freeze vs tonic immobility as distinct states.",
            "Roelofs (2017). Freeze for Action. Phil Trans Roy Soc B. Amygdala-PAG circuit mechanisms; adaptive but can become maladaptive.",
            "Neuhuber & Berthoud (2022). Lack of Evidence for Dorsal Vagal Complex as Integrative Autonomic Center. Biological Psychology. Critical neuroanatomical challenge to polyvagal theory premises.",
            "Grossman (2023). Polyvagal Theory: Systematic Critique. Biological Psychology. Systematic challenge to all 5 PVT premises.",
        ],
        "evidence_level": "MODERATE",
        "moderating_factors": "Trauma history/PTSD severity | Perceived controllability | Developmental stage at trauma | Co-regulation availability",
        "evidence_notes": "CORRECTION: Defense cascade mechanism via amygdala-PAG is well-established (Kozlowska 2015, Roelofs 2017). Polyvagal theory (Porges 2011) has been significantly challenged: Neuhuber & Berthoud (2022) refute the neuroanatomical premises; Grossman (2023) challenges all 5 core PVT premises. The freeze response as phenomenon is robust; the specific polyvagal mechanism is no longer tenable as sole explanation. Porges (2025, Clinical Neuropsychiatry) has responded to critics. Gate note: C6A-8 is qualitatively different from C6A-1 through C6A-7 -- those are volitional initiation failures; C6A-8 is involuntary autonomic override bypassing volition entirely.",
    },

    # ---- C6B: ACTION PERSISTENCE ----
    "C6B-1": {
        "gate": "C6B", "name": "Task Boredom",
        "definition": "Activity becomes monotonous; hedonic adaptation degrades reward salience, interest wanes mid-task.",
        "mechanism": "Hedonic adaptation; reward signal attenuates with repetition reducing continuation motivation.",
        "decision_comp": "IF Reward_Salience(Task) < Continuation_Threshold THEN Engagement_Drops",
        "diagnostic_signals": "'This is boring' | 'I've lost interest' | Dropout mid-task",
        "intervention": "Increase intrinsic cues, Reframe reward, Change timing.",
        "effect_size": "d = 0.4-0.6 (medium)",
        "key_studies": [
            "Frederick & Loewenstein (1999). Hedonic Adaptation. In Kahneman et al. (Eds.), Well-Being. Russell Sage. Hedonic treadmill reduces response to repeated stimuli.",
            "Luo et al. (2022). Neural Habituation to Hedonic and Eudaimonic Rewards. Psychophysiology. Reward positivity decreases for hedonic but sustains for eudaimonic rewards.",
            "Lyubomirsky (2011). Hedonic Adaptation. Oxford Handbook. Variety and novelty slow adaptation.",
            "Eastwood et al. (2012). Unengaged Mind: Defining Boredom. Perspectives on Psychological Science. Boredom = failure to engage attention satisfyingly.",
            "Tam et al. (2021). Integrated Boredom as Attentional Feedback Loop. Personality and Social Psychology Review, 25(3). Boredom signals unmet engagement needs.",
            "Nature Communications Psychology (2025). Boredom and Arousal Multilevel Meta-Analysis. 214 effects, N=6,570.",
        ],
        "evidence_level": "STRONG",
        "moderating_factors": "Task variety/novelty | Autonomous motivation | Optimal stimulation level | Social engagement during task",
        "evidence_notes": "Hedonic adaptation robust. Luo et al. (2022) recent, awaits replication. Application to task persistence is extrapolation from well-being research.",
    },
    "C6B-2": {
        "gate": "C6B", "name": "Competing Stimuli (Distraction)",
        "definition": "Susceptible to pulling away from task by external/internal temptations; alternative NAcc activation disrupts goal-shielding. Includes both involuntary attentional capture (notifications, novel stimuli) and voluntary comfort-seeking (choosing easier/more pleasant alternatives such as social media, food, rest).",
        "mechanism": "Attentional competition; alternative reward sources capture attention away from current task.",
        "decision_comp": "IF Alternative_NAcc_Activation > Task_NAcc_Activation THEN Goal_Shielding_Fails",
        "diagnostic_signals": "'I keep checking my phone' | 'I can't focus' | Task switching | 'That activity is more comfortable' | 'I'd rather do X'",
        "intervention": "Environmental control, Add friction to alternatives, Increase structure, Effort pacing.",
        "effect_size": "d = 0.5-0.7 (upper bound slightly generous)",
        "key_studies": [
            "Lavie (2005). Distracted and Confused? TICS. Spare capacity automatically processes distractors.",
            "Berridge & Robinson (2016). Incentive-Sensitization Theory. American Psychologist. Environmental cues trigger wanting.",
            "Steel (2007). Nature of Procrastination. Psychological Bulletin. Impulsiveness among strongest procrastination predictors.",
            "Hofmann et al. (2012). Everyday Temptations. JPSP. 7,827 episodes: media, leisure, sleep most yielded to.",
            "McClure et al. (2004). Separate Neural Systems Value Immediate and Delayed Rewards. Science. Limbic preferentially activated by immediately available rewards.",
            "Anderson (2021). Reward-Driven Attentional Capture Meta-Analysis. Psychological Bulletin, 147(4). Reward history captures attention involuntarily.",
            "Milyavskaya & Inzlicht (2017). Temptation Strength Predicts Goal Failure. SPPS, 8(8). Temptation exposure, not resistance, predicts failure.",
        ],
        "evidence_level": "STRONG",
        "moderating_factors": "Perceptual load | Environmental stimulus richness | Working memory capacity | ADHD/trait impulsivity | Proximity/accessibility of comfort options | Discounting steepness",
        "evidence_notes": "Strongest convergent evidence across neuroscience, cognitive psych, behavioral econ. Ego depletion explanation contested; competing-rewards mechanism well-supported. Subtypes include stimulus-driven capture (involuntary, bottom-up) and hedonic comfort-seeking (voluntary, top-down). Both share the core mechanism of alternative reward outcompeting task reward. DIFFERENTIATION: Distinct from C6A-3 (Attention Diversion). Shares the same core mechanism (attentional capture) but during sustained effort. Intervention emphasis: increase persistence structure (environmental control, task scaffolding). Both may co-occur. DIAGNOSTIC: Do you get distracted when you're trying to start, or after you've already begun working? Right at start → C6A-3. After starting → C6B-2. Both → flag both, assess which is primary.",
    },
    "C6B-3": {
        "gate": "C6B", "name": "Mood-Dependent Execution",
        "definition": "Action output fluctuates with daily mood/energy; execution contingent on mood state.",
        "mechanism": "Mood-contingent action policy; no decoupling between affective state and behavioral output.",
        "decision_comp": "IF Mood_State = LOW THEN Action_Output = LOW (regardless of plan)",
        "diagnostic_signals": "'I can only do it when I feel good' | 'Bad days I can't function'",
        "intervention": "Use implementation intentions, Increase structure, Effort pacing.",
        "effect_size": "d = 0.4-0.6 (medium)",
        "key_studies": [
            "Gendolla (2000). Impact of Mood on Behavior. Review of General Psychology. Mood-behavior model: mood signals task difficulty.",
            "Sirois & Pychyl (2013). Procrastination and Mood Regulation. SPPC. Procrastination = short-term mood repair priority.",
            "Schwarz & Clore (1983). Mood, Misattribution, and Judgments. JPSP. Affect-as-information: mood used as heuristic.",
            "Gendolla & Silvestrini (2010). Implicit 'Go.' Psychological Science. Mood effects operate even with subliminal primes.",
            "Gendolla (2025). 25-Year Mood-Effort Coupling Review. Psychological Science. Comprehensive review of mood-behavior model.",
            "Rackemann et al. (2024). Implementation Intentions in Mood Contexts. Stress and Health. II effectiveness moderated by mood.",
        ],
        "evidence_level": "STRONG",
        "moderating_factors": "Emotional regulation skills | Implementation intentions | Task intrinsic motivation | Dispositional mindfulness",
        "evidence_notes": "Decades of support. Overlap with C6A-4 (same mechanism, different temporal phase).",
    },
    "C6B-4": {
        "gate": "C6B", "name": "Single-Error Catastrophizing",
        "definition": "Single small error triggers premature task abandonment; catastrophizing triggers abandonment.",
        "mechanism": "All-or-nothing thinking; single deviation from plan treated as total failure.",
        "decision_comp": "IF Error_Count >= 1 THEN Catastrophize; Abandon_Task",
        "diagnostic_signals": "'I messed up so why bother' | 'One mistake ruins everything'",
        "intervention": "Reframe reward, Cost management, Monitor progress.",
        "effect_size": "d = 0.4-0.7 (medium to large)",
        "key_studies": [
            "Beck et al. (1979). Cognitive Therapy of Depression. Guilford. All-or-nothing thinking as core cognitive distortion.",
            "Frost et al. (1990). Dimensions of Perfectionism. CTR. 'Concern over Mistakes' as maladaptive perfectionism dimension.",
            "Burns (1980). Perfectionist's Script for Self-Defeat. Psychology Today. Giving up pattern after minor errors.",
            "Egan et al. (2011). Perfectionism as Transdiagnostic Process. CPR. Concern over mistakes predicts disengagement across disorders.",
            "Marlatt & Gordon (1985). Relapse Prevention. Guilford. Abstinence Violation Effect: single lapse triggers full relapse.",
            "Neff (2023). Self-Compassion. Annual Review of Psychology, 74. Self-compassion as error buffer against catastrophizing.",
        ],
        "evidence_level": "STRONG",
        "moderating_factors": "Trait perfectionism (Concern over Mistakes) | Depression severity | Self-compassion | Growth vs fixed mindset",
        "evidence_notes": "All-or-nothing thinking among most established cognitive distortions. Specific abandonment consequence less directly studied experimentally vs other responses.",
    },
    "C6B-5": {
        "gate": "C6B", "name": "Perfectionism",
        "definition": "Unwillingness to continue with imperfect output; perfectionistic standards halt progress.",
        "mechanism": "Quality threshold set above achievable level; imperfection intolerance blocks continuation.",
        "decision_comp": "IF Output_Quality < Perfection_Standard THEN Halt_Progress",
        "diagnostic_signals": "'It's not good enough' | 'I need to redo this' | Endless revision",
        "intervention": "Reduce cost, Reframe reward, Mini-deadlines.",
        "effect_size": "d = 0.4-0.7 (medium to large)",
        "key_studies": [
            "Hewitt & Flett (1991). Perfectionism in Self and Social Contexts. JPSP. Multidimensional Perfectionism Scale: self-oriented, other-oriented, socially prescribed.",
            "Frost et al. (1990). Dimensions of Perfectionism. CTR. Concern over Mistakes and Doubts about Actions as maladaptive dimensions.",
            "Stoeber & Otto (2006). Positive Conceptions of Perfectionism. PSPR. Strivings (adaptive) vs concerns (maladaptive): concerns predict abandonment.",
            "Curran & Hill (2019). Perfectionism Increasing Over Time: Meta-Analysis. Psychological Bulletin. All three dimensions increasing across birth cohorts 1989-2016.",
            "Bellam et al. (2025). Strivings-Concerns Distinction Confirmed. Journal of Occupational and Organizational Psychology, 28 samples, N=9,560.",
        ],
        "evidence_level": "STRONG",
        "moderating_factors": "Perfectionism dimension (concerns >> strivings) | Self-compassion | External deadlines | Social evaluation context",
        "evidence_notes": "Extensive research, two major validated scales. Strivings-concerns distinction well-replicated. Most evidence correlational. Curran & Hill adds ecological urgency.",
    },
    "C6B-6": {
        "gate": "C6B", "name": "Invisible Progress",
        "definition": "Motivation drops when progress is invisible or slower than expected; motivation contingent on visible progress signal.",
        "mechanism": "Progress feedback loop broken; without visible progress, continuation motivation decays.",
        "decision_comp": "IF Visible_Progress < Expected_Progress THEN Motivation_Decay",
        "diagnostic_signals": "'I'm not getting anywhere' | 'Nothing is changing' | Fatigue-driven avoidance",
        "intervention": "Progress feedback, Monitor progress, Mini-deadlines.",
        "effect_size": "d = 0.4-0.7 (Harkin 2016 d=0.40; upper bound with sub-milestone and visual tracking)",
        "key_studies": [
            "Harkin et al. (2016). Monitoring Goal Progress: Meta-Analysis. Psychological Bulletin. 138 studies, N=19,951, d=0.40 (95% CI [0.32, 0.48]); larger when recorded or public.",
            "Amabile & Kramer (2011). Progress Principle. Harvard Business Press. 12,000+ diary entries: meaningful progress = strongest predictor of positive inner work life.",
            "Locke & Latham (2002). Goal Setting: 35-Year Odyssey. American Psychologist. Goals without feedback = attenuated effects.",
            "Koo & Fishbach (2012). Small-Area Hypothesis. JCR. Motivation highest when focusing on smaller proportion of work.",
        ],
        "evidence_level": "STRONG",
        "moderating_factors": "Feedback frequency/tangibility | Public vs private recording | Task time horizon | External vs self-generated tracking",
        "evidence_notes": "Harkin meta-analysis highly rigorous. Amabile & Kramer observational but ecologically rich. Locke & Latham among most cited theories in organizational psych.",
    },
    "C6B-7": {
        "gate": "C6B", "name": "High Restart Cost",
        "definition": "Restarting task after interruption/break is as difficult as initial start; momentum loss creates high restart cost.",
        "mechanism": "Context-switching cost; mental state reconstruction after interruption requires full re-initiation.",
        "decision_comp": "IF Restart_Cost ~ Initiation_Cost THEN Resumption_Blocked",
        "diagnostic_signals": "'Once I stop I can't get back to it' | 'Breaks ruin my flow'",
        "intervention": "Reduce inertia, Micro-starts, Increase structure.",
        "effect_size": "d = 0.5-0.7 (medium to large)",
        "key_studies": [
            "Altmann & Trafton (2002). Memory for Goals: Activation-Based Model. Cognitive Science. Goal activation decays during interruption; resumption requires re-retrieval.",
            "Monsell (2003). Task Switching. TICS. Switch costs include reconfiguration time and residual cost (~600ms asymptote).",
            "Mark et al. (2008). Cost of Interrupted Work. CHI. After interruption, 23+ min to refocus; higher stress and frustration.",
            "Trafton et al. (2003). Preparing to Resume Interrupted Task. IJHCS. Prospective encoding reduces resumption lag.",
        ],
        "evidence_level": "STRONG",
        "moderating_factors": "Interruption duration and cognitive demand | External cues/notes for reinstatement | Working memory capacity | Task complexity/engagement depth",
        "evidence_notes": "MFG model has strong experimental support. Mark et al. provides ecological validity. Restart cost reflects genuine cognitive re-initialization, not just reluctance.",
    },
    "C6B-8": {
        "gate": "C6B", "name": "Shame-Driven Withdrawal",
        "definition": "Shame from perceived failure triggers withdrawal and avoidance instead of continued effort. Distinct from guilt, which motivates reparative action.",
        "mechanism": "Shame-driven avoidance response: perceived failure activates shame (global self-evaluation: 'I am bad'), triggering withdrawal from the goal domain. This is NOT guilt (specific behavior evaluation: 'I did something bad'), which motivates repair. Shame-proneness predicts destructive disengagement; guilt-proneness predicts constructive correction.",
        "decision_comp": "IF Perceived_Failure = TRUE AND Shame_Proneness = HIGH THEN Avoidance_Response_Activated; Withdrawal_From_Goal_Domain",
        "diagnostic_signals": "'I feel terrible about missing a day' | 'I beat myself up' | 'I withdraw instead of trying again' | 'I'm a failure'",
        "intervention": "Self-compassion training (primary), Error normalization, Cost management, Monitor progress.",
        "effect_size": "d = 0.3-0.5 (indirect evidence base)",
        "key_studies": [
            "Tangney & Dearing (2002). Shame and Guilt. Guilford. Shame = withdrawal/avoidance; guilt = reparative action. Critical distinction.",
            "Tangney et al. (2007). Moral Emotions and Moral Behavior. Annual Review of Psychology. Shame-proneness = destructive responses; guilt-proneness = constructive correction.",
            "Lickel et al. (2014). Shame and Motivation to Change. Emotion. Shame predicted desire for change but NOT constructive behavior.",
            "Neff (2003). Self-Compassion Scale. Self and Identity. Self-compassion as primary buffer against shame-driven withdrawal.",
            "Neff (2023). Self-Compassion: Theory, Method, Research, and Intervention. Annual Review of Psychology. Meta-analytic review; self-compassion training as intervention.",
        ],
        "evidence_level": "MODERATE",
        "moderating_factors": "Shame-proneness (primary moderator) | Self-compassion level | Perceived reparability | Cultural norms around failure | Public vs private failure context",
        "evidence_notes": "CORRECTION: Reframed from guilt to SHAME. The entry's own cited source (Tangney et al. 2007) clearly distinguishes them: shame → withdrawal; guilt → repair. The mechanism driving task abandonment is SHAME (global self-evaluation), not guilt (behavior-specific evaluation). Self-compassion training added as primary intervention per Neff (2023). Effect size adjusted to d=0.3-0.5 reflecting indirect evidence base.",
    },
}

# ============================================================
# DOCUMENT GENERATION
# ============================================================

def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h


def add_paragraph(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    return p


def color_for_evidence(level):
    if level == "STRONG":
        return RGBColor(0, 128, 0)
    elif level == "MODERATE":
        return RGBColor(180, 120, 0)
    elif level == "WEAK":
        return RGBColor(200, 0, 0)
    return RGBColor(128, 128, 128)


def add_evidence_summary_table(doc):
    """Add a summary table of all 56 entries with evidence levels."""
    gate_order = ["C0", "C1", "C2", "C3", "C4", "C5", "C6A", "C6B"]
    # Collect entries in gate order
    ordered = []
    for gid in gate_order:
        for eid, entry in ENTRIES.items():
            if entry["gate"] == gid:
                ordered.append((eid, entry))

    table = doc.add_table(rows=len(ordered) + 1, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['ID', 'Root Cause', 'Gate', 'Evidence Level']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = 'Calibri'

    for row_idx, (eid, entry) in enumerate(ordered, start=1):
        data = [eid, entry["name"], entry["gate"], entry["evidence_level"]]
        for col_idx, text in enumerate(data):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = ''
            run = cell.paragraphs[0].add_run(text)
            run.font.size = Pt(9)
            run.font.name = 'Calibri'
            if col_idx == 3:
                run.font.color.rgb = color_for_evidence(text)
                run.bold = True

    for row in table.rows:
        row.cells[0].width = Inches(0.7)
        row.cells[1].width = Inches(3.0)
        row.cells[2].width = Inches(0.7)
        row.cells[3].width = Inches(1.2)

    doc.add_paragraph()


def add_entry_table(doc, eid, entry):
    """Add a comprehensive entry table with all fields including new academic fields."""
    rows_data = [
        ('Definition', entry['definition']),
        ('Mechanism', entry['mechanism']),
        ('Decision Computation', entry['decision_comp']),
        ('Diagnostic Signals', entry['diagnostic_signals']),
        ('Intervention Axes', entry['intervention']),
        ('Effect Size', entry['effect_size']),
        ('Evidence Level', entry['evidence_level']),
        ('Moderating Factors', entry['moderating_factors']),
    ]

    table = doc.add_table(rows=len(rows_data) + 1, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    cell0 = table.rows[0].cells[0]
    cell0.text = ''
    run = cell0.paragraphs[0].add_run(f'{eid}: {entry["name"]}')
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    # Merge header cells
    cell0.merge(table.rows[0].cells[1])

    for row_idx, (label, value) in enumerate(rows_data, start=1):
        c0 = table.rows[row_idx].cells[0]
        c1 = table.rows[row_idx].cells[1]
        c0.text = ''
        c1.text = ''

        run0 = c0.paragraphs[0].add_run(label)
        run0.bold = True
        run0.font.size = Pt(10)
        run0.font.name = 'Calibri'

        run1 = c1.paragraphs[0].add_run(value)
        run1.font.size = Pt(10)
        run1.font.name = 'Calibri'

        if label == 'Evidence Level':
            run1.font.color.rgb = color_for_evidence(value)
            run1.bold = True

    for row in table.rows:
        row.cells[0].width = Inches(1.5)
        if len(row.cells) > 1:
            row.cells[1].width = Inches(5.5)

    doc.add_paragraph()

    # Key Studies (as formatted list below table)
    add_paragraph(doc, 'Key Studies:', bold=True)
    for study in entry['key_studies']:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(study)
        run.font.size = Pt(9)
        run.font.name = 'Calibri'

    # Evidence Notes
    if entry.get('evidence_notes'):
        p = doc.add_paragraph()
        run_label = p.add_run('Evidence Notes: ')
        run_label.bold = True
        run_label.font.size = Pt(9)
        run_label.font.name = 'Calibri'
        run_text = p.add_run(entry['evidence_notes'])
        run_text.font.size = Pt(9)
        run_text.font.name = 'Calibri'
        run_text.italic = True

    doc.add_paragraph()  # spacing


def add_neural_table(doc):
    """Add the neural systems validation table."""
    table = doc.add_table(rows=9, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['Gate', 'Neural Systems', 'Vocal Signature', 'Entry Count']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = 'Calibri'

    gate_order = ["C0", "C1", "C2", "C3", "C4", "C5", "C6A", "C6B"]
    counts = {"C0": 3, "C1": 8, "C2": 8, "C3": 6, "C4": 7, "C5": 8, "C6A": 8, "C6B": 8}

    for row_idx, gid in enumerate(gate_order, start=1):
        g = GATES[gid]
        data = [f'{gid}: {g["gate_name"]}', g["neural"], g["vocal_signature"], str(counts[gid])]
        for col_idx, text in enumerate(data):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = ''
            run = cell.paragraphs[0].add_run(text)
            run.font.size = Pt(9)
            run.font.name = 'Calibri'

    doc.add_paragraph()


def build_document():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # ---- TITLE PAGE ----
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_heading('VOLITIONAL CHAIN MODEL', level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)

    add_paragraph(doc, 'Complete Root Cause Database -- CORRECTED', bold=True)
    add_paragraph(doc, '56 Root Causes Across 8 Constraint Gates')
    add_paragraph(doc, '3 Pre-Volitional Screening Entries (C0) + 53 Volitional Root Causes (C1-C6B)')
    doc.add_paragraph()
    add_paragraph(doc, 'February 2026 -- CONFIDENTIAL')
    add_paragraph(doc, 'Nunez Medical Services Inc. / MindMetrics\u2122')
    doc.add_paragraph()
    add_paragraph(doc, 'CORRECTED VERSION -- Phase 2 Additions:', bold=True)
    add_paragraph(doc, '- Key Studies: 2-4 peer-reviewed citations per root cause entry')
    add_paragraph(doc, '- Evidence Level: STRONG / MODERATE / WEAK rating per entry')
    add_paragraph(doc, '- Moderating Factors: Variables that influence mechanism strength')
    add_paragraph(doc, '- Evidence Notes: Caveats, replication concerns, and methodological notes')
    doc.add_page_break()

    # ---- SECTION 1: DOCUMENT OVERVIEW ----
    add_heading(doc, '1. DOCUMENT OVERVIEW', level=1)
    add_paragraph(doc, 'This document contains the complete Root Cause Database for the Volitional Chain Model (VCM). It catalogs 56 distinct failure modes organized across 8 sequential constraint gates that map the perception-to-action pipeline for goal-directed behavior.')
    add_paragraph(doc, 'This corrected version adds academic validation fields to each root cause entry: Key Studies (peer-reviewed citations supporting the mechanism), Evidence Level (STRONG/MODERATE/WEAK rating), Moderating Factors (variables that influence mechanism strength), and Evidence Notes (caveats, replication concerns, and methodological qualifications).')
    add_paragraph(doc, 'Structural note: Gate C0 (Environmental Permeability) serves as a pre-volitional screening gate. It identifies non-volitional blockers -- biological, environmental, and structural constraints -- that must be ruled out before diagnosing failure within the volitional chain proper (C1-C6B).')

    # ---- SECTION 2: EVIDENCE LEVEL SUMMARY ----
    add_heading(doc, '2. EVIDENCE LEVEL SUMMARY', level=1)

    # Counts
    strong = sum(1 for e in ENTRIES.values() if e["evidence_level"] == "STRONG")
    moderate = sum(1 for e in ENTRIES.values() if e["evidence_level"] == "MODERATE")
    weak = sum(1 for e in ENTRIES.values() if e["evidence_level"] == "WEAK")

    add_paragraph(doc, f'STRONG: {strong} entries | MODERATE: {moderate} entries | WEAK: {weak} entries')
    doc.add_paragraph()
    add_evidence_summary_table(doc)

    # ---- SECTION 3: CRITICAL FLAGS ----
    add_heading(doc, '3. CRITICAL FLAGS', level=1)
    add_heading(doc, '3.1 C3-1: Self-Control Fatigue -- UPGRADED TO MODERATE (Feb 2026)', level=2)
    p = doc.add_paragraph()
    run = p.add_run('REVISED: ')
    run.bold = True
    run.font.color.rgb = RGBColor(0, 100, 180)
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    run2 = p.add_run('C3-1 was originally rated WEAK due to two large-scale replication failures of the ego depletion resource model. It has been upgraded to MODERATE based on new evidence and mechanism revision:')
    run2.font.size = Pt(11)
    run2.font.name = 'Calibri'

    add_paragraph(doc, 'Replication history:')
    add_paragraph(doc, '- Hagger et al. (2016): 23 labs, N=2,141, d=0.04 (null with brief 5-10 min protocols)')
    add_paragraph(doc, '- Vohs et al. (2021): 36 labs, N=3,531, d=0.06 (null with brief protocols)')
    add_paragraph(doc, '- Dang et al. (2025): 14 samples, N=2,078, d=0.31-0.35 (significant with intensive 30-40 min protocols, I-squared=0)')
    add_paragraph(doc, 'Resolution: The phenomenon is real under specific conditions (prolonged, intensive cognitive demands) but the mechanism is motivational shift, not resource depletion. Entry renamed from "Energy Depletion" to "Self-Control Fatigue" and mechanism revised to the process model (Inzlicht & Schmeichel, 2012; Berkman et al., 2017). Key moderators: task duration/intensity, beliefs about willpower, incentive level (increased incentives reverse the effect).')

    add_heading(doc, '3.2 Evidence Caveats by Rating', level=2)
    add_paragraph(doc, 'MODERATE entries (16 total, including upgraded C3-1): These have supporting evidence but with qualifications -- typically extrapolation from adjacent domains, small effect sizes in meta-analyses, or replication concerns. They are defensible in an academic context but should be presented with noted limitations. C3-1 (Self-Control Fatigue) was upgraded from WEAK after Dang et al. (2025) provided replication support with intensive protocols.')
    add_paragraph(doc, 'STRONG entries (40 total): These have multiple direct peer-reviewed studies supporting the specific mechanism described. They form the backbone of the VCM\'s empirical foundation.')
    add_paragraph(doc, 'WEAK entries: None remaining. C3-1 was the sole WEAK entry; upgraded to MODERATE following mechanism revision and new supporting evidence.')

    # ---- SECTION 4: GATE ARCHITECTURE ----
    add_heading(doc, '4. GATE ARCHITECTURE', level=1)
    add_neural_table(doc)

    # ---- SECTION 5: COMPLETE ROOT CAUSE ENTRIES ----
    add_heading(doc, '5. COMPLETE ROOT CAUSE ENTRIES', level=1)
    add_paragraph(doc, 'Each entry below contains: Definition, Mechanism, Decision Computation, Diagnostic Signals, Intervention Axes, Effect Size, Evidence Level, Moderating Factors, Key Studies, and Evidence Notes.')

    gate_order = ["C0", "C1", "C2", "C3", "C4", "C5", "C6A", "C6B"]

    for gid in gate_order:
        g = GATES[gid]
        add_heading(doc, f'{gid}: {g["gate_name"]}', level=2)
        add_paragraph(doc, f'Constraint Type: {g["constraint_type"]}')
        add_paragraph(doc, f'Diagnostic Question: {g["question"]}')
        add_paragraph(doc, f'Neural Systems: {g["neural"]}')
        add_paragraph(doc, f'Vocal Signature: {g["vocal_signature"]}')
        if g.get("framing_note"):
            add_paragraph(doc, f'Framing Note: {g["framing_note"]}', italic=True)
        doc.add_paragraph()

        # Add all entries for this gate
        for eid, entry in ENTRIES.items():
            if entry["gate"] == gid:
                add_heading(doc, f'{eid}: {entry["name"]}', level=3)
                add_entry_table(doc, eid, entry)

    # ---- SECTION 6: INTERVENTION AXES REFERENCE ----
    add_heading(doc, '6. INTERVENTION AXES REFERENCE', level=1)
    add_paragraph(doc, 'The VCM defines 37 distinct intervention axes. Each root cause maps to 1-3 primary intervention axes. The full intervention taxonomy is preserved from the original database. Key intervention categories:')

    interventions = [
        "Rebuild efficacy", "Increase specificity", "Mastery experiences",
        "Reframe reward", "Reduce cost", "Simulate context",
        "Increase value alignment", "Pre-commitment", "Mini-deadlines",
        "Increase perceived control", "Change structure", "Environmental redesign",
        "Reduce activation cost", "Micro-starts", "Reduce inertia",
        "Use implementation intentions", "Habit stacking", "Immediate cues",
        "Environmental control", "Add friction to alternatives", "Reduce friction",
        "Increase structure", "Reduce goal scope", "Change timing",
        "Monitor progress", "Raise abandonment cost", "Ulysses pact",
        "Effort pacing", "Energy pacing", "Cost management",
        "Progress feedback", "Reward persistence", "Increase intrinsic cues",
        "Alter environment",
    ]

    for idx, name in enumerate(interventions, 1):
        p = doc.add_paragraph()
        run = p.add_run(f'{idx}. {name}')
        run.font.size = Pt(10)
        run.font.name = 'Calibri'

    # ---- SECTION 7: METHODOLOGY NOTES ----
    add_heading(doc, '7. METHODOLOGY NOTES', level=1)
    add_paragraph(doc, 'Evidence Level Criteria:', bold=True)
    add_paragraph(doc, 'STRONG: Multiple peer-reviewed studies directly supporting the claimed mechanism. Meta-analyses available. Mechanism well-established in neuroscience or behavioral science literature. Replication record positive.')
    add_paragraph(doc, 'MODERATE: At least one direct study or multiple indirect studies. Mechanism plausible and consistent with established theory. May involve extrapolation from adjacent domains. Some replication concerns or limited sample sizes.')
    add_paragraph(doc, 'WEAK: Only indirect evidence, clinical observation, or theoretical extrapolation. Major replication failures for core mechanism. Requires revision before academic presentation.')
    doc.add_paragraph()
    add_paragraph(doc, 'Citation Selection Criteria:', bold=True)
    add_paragraph(doc, 'Citations were selected for: (1) direct relevance to the specific mechanism described, (2) publication in peer-reviewed journals, (3) citation impact and replication record, (4) recency where recent work supersedes or qualifies earlier findings.')
    doc.add_paragraph()
    add_paragraph(doc, 'Moderating Factor Selection:', bold=True)
    add_paragraph(doc, 'Moderating factors were identified from: (1) meta-analytic moderator analyses, (2) individual difference variables shown to affect mechanism strength, (3) contextual factors identified in empirical literature, (4) clinical observations documented in treatment research.')
    doc.add_paragraph()
    add_paragraph(doc, 'Cross-Cultural Validity:', bold=True)
    add_paragraph(doc, 'The majority of studies cited in this database were conducted with WEIRD (Western, Educated, Industrialized, Rich, Democratic) samples, predominantly North American and European university students. This is a known limitation of the behavioral science evidence base.')
    add_paragraph(doc, 'Implications for VCM: (1) Gate architecture (sequential constraint processing) is hypothesized to be universal, as it maps to conserved neural systems. However, which root causes are most prevalent likely varies by culture. (2) Entries grounded in basic neuroscience (basal ganglia gating, dopaminergic reward, amygdala threat response) are more likely cross-culturally valid than entries grounded in social-cognitive constructs (identity, values, social monitoring). (3) Effect sizes may differ across cultural contexts, particularly for interventions involving self-monitoring, social accountability, and value alignment. (4) Validation with non-WEIRD populations is recommended before clinical deployment in diverse settings.')

    # ---- SAVE ----
    output_path = r'c:\Users\randy\.claude\projects\kdenz\Claude code created\VCM_Root_Cause_Database_CORRECTED.docx'
    doc.save(output_path)
    print(f'Document saved to: {output_path}')

    # Verification
    total = len(ENTRIES)
    print(f'Total entries: {total}')
    strong = sum(1 for e in ENTRIES.values() if e["evidence_level"] == "STRONG")
    moderate = sum(1 for e in ENTRIES.values() if e["evidence_level"] == "MODERATE")
    weak = sum(1 for e in ENTRIES.values() if e["evidence_level"] == "WEAK")
    print(f'STRONG: {strong}, MODERATE: {moderate}, WEAK: {weak}')

    # Gate counts
    gate_counts = {}
    for e in ENTRIES.values():
        g = e["gate"]
        gate_counts[g] = gate_counts.get(g, 0) + 1
    for gid in ["C0", "C1", "C2", "C3", "C4", "C5", "C6A", "C6B"]:
        print(f'  {gid}: {gate_counts.get(gid, 0)} entries')

    # Count total studies
    total_studies = sum(len(e["key_studies"]) for e in ENTRIES.values())
    print(f'Total key studies cited: {total_studies}')


if __name__ == '__main__':
    build_document()
